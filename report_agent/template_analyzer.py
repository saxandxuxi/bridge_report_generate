# -*- coding: utf-8 -*-
"""模板识别：扫描 .docx 模板，找出需要替换的数据位置。

支持的占位符：
  {{stats.<指标>}}      统计值，例如 {{stats.temperature.max}}
  {{date.<字段>}}       报告期日期，例如 {{date.period_start}}
  {{chart.<图表ID>}}    图表位置（独占一行）
  {{rows.<数据集>}}     表格中可重复行的模板标记
  {{col.<字段>}}        可重复行内的列字段
  {{key|default:值}}    带默认值的占位符
  {{expr:表达式}}        算术表达式计算
  {{?condition}}...{{?}} 条件渲染块
"""

import re
from typing import Dict, List

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# 与 report_builder.MARKER_RE 保持一致
# key 用 [^:|}]+ 匹配除分隔符外的任意字符（含中文、#、（）等），避免特殊字符占位符"隐形"
MARKER_RE = re.compile(r"\{\{([^:|}]+)(?::([^}|]+))?(?:\|default:([^}]*))?\}\}")
COND_RE = re.compile(r"\{\{\?(.+?)\}\}(.*?)\{\{\?\}\}", re.DOTALL)
NUMBER_RE = re.compile(r"(?<![\w℃%])([+-]?\d+(?:\.\d+)?)(?![\w%])")


def iter_block_items(parent):
    """按文档顺序遍历正文中的段落和表格。"""
    from docx.oxml.ns import qn

    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        body = parent.element.body
    else:
        body = parent._tc
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def iter_paragraphs(doc: Document):
    """遍历正文、页眉、页脚以及表格单元格里的所有段落。"""
    for para in iter_block_items(doc):
        if isinstance(para, Paragraph):
            yield para
        elif isinstance(para, Table):
            for row in para.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def _classify(key: str) -> str:
    if key.startswith("stats."):
        return "stats"
    if key.startswith("date."):
        return "date"
    if key.startswith("chart."):
        return "chart"
    if key.startswith("rows."):
        return "rows"
    if key.startswith("col."):
        return "col"
    if key.startswith("expr:"):
        return "expr"
    return "unknown"


def analyze_template(template_path: str) -> Dict:
    """识别模板中所有动态内容，返回结构化结果。"""
    doc = Document(template_path)
    result = {
        "template": template_path,
        "placeholders": [],
        "charts": [],
        "rows": [],
        "candidate_numbers": [],
        "tables": [],
        "conditional_blocks": [],
    }

    for idx, para in enumerate(iter_paragraphs(doc)):
        text = para.text
        # 检测条件渲染块
        for cm in COND_RE.finditer(text):
            result["conditional_blocks"].append(
                {
                    "condition": cm.group(1).strip(),
                    "content_snippet": cm.group(2).strip()[:60],
                    "index": idx,
                }
            )

        for m in MARKER_RE.finditer(text):
            key = m.group(1)
            spec = m.group(2)
            default = m.group(3)
            entry = {
                "marker": m.group(0),
                "key": key,
                "type": _classify(key),
                "format": spec.strip() if spec else None,
                "default": default,
                "snippet": text.strip()[:80],
                "index": idx,
            }
            result["placeholders"].append(entry)
            if key.startswith("chart."):
                result["charts"].append(key.split(".", 1)[1])
            elif key.startswith("rows."):
                result["rows"].append(key.split(".", 1)[1])

        if not MARKER_RE.search(text) and not COND_RE.search(text):
            for m in NUMBER_RE.finditer(text):
                result["candidate_numbers"].append(
                    {"number": m.group(0), "snippet": text.strip()[:80], "index": idx}
                )

    for t_idx, table in enumerate(doc.tables):
        info = {
            "table_index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "markers": [],
        }
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for m in MARKER_RE.finditer(cell.text):
                    info["markers"].append(
                        {
                            "row": row_idx,
                            "col": col_idx,
                            "marker": m.group(0),
                            "key": m.group(1),
                        }
                    )
        result["tables"].append(info)

    # 去重图表/行数据集
    result["charts"] = list(dict.fromkeys(result["charts"]))
    result["rows"] = list(dict.fromkeys(result["rows"]))
    return result


def print_analysis(result: Dict) -> None:
    """以中文友好格式打印识别结果。"""
    print("=" * 60)
    print("模板动态内容识别结果")
    print("=" * 60)
    print(f"模板文件: {result['template']}")
    print(f"\n[占位符] 共 {len(result['placeholders'])} 处")
    for e in result["placeholders"]:
        fmt = f" (格式 {e['format']})" if e.get("format") else ""
        default = f" (默认: {e['default']})" if e.get("default") else ""
        print(f"  {e['marker']}  [{e['type']}]{fmt}{default}  段落#{e['index']}")

    # 条件渲染块
    cond_blocks = result.get("conditional_blocks", [])
    if cond_blocks:
        print(f"\n[条件渲染块] 共 {len(cond_blocks)} 处")
        for cb in cond_blocks:
            print(f"  {{?{cb['condition']}}}  <- \"{cb['content_snippet']}\"  段落#{cb['index']}")

    print(f"\n[图表] 共 {len(result['charts'])} 张")
    for cid in result["charts"]:
        print(f"  {{{{chart.{cid}}}}}")

    print(f"\n[可重复行] 共 {len(result['rows'])} 个数据集")
    for rid in result["rows"]:
        print(f"  {{{{rows.{rid}}}}}")

    print(f"\n[表格] 共 {len(result['tables'])} 张")
    for t in result["tables"]:
        marker_keys = ", ".join(m["key"] for m in t["markers"][:8])
        print(f"  表#{t['table_index']}: {t['rows']}行 x {t['cols']}列 -> {marker_keys}")

    cand = result.get("candidate_numbers", [])
    print(f"\n[候选数值]（正文中的普通数字，可能是需要动态替换的数据，共 {len(cand)} 处）")
    for c in cand[:20]:
        print(f"  {c['number']:>10}  <- \"{c['snippet']}\"")
    if len(cand) > 20:
        print(f"  ... 其余 {len(cand) - 20} 处略")
