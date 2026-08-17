# -*- coding: utf-8 -*-
"""成品报告解析识别：分析一份已完成的 DOCX / PDF 报告，判断

1. 哪些图片是动态数据图（建议替换为新图表），哪些是固定图（如 CAD 图、
   示意图、logo，保留）；
2. 哪些数字是动态统计值（建议替换为 {{stats.*}} 占位符），哪些是固定参数
   （如"桥长123米"这类设计常量，保留）。

识别流程（两轮筛选）：
  第一轮 — 关键词打分：基于上下文关键词、图题、文本长度进行启发式打分，
           给出 replace / keep / review 初步结论。
  第二轮 — LLM 完整性校验（llm_classifier.verify_and_complete，默认接入
           Qwen/DashScope）：
           1. 检查第一轮提取结果是否漏掉动态值（missed）
           2. 检查是否有静态值被误判为动态（wrong）
           3. 对待确认项（review）给出最终 replace / keep 判定
           4. 识别季度/时间表述并建议动态化（text_replacements）
           若全部正确，LLM 只回复"是"。
           LLM 不可用时自动降级为文本长度启发式。

输出统一的识别 JSON；对 DOCX 还可以生成一份"标注草稿"，把动态项直接改写为
统一标准的占位符，供人工确认后作为模板交给 run_agent.py 使用。

占位符统一标准：{{stats.<指标英文名>.<统计类型>}}
  指标英文名：temperature/humidity/deflection/displacement/rotation/strain/
             stress/cable_force/cable_clamp/vehicle_load/wind_load/wind_speed/
             earthquake_load/structural_temp/vehicle_count/settlement ...
  统计类型：max/min/avg/median/std/range

注意：这是基于关键词/上下文/图题/LLM 的启发式识别，结果需要人工复核；
不确定的项会标记为 review。
"""

import datetime as dt
import json
import logging
import os
import re
from functools import lru_cache
from typing import Dict, List, Optional, Tuple

log = logging.getLogger("report-agent.recognizer")

# 词库文件路径（项目根目录）
_GLOSSARY_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "keyword_glossary.json",
)

# 内置 fallback 信号（即使词库文件不存在也能运行）
_FALLBACK_DYNAMIC_TITLES = [
    "监测统计", "监测结果", "测值统计", "测值表", "动态数据",
    "变化监测", "时程监测", "频率分布", "分布统计", "变化统计",
]
_FALLBACK_STATIC_TITLES = [
    "对照表", "阈值表", "监测阈值", "作用阈值", "报警阈值", "控制值",
    "测点表", "测点布置", "测点布设", "测点一览", "监测内容表",
    "监测要素表", "监测项目表", "装备清单", "设备清单", "装备设置",
    "装备表", "设备表", "传感器类型", "分级对照", "等级对照",
    "风力分级", "风级对照", "风力对照", "规范表", "标准表",
    "技术参数", "车道分布", "车速统计", "车型统计", "交通量统计",
]

# 标题关键词 → metric 名
_TITLE_TO_METRIC = [
    # (匹配词, metric, 中文标签)
    ("结构温度", "structure_temperature", "结构温度"),
    ("温度", "temperature", "温度"),
    ("湿度", "humidity", "湿度"),
    ("应变", "strain", "应变"),
    ("挠度", "deflection", "挠度"),
    ("转角", "rotation", "转角"),
    ("倾角", "rotation", "倾角"),
    ("索夹", "cable_clamp", "索夹"),
    ("索力", "cable_force", "索力"),
    ("风速", "wind_speed", "风速"),
    ("风向", "wind_dir", "风向"),
    ("地震", "earthquake_load", "地震"),
    ("位移", "displacement", "位移"),
    ("变位", "displacement", "变位"),
    ("偏位", "displacement", "偏位"),
    ("振动", "vibration", "振动"),
    ("应力", "stress", "应力"),
    ("沉降", "settlement", "沉降"),
    ("车辆", "vehicle_count", "车辆"),
    ("交通", "vehicle_count", "交通"),
    ("荷载", "vehicle_count", "荷载"),
    ("车流", "vehicle_count", "车流"),
    ("裂缝", "crack", "裂缝"),
    ("桥面", "displacement", "桥面位移"),
    ("支座", "bearing_displacement", "支座位移"),
]

# 列头关键词 → stat 名（中文列头可能含"平均"，"最大"，"温差"等）
_HEADER_TO_STAT = [
    # (匹配词列表, stat, 优先级)
    (["平均温度", "平均值", "均值", "日均", "平均"], "avg", 0),
    (["最大温差", "最高温差"], "range", 0),
    (["最大", "最高"], "max", 0),
    (["最小", "最低"], "min", 0),
    (["绝对最大"], "abs_max", 0),
    (["均方根"], "rms", 0),
    (["应变差"], "range", 0),          # 最大应变差/με
    (["剔除温度最大"], "temp_rm_max", 0),   # 剔除温度效应后的应变最大值
    (["剔除温度最小"], "temp_rm_min", 0),   # 剔除温度效应后的应变最小值
    (["相关性系数", "相关系数"], "corr", 0), # 应变-温度相关性系数
    (["温差", "差值", "极差"], "range", 0),
    (["标准差", "std"], "std", 0),
    (["中位数", "median"], "median", 0),
    (["合计", "累计", "总和"], "sum", 0),
    (["数值"], "count", 0),       # 交通荷载“数值/辆”
    (["比例", "占比"], "ratio", 0),  # 交通荷载“比例/%”
]


@lru_cache(maxsize=1)
def load_glossary() -> Dict:
    """加载术语库（含动态/静态表格信号）。"""
    g: Dict = {
        "dynamic_signals": list(_FALLBACK_DYNAMIC_TITLES),
        "static_signals": list(_FALLBACK_STATIC_TITLES),
        "components": [],
        "bridge_names": [],
        "metric_keywords": [],
    }
    try:
        if os.path.isfile(_GLOSSARY_PATH):
            with open(_GLOSSARY_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            for k in g:
                if k in data:
                    g[k] = data[k]
    except (OSError, json.JSONDecodeError) as exc:
        log.warning("加载术语库失败，使用 fallback: %s", exc)
    return g


def slugify_label(label: str) -> str:
    """把行标签/列头清洗为可读占位符片段。

    例："岳阳索塔上游侧塔冠" → "岳阳索塔上游侧塔冠"（保留中文）
        "平均温度/℃" → "平均温度"（去除单位符号）
    """
    if not label:
        return ""
    out = label.strip()
    # 去除单位括号：平均温度/℃ → 平均温度
    # 注意不能用 \w（会匹配中文，把“跨中1/2截面”的 /2截面 误当单位剥掉），
    # 只剥 ASCII 字母数字 + ℃%°/ 组成的单位后缀
    out = re.sub(r"/[A-Za-z0-9℃%°/]+$", "", out).strip()
    # 去除纯符号
    # 位置里的 “/”（如 跨中1/2截面）必须保留，才能与名称对照/图库目录匹配
    out = re.sub(r"[\s\\()\[\]{}]+", "_", out)
    out = re.sub(r"_+", "_", out).strip("_")
    return out or label.strip()


def metric_from_title(title: str) -> str:
    """从表格标题推断 metric 名。"""
    if not title:
        return "generic"
    # 长度优先匹配（"塔冠环境温度监测统计" 中 "塔冠" 在前，但 "温度" 更具体）
    # 先按特定关键词，再按通用关键词
    for kw, metric, _ in _TITLE_TO_METRIC:
        if kw in title:
            return metric
    return "generic"


def stat_from_col_header(header: str) -> str:
    """从列头推断统计类型。"""
    if not header:
        return "value"
    # 通用“差”规则：最大湿度差/最大温度差/最大应变差/差值/极差 → range。
    # “标准差/标准偏差” 除外（→ std）。
    if "标准" in header and ("差" in header or "偏差" in header):
        return "std"
    if "差" in header:
        return "range"
    # 用最长匹配（"绝对最大值" 优先于 "最大"）
    candidates = []
    for keys, stat, _ in _HEADER_TO_STAT:
        for k in keys:
            if k in header:
                candidates.append((len(k), stat))
                break
    if candidates:
        candidates.sort(key=lambda x: -x[0])
        return candidates[0][1]
    # 启发式：纯数字列头可能是"测值"
    return "value"


def is_dynamic_table_title(title: str) -> bool:
    """判断表格标题是否属于动态监测表（值需替换）。"""
    g = load_glossary()
    t = title or ""
    # 静态信号胜出（更具体）
    for sig in g["static_signals"]:
        if sig in t:
            return False
    # 动态信号匹配
    for sig in g["dynamic_signals"]:
        if sig in t:
            return True
    return False




# 旧的 signal-driven 标题识别（保留供数字识别等使用）
TABLE_TITLE_SIGNALS_LEGACY = [
    "监测内容表", "监测项目表", "测点布设表", "测点布置表",
    "阈值表", "监测阈值", "作用监测阈值", "结构响应监测阈值",
    "报警阈值", "控制值表",
    "对照表", "分级表", "分级对照", "等级对照",
    "风力分级", "风级对照",
    "监测部位表", "设备清单", "传感器类型表",
    "规范表", "标准表",
]


from .llm_classifier import (
    LLMClassifier,
    SHORT_CONTEXT_THRESHOLD,
    LONG_CONTEXT_THRESHOLD,
    _text_length_heuristic,
)

log = logging.getLogger("report-agent.recognizer")

NUMBER_RE = re.compile(r"\d+(?:\.\d+)?")
MARKER_RE = re.compile(r"\{\{[^}]+\}\}")
SEQ_RE = re.compile(r"[图表第]\s*\d+")
DATE_RE = re.compile(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}")
KILO_POST_RE = re.compile(r"[Kk]\s*\d+\s*\+?\s*\d*")
CAPTION_RE = re.compile(r"^\s*(图|表|Figure|Fig\.?)\s*\d*")
# 图表文本占位：报告里 "xxx_time_series" / "xxx_histogram" 等行表示该处应插入图表
CHART_TEXT_RE = re.compile(
    r"_(time_series|histogram|curve|frequency|scatter|boxplot)$"
    r"|_[^_]+(时程曲线图|频率分布直方图|频率分布图|时程曲线|时间序列图)$"
    r"|_(时程曲线图|频率分布直方图|频率分布图|时程曲线|时间序列图|直方图)$"
    r"|_(时程曲线图|频率分布直方图|频率分布图|时程曲线|时间序列图|直方图)_\d+$"
)
# 多传感器合并行：位移*906*907*908*时程曲线_x1 / 位移*906_907_908*频率分布_3x1
CHART_MULTI_RE = re.compile(
    r"^\s*位移[_＊*]([\d_＊*]+)[_＊*](时程曲线|频率分布)[_＊*]?(\d*)x(\d+)\s*$"
)
# “编号(特征)_图型”行（309(xJsd)_时程曲线）：转为位置展开的占位行，
# 不再生成 chart_sensor_<编号>_<图型>（位置/方位由表格补全）
CHART_LINE_RE = re.compile(
    r"^\s*(\d{1,5})\s*\(([^()]+)\)\s*_(时程曲线|频率分布直方图|时间序列图|时间序列|直方图|时程|曲线)\s*$"
)
# 特征编码 -> 指标（xJsd/yJsd/zJsd 是振动三向；xJd/yJd 是倾角）
FEATURE_METRIC = {
    "xjsd": "vibration", "yjsd": "vibration", "zjsd": "vibration",
    "xjd": "rotation", "yjd": "rotation",
    "temp": "temperature", "rh": "humidity",
    "rsg": "strain", "nd": "deflection", "sl": "cable_force",
    "wy": "displacement", "lf": "crack",
}

# 纯图题识别：图题末尾含"图"且包含动态图表关键词
# 例：交通荷载监测_车辆累计通过数量统计图、梁端转角测点变化时程曲线图
BARE_CHART_CAPTION_RE = re.compile(
    r"^(?P<prefix>.+?)(?P<kind>(曲线|时程|直方|柱状|分布|趋势|占比|变化|统计|频谱|频次|相关|车流|载荷|荷载)图?)$"
)
# 表格 Excel 单元格引用占位符：H(行,列) / M(行,列) / J(行,列)
# 字母前缀表示表格（指标），行号从 2 开始（标题行算 1），列号从 2 开始（标签列算 1）
# 例：H(10,2) = 索夹表 87-L 索夹 1 的平均值列
CELL_REF_RE = re.compile(r"([A-Z])\((\d+)\s*,\s*(\d+)\)")


def _metric_from_chart_text(text: str) -> str:
    """从图表文本/图题推断指标名（长关键词优先，避免“结构温度”误判为“温度”）。"""
    t = text or ""
    ordered = [
        ("结构温度", "structure_temperature"),
        ("环境温度", "temperature"),
        ("环境湿度", "humidity"),
        ("梁端倾角", "rotation"),
        ("倾角", "rotation"),
        ("挠度", "deflection"),
        ("应变", "strain"),
        ("转角", "rotation"),
        ("索夹", "cable_clamp"),
        ("索力", "cable_force"),
        ("位移", "displacement"),
        ("变位", "displacement"),
        ("偏位", "displacement"),
        ("裂缝", "crack"),
        ("风速", "wind_speed"),
        ("风向", "wind_dir"),
        ("振动", "vibration"),
        ("应力", "stress"),
        ("沉降", "settlement"),
        ("车辆", "vehicle_count"),
        ("交通", "vehicle_count"),
        ("车流", "vehicle_count"),
        ("荷载", "load"),
        ("温度", "temperature"),
        ("湿度", "humidity"),
    ]
    for kw, metric in ordered:
        if kw in t:
            return metric
    return "chart"

# 动态统计关键词（出现则倾向于替换）
DYNAMIC_WORDS = [
    "最高", "最低", "最大", "最小", "平均", "均值", "中位数", "标准差", "极差", "方差",
    "占比", "比例", "增长率", "同比", "环比", "上涨", "下降", "超标", "超过", "低于", "高于",
    "累计", "合计", "总计", "共", "趋势", "温度", "湿度", "流量", "水位", "浓度", "沉降",
    "位移", "应变", "应力", "裂缝", "挠度", "温差", "雨量", "风速", "合格率", "通过率",
    "偏差", "波动", "监测", "统计", "指标", "产值", "营收", "销售额", "产量", "利润",
    "天数", "小时", "频次", "次数", "超标日",
    # 桥梁健康监测专用
    "索力", "索夹", "转角", "变位", "振动", "车辆", "地震", "加速度",
    "差值", "最大值", "最小值", "平均风速", "最大差值", "最小差值", "通过数量",
]

# 固定参数关键词（出现则倾向于保留）
# 注意：只放"真正的静态参数指标"，位置词（塔冠/上游侧/测点等）不放——
# 它们在动态上下文里也大量出现，会错误地压低动态值的分数。
STATIC_WORDS = [
    "桥长", "桥宽", "跨径", "孔径", "桩长", "桩径", "桩号", "设计", "标准", "规范",
    "编号", "图号", "图纸", "坐标", "高程", "比例尺", "型号", "规格", "地址", "电话",
    "传真", "合同", "标段", "里程", "日期", "年份", "版本", "负责人", "单位", "项目名称",
    "车道", "限速", "荷载等级", "强度等级", "材料", "材质",
    "面积", "容积", "功率", "电压", "频率", "重量", "单价", "金额", "造价", "预算", "投资",
    "预应力", "钢束", "锚具", "斜拉索",
    # 检查/检测类静态叙述：螺栓抽取数、扭力合格范围、节点板编号等
    "螺栓", "扭力", "合格范围", "节点板", "大桩号", "小桩号", "主桁架",
    "抽取", "检查", "复检", "抽检",
    # 签字表 / 机构信息
    "检测有限公司", "资质", "证书", "职称", "签字", "姓名", "岗位", "职业资格",
    # 规范引用（后面跟编号）
    "交办公路", "指南", "规程", "规范编号",
]

# 图片图题关键词
IMAGE_REPLACE_WORDS = [
    "趋势", "曲线", "直方", "柱状", "折线", "饼图", "分布", "对比", "统计", "变化",
    "走势", "箱线", "散点", "相关性", "时序", "频率", "占比", "监测曲线",
]
IMAGE_KEEP_WORDS = [
    "示意", "图纸", "CAD", "设计图", "平面图", "立面图", "剖面图", "效果图", "照片",
    "现场", "logo", "标志", "流程图", "架构", "拓扑", "布置图", "总图", "地形图",
    "地质图", "断面图", "构造图", "配筋图",
]

METRIC_WORDS = {
    "温度": "temperature", "湿度": "humidity", "流量": "flow", "水位": "water_level",
    "浓度": "concentration", "沉降": "settlement", "位移": "displacement",
    "应变": "strain", "应力": "stress", "裂缝": "crack", "挠度": "deflection",
    "雨量": "rainfall", "风速": "wind_speed",
    # 桥梁健康监测专用（与 llm_classifier.METRIC_STANDARD 保持一致）
    "转角": "rotation", "倾角": "rotation", "索力": "cable_force", "索夹": "cable_clamp",
    "变位": "displacement", "支座位移": "bearing_displacement",
    "振动": "vibration", "地震": "earthquake_load", "车辆": "vehicle_count",
    "荷载": "load", "结构温度": "structure_temperature",
}
STAT_WORDS = {
    "最高": "max", "最大": "max", "最低": "min", "最底": "min", "最小": "min",
    "平均": "avg", "均值": "avg", "中位数": "median", "标准差": "std", "极差": "range",
    # 长词优先（suggest_placeholder 会按关键词长度优先），“最大差值”应映射 range
    "最大差值": "range", "最小差值": "range", "最大变化": "range",
    "绝对最大": "abs_max",
}


def _window(text: str, start: int, end: int, width: int = 20) -> str:
    return text[max(0, start - width):end + width]


def suggest_placeholder(before: str, after: str):
    """根据数字前后文本中距离最近的关键词建议占位符。

    优先取数字“前面”的关键词（指标词通常写在数字前，如“最高温度26.0℃”）；
    仅当前面不足时再参考后面。stat 词距离阈值 6 字、metric 词 10 字。
    """
    combined = {**METRIC_WORDS, **STAT_WORDS}
    best_metric = best_stat = None
    best_md = best_sd = None
    best_stat_kw = best_metric_kw = ""

    def consider(text, offset_penalty, from_end=False):
        nonlocal best_metric, best_stat, best_md, best_sd, best_stat_kw, best_metric_kw
        # 长关键词优先：同一位置命中时，“最大差值”优先于“最大”
        for kw, v in sorted(combined.items(), key=lambda kv: -len(kv[0])):
            idx = text.rfind(kw) if from_end else text.find(kw)
            if idx == -1:
                continue
            d = (len(text) - idx if from_end else idx) + offset_penalty
            if v in METRIC_WORDS.values() and d <= 20 and (
                    best_md is None or d < best_md
                    or (d == best_md and len(kw) > len(best_metric_kw))):
                best_metric, best_md, best_metric_kw = v, d, kw
            elif v in STAT_WORDS.values() and d <= 8 and (
                    best_sd is None or d < best_sd
                    or (d == best_sd and len(kw) > len(best_stat_kw))):
                best_stat, best_sd, best_stat_kw = v, d, kw

    consider(before, 0, from_end=True)
    # “结构温度”比“温度”更具体，覆盖短词“温度”命中，
    # 避免“结构温度…最高温度”被抢成 stats.temperature.*
    if "结构温度" in before:
        best_metric = "structure_temperature"
        best_md = 0
    if "支座位移" in before:
        best_metric = "bearing_displacement"
        best_md = 0
    if best_metric is None or best_stat is None:
        consider(after, 30)  # 数字之后的词仅作补充，且带较大距离惩罚
    if best_metric and best_stat:
        return f"stats.{best_metric}.{best_stat}"
    return None


def classify_number(value: str, context: str, before: str = "", after: str = "") -> dict:
    """判断一个数字是否应替换。

    返回 {"verdict": replace|keep|review, "confidence": 0-1,
          "reasons": [...], "placeholder": 建议占位符}
    """
    score = 0.5
    reasons = []

    hits_dyn = [w for w in DYNAMIC_WORDS if w in context]
    hits_static = [w for w in STATIC_WORDS if w in context]
    if hits_dyn:
        score += 0.28
        reasons.append("上下文含动态关键词：" + "、".join(hits_dyn[:3]))
    if hits_static:
        score -= 0.32
        reasons.append("上下文含固定参数关键词：" + "、".join(hits_static[:3]))

    # 计算用于序号检测的干净上下文
    ctx_stripped = context.strip()

    # 数字后面紧跟单位
    if context and context[-1] in "℃%":
        score += 0.10
        reasons.append("带测量单位")

    if SEQ_RE.search(context):
        score -= 0.40
        reasons.append("图/表/第 序号")
    if DATE_RE.search(context):
        score -= 0.35
        reasons.append("日期格式")
    if re.fullmatch(r"\d{4}", value) and after.startswith("年"):
        score -= 0.55
        reasons.append("年份（后跟年字）")
    elif re.fullmatch(r"\d{4}", value) and 1950 <= int(value) <= 2100:
        score -= 0.30
        reasons.append("疑似年份")
    # 月份：年份后的 "01月-03月" 中的 01/03
    if re.fullmatch(r"\d{1,2}", value) and after.startswith("月") and re.search(r"年\d*$", before):
        score -= 0.55
        reasons.append("月份（年份后）")
    if KILO_POST_RE.search(context):
        score -= 0.50
        reasons.append("桩号/里程")

    # 测点布设数量（静态设计信息）："共布置11个测点" / "布设1个测点"
    if "共布置" in before or "共布设" in before or before.endswith(("布置", "布设")):
        score -= 0.50
        reasons.append("测点布设数量")

    # 测点编号：数字前后紧跟 "测点" / "测点N" / "N号测点"。
    # 只作用于整数（测点编号是 1/2/3…）；"测点的最大值为39.96" 这类小数是测量值，
    # 不能被误判成测点编号。
    if re.fullmatch(r"\d+", value) and (
            re.search(r"测\s*点\s*$", before) or re.search(r"测\s*点\s*$", ctx_stripped)):
        score -= 0.55
        reasons.append("测点编号（前）")
    elif re.fullmatch(r"\d+", value) and re.search(r"^\s*测\s*点", after):
        score -= 0.55
        reasons.append("测点编号（后）")
    if re.fullmatch(r"\d+", value) and (
            re.search(r"^\s*号\s*测\s*点", after) or re.search(r"号\s*测\s*点$", before)):
        score -= 0.55
        reasons.append("测点编号")

    # 车辆总数：数字后面紧跟“辆”（如 “…车辆总数为378993辆”），
    # 与车道号（“车道1、车道2”）区分，保证结论段三个车辆数都判为动态。
    if re.search(r"^\s*辆", after):
        score += 0.35
        reasons.append("车辆总数（后跟辆）")

    # 索夹/吊索编号："87-L" / "87-R" / "88-L" / "88-R" 等字母结尾
    if re.search(r"[A-Za-z]$", after) and re.fullmatch(r"\d+", value):
        score -= 0.60
        reasons.append("索夹/吊索编号")
    if re.search(r"^\s*-\s*[A-Za-z]$", after) and re.fullmatch(r"\d+", value):
        score -= 0.60
        reasons.append("索夹/吊索编号")
    # 字母前缀编号（如 RE24-RE25 节点板、K814+091 桩号）：数字紧跟在字母后
    if re.search(r"[A-Za-z]$", before) and re.fullmatch(r"\d+", value):
        score -= 0.60
        reasons.append("编号（字母前缀）")

    # 表格标题信号：context 含 "监测内容表" / "阈值表" / "对照表" / "分级表" 等
    table_title_signals = [
        "监测内容表", "监测内容", "监测项目表", "测点布设",
        "阈值表", "监测阈值", "作用监测阈值", "结构响应监测阈值", "报警阈值",
        "对照表", "分级表", "分级对照",
        "监测部位", "监测要素", "传感器类型", "设备清单",
    ]
    if any(sig in context for sig in table_title_signals):
        score -= 0.45
        reasons.append("表格标题属于规范/配置/阈值表")

    # 章节/列表序号（只作用于"数字本身"是序号的情况，不影响同段落其他数字）
    if not before.strip() and re.match(r"^\d+[\.、．]", ctx_stripped):
        score -= 0.60
        reasons.append("章节号（段落开头）")
    elif before.endswith(("（", "(", "第")) and after.startswith(("）", ")", "项")):
        score -= 0.60
        reasons.append("列表序号")
    elif re.search(r"\d\.\d*\.\s*$", before) or re.search(r"\.\d+\s*$", before):
        score -= 0.60
        reasons.append("多级章节号")

    # 文本长度影响：极短上下文倾向 keep（标签/序号），长描述倾向 replace（数据叙述）
    ctx_len = len(ctx_stripped)
    if ctx_len < SHORT_CONTEXT_THRESHOLD:
        score -= 0.08
        reasons.append(f"极短上下文（{ctx_len}字），疑似标签/序号")
    elif ctx_len > LONG_CONTEXT_THRESHOLD:
        score += 0.05
        reasons.append(f"长描述性上下文（{ctx_len}字），疑似数据叙述")

    score = max(0.05, min(0.95, score))
    if score >= 0.68:
        verdict = "replace"
    elif score <= 0.38:
        verdict = "keep"
    else:
        verdict = "review"

    placeholder = None
    if verdict == "replace":
        placeholder = suggest_placeholder(before, after)

    return {
        "verdict": verdict,
        "confidence": round(score, 2),
        "reasons": reasons,
        "placeholder": placeholder,
    }


def classify_image(caption: str, w_in: float = 0, h_in: float = 0,
                   in_header_footer: bool = False) -> dict:
    """判断一张图片是否应替换。"""
    if in_header_footer:
        return {"verdict": "keep", "confidence": 0.92,
                "reasons": ["位于页眉页脚，通常是 logo/装饰"], "placeholder": None}

    caption = (caption or "").strip()
    if not caption:
        if w_in and h_in and max(w_in, h_in) < 1.5:
            return {"verdict": "keep", "confidence": 0.75,
                    "reasons": ["无图题且尺寸小，疑似 logo/图标"], "placeholder": None}
        return {"verdict": "review", "confidence": 0.45,
                "reasons": ["无图题，需人工确认"], "placeholder": None}

    score = 0.5
    reasons = []
    if CAPTION_RE.match(caption):
        reasons.append("带图题")
        hits_r = [w for w in IMAGE_REPLACE_WORDS if w in caption]
        hits_k = [w for w in IMAGE_KEEP_WORDS if w in caption]
        if hits_r and not hits_k:
            score += 0.35
            reasons.append("图题含动态图表关键词：" + "、".join(hits_r[:3]))
        elif hits_k:
            score -= 0.42
            reasons.append("图题含固定图关键词：" + "、".join(hits_k[:3]))
        else:
            score -= 0.10
            reasons.append("图题无明确关键词")
    else:
        score -= 0.20
        reasons.append("无规范图题（图1/Figure 1 格式）")

    # 文本长度影响：极短图题可能是编号，长图题更可能含描述性关键词
    cap_len = len(caption.strip())
    if cap_len < SHORT_CONTEXT_THRESHOLD and score > 0.40:
        score -= 0.05
        reasons.append(f"极短图题（{cap_len}字）")
    elif cap_len > LONG_CONTEXT_THRESHOLD:
        score += 0.04
        reasons.append(f"长描述性图题（{cap_len}字）")

    score = max(0.05, min(0.95, score))
    if score >= 0.68:
        verdict = "replace"
    elif score <= 0.40:
        verdict = "keep"
    else:
        verdict = "review"
    return {"verdict": verdict, "confidence": round(score, 2),
            "reasons": reasons, "placeholder": None}


def _extract_numbers(text: str, base: dict, header_context: str = "") -> list:
    """从一段文本里提取所有数字并分类。"""
    out = []
    markers = list(MARKER_RE.finditer(text))
    for m in NUMBER_RE.finditer(text):
        if any(mo.start() <= m.start() and mo.end() >= m.end() for mo in markers):
            continue  # 跳过占位符内部
        value = m.group(0)
        ctx = _window(text, m.start(), m.end())
        before = text[max(0, m.start() - 20):m.start()]
        after = text[m.end():m.end() + 4]
        if header_context:
            ctx = header_context + " | " + ctx
            if not any(k in before for k in {**METRIC_WORDS, **STAT_WORDS}):
                before = header_context + " " + before
        info = classify_number(value, ctx, before, after)
        info.update(
            {
                "value": value,
                "context": ctx,          # 数字附近的真实上下文（含表格标题）
                "snippet": text.strip()[:80],
                "position": m.start(),
                **base,
            }
        )
        out.append(info)
    return out


# 原文占位标记：如 [A-MAX]、[A-MAX-LOC]、B-MIN、G-MAX-X 等。
# 部分成品报告用这类标记代替真实数值（由报告生成器占位），识别时应视为动态值。
TAG_TOKEN_RE = re.compile(r"[A-Z][A-Z0-9]*(?:-[A-Z0-9]+)+")
TAG_RE = re.compile(
    r"\[(" + TAG_TOKEN_RE.pattern + r")\]"
    r"|(?<![A-Za-z0-9\]])" + TAG_TOKEN_RE.pattern + r"(?![A-Za-z0-9\[])",
)


def _nearest_word(text: str, pos: int, win: int, words: dict, skip=()) -> Optional[str]:
    """在 text[pos-win:pos] 里找最近的指标/统计词，长词优先。

    短词是长词的子串且落在长词内部时（如 “最大” 在 “绝对最大/最大差值” 内、
    “温度” 在 “结构温度” 内），必须让长词胜出，否则会把 绝对最大/最大差值
    错判成 max、把 结构温度 错判成 环境温度。
    skip: 需要排除的词类值（如 “荷载/车辆” 太泛，不适合推断监测指标）。
    """
    ordered = sorted(words.items(), key=lambda kv: -len(kv[0]))
    best = None
    for kw, v in ordered:
        if v in skip:
            continue
        i = text.rfind(kw, max(0, pos - win), pos)
        if i == -1:
            continue
        # 该位置被更长关键词覆盖（长词起点 <= 短词起点 < 长词终点）→ 跳过短词
        covered = False
        for lkw, lv in ordered:
            if len(lkw) <= len(kw) or lv in skip:
                continue
            j = text.rfind(lkw, max(0, pos - win), pos)
            if j != -1 and j <= i < j + len(lkw):
                covered = True
                break
        if covered:
            continue
        d = pos - i
        if best is None or d < best[0]:
            best = (d, v)
    return best[1] if best else None


# 故障/异常位置判断句：位置+测点 一直为0/数据异常/传感器故障导致 等。
# 例：“上游炎陵侧边跨跨中截面顶板测点2、汝城侧边跨跨中截面3个测点一直为0℃，
#     为传感器故障导致。其他测点均在正常范围内，未对桥梁结构造成影响。”
SUMMARY_FAULT_RE = re.compile(
    r"[^。；]*?(?:"
    r"(?:一直|始终|持续|长期|连续)为\s*0(?:\s*[℃%％]?)|"
    r"由设备设置错误引起|为传感器故障导致|传感器故障|监测数据异常|数据出现异常"
    r")[^。]*。"
    r"(?:\s*其他[^。]*?(?:正常|无异常|未对[^。]*影响)[^。]*。)?"
)

# 判断句需含位置/测点词，避免误伤普通数据描述句
SUMMARY_POS_RE = re.compile(
    r"测点|监测部位|位置|截面|索塔|塔冠|塔|墩|跨|梁端|锚固区|箱梁|桥面|"
    r"主缆|加劲梁|索鞍|塔底|塔顶"
)


def _metrics_in_text(text: str, skip=("load", "vehicle_count")) -> list:
    """按出现顺序返回文本中的指标键（长词优先；结构温度>温度、支座位移>位移）。
    避免“结构温度最高为…；温度最低为…”里裸“温度”被误判成环境温度。"""
    kws = sorted(((k, v) for k, v in METRIC_WORDS.items() if v not in skip),
                 key=lambda kv: -len(kv[0]))
    pat = re.compile("|".join(re.escape(k) for k, _ in kws))
    seen = []
    for m in pat.finditer(str(text)):
        v = next((v for k, v in kws if k == m.group(0)), None)
        if v and v not in seen:
            seen.append(v)
    if "structure_temperature" in seen and "temperature" in seen:
        seen.remove("temperature")
    if "bearing_displacement" in seen and "displacement" in seen:
        seen.remove("displacement")
    return seen


def detect_summary_placeholders(texts_all: list) -> dict:
    """识别“…位置一直为0/数据异常…为传感器故障导致…”总结句，
    按该句附近（前 60 字窗口）的特征词生成 {{summary.<metric>}} 占位符。
    一段同时总结多个特征（如 挠度+结构温度）时，按特征各生成一个占位符：
    第一个原位替换整句，其余紧跟其后插入。
    返回 {段落索引: [(start, end, marker), ...]}（坐标基于原文）。"""
    out = {}
    for i, t in enumerate(texts_all or []):
        t = str(t)
        if not t:
            continue
        for m in SUMMARY_FAULT_RE.finditer(t):
            seg = m.group(0)
            if not SUMMARY_POS_RE.search(seg):
                continue
            window = t[max(0, m.start() - 120):m.end()]
            metrics = _metrics_in_text(window)
            if not metrics:
                metrics = _metrics_in_text(t)
            if not metrics:
                continue
            spans = []
            for k, metric in enumerate(metrics):
                if k == 0:
                    spans.append((m.start(), m.end(),
                                  f"{{{{summary.{metric}}}}}"))
                else:
                    # 其余特征占位符紧跟整句之后插入
                    spans.append((m.end(), m.end(),
                                  f"{{{{summary.{metric}}}}}"))
            out.setdefault(i, []).extend(spans)
    return out


def _extract_tags(text: str, base: dict) -> list:
    """提取 [A-MAX] / A-MAX-LOC 等原文占位标记。"""
    out = []
    for m in TAG_RE.finditer(text):
        token = m.group(1) or m.group(0)
        # 规范编号（CECS3-2012 / JTG3362-2018 / T1037-2016 等）以 4 位年份结尾，
        # 是固定参数，不是动态占位标记；上下文含 规范/标准/规程/编号 时同样保留
        before_ctx = text[max(0, m.start() - 30):m.start()]
        if re.search(r"-\d{4}$", token) or any(
                w in before_ctx for w in ("规范", "标准", "规程", "编号", "JTG", "GB", "CECS")):
            continue
        # 真正的占位标记含 MAX/MIN/LOC 等统计词；RE24-RE25、K814+091 这类
        # 节点板/桩号编号不该被当作动态标记
        if not re.search(r"(MAX|MIN|LOC|ABS|AVG|STD|RANGE|MEDIAN)", token):
            continue
        before = text[max(0, m.start() - 20):m.start()]
        after = text[m.end():m.end() + 4]
        out.append({
            "token": token,
            "value": token,
            "position": m.start(),
            "end": m.end(),
            "before": before,
            "after": after,
            "verdict": "replace",
            "confidence": 0.9,
            "reasons": ["原文占位标记（如 A-MAX），应动态替换"],
            "placeholder": None,
            **base,
        })
    return out


def _classify_tag(tag: dict, texts: list) -> None:
    """根据上下文把 [A-MAX] 这类标记映射为 stats.* 占位符。

    - 值标记（A-MAX）→ {{stats.<指标>.<统计>}}（重算）
    - 位置标记（A-MAX-LOC）→ {{stats.<指标>.<统计>.loc}}（对应测点位置，运行时取
      最值传感器的监测部位）；推断不到时退回 {{data.N}} 原文回填。
    """
    para_idx = tag.get("paragraph")
    pos = tag.get("position")
    if not isinstance(para_idx, int) or not isinstance(pos, int):
        return
    t = str(texts[para_idx]) if 0 <= para_idx < len(texts) else ""
    if not t:
        return
    token = str(tag.get("token", ""))
    is_loc = token.endswith("-LOC")
    # 统计词扫描窗口：位置标记（…对应测点位置为[A-MAX-LOC]）里统计词离得较远
    stat_win = 40 if is_loc else 12
    # 只在本句（最后句号之后）里找指标词；分号后子句若自带指标词
    # （“…；地震监测数据的绝对最大值为[D-MAX]”）则按子句归属，
    # 否则沿用整句（“…索力最大值为…；最小值为”）
    sent_start = max(t.rfind("。", 0, pos), t.rfind("！", 0, pos),
                     t.rfind("？", 0, pos)) + 1
    sentence = t[sent_start:pos]
    semi = sentence.rfind("；")
    if semi != -1:
        clause = sentence[semi + 1:]
        clause_metric = _nearest_word(clause, len(clause), 60, METRIC_WORDS,
                                      skip={"load", "vehicle_count"})
        if clause_metric:
            sentence = clause
    best_metric = _nearest_word(sentence, len(sentence), 60, METRIC_WORDS,
                                skip={"load", "vehicle_count"})
    # F-MAX / F-MIN：剔除温度效应后的应变最大/最小差值（残差），
    # 与普通 E-MAX 不同，映射到 temp_rm_range / temp_rm_min；
    # 对应位置标记取“差值”本身的最值位置（temp_rm_range.loc），
    # 而不是剔除温度残差最大值的传感器位置（temp_rm_max.loc）。
    if re.search(r"^F-(MAX|MIN)(-LOC)?$", token):
        if best_metric == "strain":
            if token == "F-MAX-LOC":
                tag["placeholder"] = "{{stats.strain.temp_rm_range.loc}}"
            elif token == "F-MAX":
                tag["placeholder"] = "{{stats.strain.temp_rm_range}}"
            elif token == "F-MIN-LOC":
                tag["placeholder"] = "{{stats.strain.temp_rm_min.loc}}"
            else:
                tag["placeholder"] = "{{stats.strain.temp_rm_min}}"
            return
    best_stat = _nearest_word(sentence, len(sentence), stat_win, STAT_WORDS)
    TAG_METRICS = {
        "temperature", "humidity", "structure_temperature", "wind_speed",
        "strain", "deflection", "rotation", "displacement", "bearing_displacement",
        "cable_force", "crack", "earthquake_load", "vibration",
    }
    if best_metric in TAG_METRICS and best_stat:
        # 方向后缀：如 [G-MAX-X] / [G-MAX-Y] / [G-MAX-Z]（GNSS 三方向）、
        # [I-MAX-X] 等，生成 direction 后缀，运行时按方向取对应特征
        direction = ""
        dm = re.search(r"-(X|Y|Z)(?:-LOC)?$", token)
        if dm:
            direction = "_" + dm.group(1).lower()
        suffix = ".loc" if is_loc else ""
        tag["placeholder"] = (
            f"{{{{stats.{best_metric}{direction}.{best_stat}{suffix}}}}}")


# ---------------------------------------------------------------------------
# DOCX 解析
# ---------------------------------------------------------------------------

def _images_in_paragraph(p_elm) -> list:
    from docx.oxml.ns import qn

    imgs = []
    nodes = p_elm.findall(".//" + qn("wp:inline")) + p_elm.findall(".//" + qn("wp:anchor"))
    for node in nodes:
        ext = node.find(qn("wp:extent"))
        cx = cy = 0
        if ext is not None:
            try:
                cx, cy = int(ext.get("cx", 0)), int(ext.get("cy", 0))
            except (TypeError, ValueError):
                pass
        blip = node.find(".//" + qn("a:blip"))
        rid = blip.get(qn("r:embed")) if blip is not None else None
        imgs.append({"rid": rid, "w_emu": cx, "h_emu": cy})
    return imgs


def _is_table_title(text: str) -> bool:
    """判断段落是否为表格标题。"""
    t = text.strip()
    if not t:
        return False
    # 形如 "表1 xx表" / "表2.3 监测阈值表" / "Xxx对照表"
    if re.match(r"^表\s*\d+", t):
        return True
    title_keywords = [
        "监测内容表", "监测项目表", "测点布设表", "测点布置表",
        "阈值表", "监测阈值", "作用监测阈值", "结构响应监测阈值",
        "报警阈值", "控制值表",
        "对照表", "分级表", "分级对照", "等级对照",
        "风力分级", "风级对照",
        "监测部位表", "设备清单", "传感器类型表",
        "规范表", "标准表",
    ]
    return any(kw in t for kw in title_keywords)


def parse_docx(path: str) -> dict:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from .report_builder import iter_block_items

    doc = Document(path)
    paragraphs = []  # 所有段落文本（含表格内、页眉页脚）
    numbers = []
    tags = []
    images = []
    # 表格布局：每个 cell_ref 段落 → (title, row_idx, col_idx, row_label, col_header)
    # 供后续 cell_ref 处理使用，避免只用 letter+position
    table_layout_for_para: Dict[int, Dict] = {}
    table_layout_for_para_pos: Dict[Tuple[int, int], Dict] = {}

    def handle_paragraph(p, header_footer=False, header_context="", table_title=""):
        idx = len(paragraphs)
        text = p.text
        paragraphs.append(text)
        imgs = _images_in_paragraph(p._p)
        if text.strip():
            ctx_parts = [x for x in (table_title, header_context) if x]
            full_ctx = " | ".join(ctx_parts) if ctx_parts else ""
            numbers.extend(
                _extract_numbers(
                    text, {"page": 1, "paragraph": idx, "table": None,
                           "table_title": table_title},
                    header_context=full_ctx,
                )
            )
            tags.extend(_extract_tags(text, {"page": 1, "paragraph": idx}))
        for img in imgs:
            images.append(
                {
                    "index": len(images),
                    "paragraph": idx,
                    "page": 1,
                    "w_in": img["w_emu"] / 914400.0,
                    "h_in": img["h_emu"] / 914400.0,
                    "rid": img["rid"],
                    "in_header_footer": header_footer,
                    "caption": "",
                }
            )

    last_table_title = ""  # 跟踪最近的表标题
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if _is_table_title(text) or is_dynamic_table_title(text):
                # 既要识别传统静态表标题，也要识别动态监测表标题
                last_table_title = text or last_table_title
            handle_paragraph(item, table_title=last_table_title)
        elif isinstance(item, Table):
            header_cells = [c.text.strip() for c in item.rows[0].cells] if item.rows else []
            # 收集每一列的"列头"——最后一行非空子串的列头信息
            for r_i, row in enumerate(item.rows):
                row_label = row.cells[0].text.strip() if row.cells else ""
                for c_i, cell in enumerate(row.cells):
                    col_header = header_cells[c_i] if c_i < len(header_cells) else ""
                    cell_info = {
                        "title": last_table_title or "",
                        "row_idx": r_i + 1,       # 1-based, 含表头
                        "col_idx": c_i + 1,       # 1-based, 含左标签列
                        "row_label": row_label,
                        "col_header": col_header,
                    }
                    for p in cell.paragraphs:
                        combined_title = last_table_title or ""
                        if c_i > 0 and row_label:
                            combined_title = f"{combined_title} {row_label}".strip()
                        elif header_cells and c_i < len(header_cells):
                            combined_title = f"{combined_title} {header_cells[c_i]}".strip()
                        idx = len(paragraphs)
                        handle_paragraph(p, table_title=combined_title)
                        # 记录段落 → 表格单元格布局
                        if c_i > 0:  # 标签列不需要（cell_ref 只在数据区）
                            table_layout_for_para[idx] = cell_info
            last_table_title = ""  # 重置

    for section in doc.sections:
        for p in section.header.paragraphs:
            handle_paragraph(p, header_footer=True)
        for p in section.footer.paragraphs:
            handle_paragraph(p, header_footer=True)

    # 图题关联：图片所在段之后第一个非空段落
    for img in images:
        nxt = img["paragraph"] + 1
        while nxt < len(paragraphs) and not paragraphs[nxt].strip():
            nxt += 1
        if nxt < len(paragraphs):
            img["caption"] = paragraphs[nxt].strip()

    # 预计算：每段对应的"最近且相邻"的表标题（用于 cell_ref 的上下文）
    # 只在距离 8 段以内的标题才算"本表"，避免远处表格标题污染
    last_table_title_for_para = {}
    current_title = ""
    gap = 0
    MAX_GAP = 8  # 距离阈值
    for idx_p, t in enumerate(paragraphs):
        t_strip = t.strip()
        if _is_table_title(t_strip):
            current_title = t_strip
            gap = 0
        else:
            gap += 1
        if current_title and gap <= MAX_GAP:
            last_table_title_for_para[idx_p] = current_title

    # 图表文本占位检测：报告中的 "xxx_time_series" / "xxx_histogram" 行
    # 以及"纯图题"（末尾含"曲线图/直方图/分布图/统计图"等关键词的段落）
    # 以及"表格 Excel 单元格引用"（H(行,列) / M(行,列) / J(行,列)）
    chart_texts = []
    seen_para = set()  # 避免同一段被多次加入
    # 车辆累计通过数量统计表：结构化表格（行=数值/比例，列=车道1..N）。
    # 源表没有“字母(行,列)”引用，单独生成带 位置(车道X)+特征(交通荷载) 的 cell 占位符，
    # 避免退化成无语义的 {{data.N}}。
    for _pidx, _info in table_layout_for_para.items():
        col_h = str(_info.get("col_header") or "")
        row_l = str(_info.get("row_label") or "")
        # 行标签可能出现全角 ％（“比例/％”），先归一化
        row_norm = row_l.replace("％", "%")
        if re.fullmatch(r"车道\d+", col_h) and re.fullmatch(r"(数值/辆|比例/%|比例%)", row_norm):
            chart_texts.append({
                "paragraph": _pidx,
                "kind": "cell_ref",
                "chart_id": f"cell_vehicle_{_info['row_idx']}_{_info['col_idx']}",
                "metric": "vehicle_count",
                "text": str(paragraphs[_pidx] if _pidx < len(paragraphs) else ""),
                "table_letter": "V",
                "row": _info.get("row_idx", 0),
                "col": _info.get("col_idx", 0),
                "row_label": col_h,          # 位置：车道1
                "col_header": row_norm,       # 统计：数值/辆 或 比例/%
                "table_title": _info.get("title", ""),
                "position": 0,
                "verdict": "replace",
                "source": "cell_ref",
                "vehicle": True,
            })
    for i, t in enumerate(paragraphs):
        t = t.strip()
        if not t or i in seen_para:
            continue
        m3 = CHART_LINE_RE.match(t)
        if m3:
            # “309(xJsd)_时程曲线”：传感器编号行。不再生成 chart_sensor_<编号>，
            # 改为 explicit_suffix 占位行，位置/方位由下方表格补全（上游/下游等）。
            feature_raw = m3.group(2).strip()
            kind = ("histogram" if ("直方" in m3.group(3) or "频率" in m3.group(3))
                    else "time_series")
            chart_id = "trend" if kind == "time_series" else "histogram"
            metric_en = (FEATURE_METRIC.get(_norm(feature_raw), "")
                         or _metric_from_chart_text(t)
                         or "chart")
            chart_texts.append({
                "paragraph": i,
                "kind": kind,
                "chart_id": chart_id,
                "metric": metric_en,
                "text": t,
                "source": "explicit_suffix",
                "sensor_ids": m3.group(1),
                "feature": feature_raw,
            })
            seen_para.add(i)
            continue
        m = CHART_TEXT_RE.search(t)
        if m:
            kind_raw = next((g for g in m.groups() if g), None) or ""
            # 中英文图型归一化（“时程曲线图/频率分布直方图” 与 time_series/histogram 等价）
            if kind_raw in ("time_series", "curve", "时程曲线", "时程曲线图", "时间序列图"):
                kind = "time_series"
            elif kind_raw in ("histogram", "频率分布直方图", "频率分布图", "直方图"):
                kind = "histogram"
            else:
                kind = kind_raw
            # 从文本前缀提取指标关键字，生成更有意义的图表 ID
            prefix = t[:m.start()].strip()
            metric_en = _metric_from_chart_text(prefix)
            chart_id = "trend" if kind in ("time_series", "curve") else "histogram"
            chart_texts.append({
                "paragraph": i,
                "kind": kind,
                "chart_id": chart_id,
                "metric": metric_en,
                "text": t,
                "source": "explicit_suffix",
            })
            seen_para.add(i)
            continue
        m2 = CHART_MULTI_RE.match(t)
        if m2:
            # “位移*906*907*908*时程曲线_x1”：多个传感器编号合并一行，
            # 视为该位置的 时程/直方 占位行（位置由下方表格补全，多传感器合并成一张图）
            kind = "time_series" if "时程" in m2.group(2) else "histogram"
            chart_id = "trend" if kind == "time_series" else "histogram"
            metric_en = _metric_from_chart_text(t) or "displacement"
            chart_texts.append({
                "paragraph": i,
                "kind": kind,
                "chart_id": chart_id,
                "metric": metric_en,
                "text": t,
                "source": "explicit_suffix",
                "sensor_ids": m2.group(1),
            })
            seen_para.add(i)
            continue
        # === 纯图题识别：段落文本本身是图题且包含"图"和动态指标词 ===
        # 例："交通荷载监测_车辆累计通过数量统计图"、"梁端转角测点变化时程曲线图"
        if t.endswith("图") and any(
            kw in t for kw in ("曲线", "时程", "直方", "柱状", "分布", "趋势",
                               "占比", "变化", "统计", "频谱", "频次", "相关", "车流")
        ) and any(
            kw in t for kw in ("温度", "湿度", "风速", "风向", "转角", "索夹", "位移",
                               "车辆", "交通", "荷载", "应变", "应力", "挠度", "索力",
                               "沉降", "振动", "结构", "桥面", "主梁", "塔", "钢桁")
        ):
            # 从图题中推断指标
            metric_en = _metric_from_chart_text(t)
            # 推断图表类型
            if "直方" in t or "分布" in t or "频次" in t or "频谱" in t:
                kind = "histogram"
                chart_id = "histogram"
            elif "柱状" in t or "占比" in t or "车流" in t or "通过数量" in t or "累计" in t:
                kind = "bar"
                chart_id = "bar"
            elif "散点" in t or "相关" in t:
                kind = "scatter"
                chart_id = "scatter"
            else:
                kind = "trend"
                chart_id = "trend"
            chart_texts.append({
                "paragraph": i,
                "kind": kind,
                "chart_id": chart_id,
                "metric": metric_en,
                "text": t,
                "source": "bare_caption",
            })
            seen_para.add(i)
            continue
        # === 表格 Excel 单元格引用占位符识别 ===
        # 例："H(10,2)" "M(2,2)" "J(2,2)" 等
        # 表格字母 → 指标（与模板约定）+ 是否动态
        letter_to_metric = {
            # 动态表（监测统计结果）
            "H": ("cable_clamp", True),    # 索夹滑动 Hangers
            "M": ("rotation", True),       # 转角 Measurement
            "J": ("wind_speed", True),     # 风速
            # 静态表（规范/阈值/设备）
            "A": ("static_a", False), "B": ("static_b", False),
            "C": ("static_c", False), "D": ("static_d", False),
            "E": ("static_e", False), "F": ("static_f", False),
            "G": ("static_g", False),
        }
        # 表格标题信号（属于静态表的特征）
        static_table_signals = [
            "监测内容表", "监测内容", "监测项目表", "测点布设表", "测点布置表",
            "测点一览表", "传感器类型表", "设备清单", "装备设置", "装备清单",
            "阈值表", "监测阈值", "作用监测阈值", "结构响应监测阈值",
            "报警阈值", "控制值表", "作用阈值",
            "对照表", "分级表", "分级对照", "等级对照",
            "风力分级", "风级对照", "风力对照",
            "监测部位表", "监测要素表",
            "规范表", "标准表",
            "交通量统计表", "车道分布",
            "车速统计", "车型统计",
        ]
        if CELL_REF_RE.search(t):
            cell_refs = list(CELL_REF_RE.finditer(t))
            for cm in cell_refs:
                table_letter = cm.group(1)
                row_idx = int(cm.group(2))   # letter 模板给定的"行"
                col_idx = int(cm.group(3))   # letter 模板给定的"列"
                # 上下文：先看 table_layout（更准），再看 last_table_title_for_para
                layout = table_layout_for_para.get(i, {})
                ctx_title = layout.get("title") or last_table_title_for_para.get(i, "")
                row_label = layout.get("row_label", "")
                col_header = layout.get("col_header", "")
                # 推断 metric：用 table title 优先，再用 letter 兜底
                if ctx_title and is_dynamic_table_title(ctx_title):
                    metric_en = metric_from_title(ctx_title)
                    is_static_table = False
                else:
                    # 字母兜底（保留旧约定）
                    letter_info = {
                        "H": ("cable_clamp", True),
                        "M": ("rotation", True),
                        "J": ("wind_speed", True),
                    }
                    metric_en, is_dynamic_by_letter = letter_info.get(
                        table_letter, (f"table_{table_letter.lower()}", True)
                    )
                    if not is_dynamic_by_letter:
                        is_static_table = True
                    elif ctx_title and not is_dynamic_table_title(ctx_title):
                        is_static_table = True
                    else:
                        is_static_table = False
                        # 静态表但有合法 cell_ref，兜底用 letter title
                        if metric_en.startswith("table_"):
                            metric_en = ctx_title or metric_en
                # 推断 stat：用列头
                stat_from_header = stat_from_col_header(col_header)
                # 推断 row 标识：清洗 row_label
                row_slug = slugify_label(row_label) or f"r{row_idx}"
                col_slug = slugify_label(col_header) or f"c{col_idx}"
                verdict = "keep" if is_static_table else "replace"
                chart_texts.append({
                    "paragraph": i,
                    "kind": "cell_ref",
                    "chart_id": f"cell_{table_letter.lower()}_{row_idx}_{col_idx}",
                    "metric": metric_en,
                    "text": cm.group(0),
                    "table_letter": table_letter,
                    "row": row_idx,
                    "col": col_idx,
                    "row_label": row_label,
                    "col_header": col_header,
                    "table_title": ctx_title,
                    "position": cm.start(),
                    "verdict": verdict,
                    "source": "cell_ref",
                })

    return {
        "images": images,
        "numbers": numbers,
        "tags": tags,
        "_table_layout": table_layout_for_para,
        "_chart_texts_raw": chart_texts,

        "numbers": numbers,
        "texts": paragraphs,
        "chart_texts": chart_texts,
    }


# ---------------------------------------------------------------------------
# PDF 解析
# ---------------------------------------------------------------------------

def _pdf_lines(page):
    """把一页文字按行归组，返回 [(top, text), ...]。"""
    words = page.extract_words(x_tolerance=2, y_tolerance=3) or []
    lines = {}
    for w in words:
        key = round(w["top"] / 4.0) * 4
        lines.setdefault(key, []).append(w)
    result = []
    for top, ws in sorted(lines.items()):
        ws.sort(key=lambda x: x["x0"])
        result.append((top, " ".join(w["text"] for w in ws)))
    # 合并被换行截断的数字：前一行以 "." 结尾且下一行以数字开头（如 "1." + "18℃"）
    merged = []
    for top, txt in result:
        if merged and merged[-1][1].endswith(".") and re.match(r"\d", txt):
            merged[-1] = (merged[-1][0], merged[-1][1] + txt)
        else:
            merged.append((top, txt))
    return merged


def parse_pdf(path: str) -> dict:
    import pdfplumber

    numbers = []
    images = []
    all_texts = []
    with pdfplumber.open(path) as pdf:
        for pno, page in enumerate(pdf.pages, 1):
            lines = _pdf_lines(page)
            for top, txt in lines:
                numbers.extend(
                    _extract_numbers(txt, {"page": pno, "paragraph": int(top), "table": None})
                )
                all_texts.append(txt)

            for im in page.images:
                x0, x1 = im.get("x0", 0), im.get("x1", 0)
                y0, y1 = im.get("top", 0), im.get("bottom", 0)
                caption = None
                for top, txt in sorted(lines):
                    if top >= y1 - 2 and top <= y1 + 48 and CAPTION_RE.match(txt):
                        caption = txt
                        break
                images.append(
                    {
                        "index": len(images),
                        "page": pno,
                        "w_in": (x1 - x0) / 72.0,
                        "h_in": (y1 - y0) / 72.0,
                        "bbox": [round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1)],
                        "in_header_footer": False,
                        "caption": caption or "",
                    }
                )
    return {"images": images, "numbers": numbers, "texts": all_texts}


# ---------------------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------------------

def _build_doc_text(parsed: dict, max_chars: int = 22000) -> str:
    """把解析出的段落文本整理为 LLM 校验用的报告文本摘录。

    只保留含数字的段落（漏掉的动态值必然是数字），并按行截断控制提示词规模。
    """
    texts = parsed.get("texts") or []
    lines = []
    total = 0
    for i, t in enumerate(texts):
        t = t.strip()
        if not t or not re.search(r"\d", t):
            continue
        # 行内截断，保留数字附近的上下文
        if len(t) > 150:
            t = t[:150] + "…"
        lines.append(f"段落{i}: {t}")
        total += len(t) + 8
        if total > max_chars:
            lines.append(f"……（文本过长已截断，共 {len(texts)} 段）")
            break
    return "\n".join(lines)


def _build_doc_chunks(parsed: dict, max_chars: int = 10000) -> List[tuple]:
    """按节标题把报告文本切分为若干分批片段。

    每段尽量从“数字编号节标题”开始（跨小节合并到接近 max_chars），
    避免把整份报告一次性塞给大模型。
    返回 [(文本, 起始段落, 结束段落)]。
    """
    texts = parsed.get("texts") or []
    heading_re = re.compile(r"^\d+(?:\.\d+){0,3}\.?(?=[\u4e00-\u9fa5\s])")
    chunks = []
    cur_start = 0
    cur_lines = []
    cur_len = 0
    for i, t in enumerate(texts):
        ts = str(t).strip()
        if not ts or not re.search(r"\d", ts):
            continue
        is_heading = bool(heading_re.match(ts)) and len(ts) <= 60
        # 遇到新节标题且当前片段已超过一半上限 -> 切段（新标题作为下一段开头）
        if is_heading and cur_lines and cur_len >= max_chars * 0.5:
            chunks.append((cur_start, i, "\n".join(cur_lines)))
            cur_start = i
            cur_lines = []
            cur_len = 0
        line = ts if len(ts) <= 150 else ts[:150] + "…"
        cur_lines.append(f"段落{i}: {line}")
        cur_len += len(line) + 8
        if cur_len >= max_chars:
            chunks.append((cur_start, i + 1, "\n".join(cur_lines)))
            cur_start = i + 1
            cur_lines = []
            cur_len = 0
    if cur_lines:
        chunks.append((cur_start, len(texts), "\n".join(cur_lines)))
    return chunks


def _locate_value(value: str, snippet: str, texts: list) -> Optional[tuple]:
    """在段落文本中定位某个数字，返回 (paragraph_index, position) 或 None。

    优先用 snippet 匹配，其次用 value 匹配。
    """
    if not texts:
        return None
    # 尝试按 snippet 片段匹配
    if snippet:
        target = snippet.strip()
        # snippet 可能被截断，用前 30 字匹配
        probe = target[:30]
        for i, t in enumerate(texts):
            if probe and probe in t:
                pos = t.find(value)
                if pos != -1:
                    return i, pos
    # 退而求其次：按 value 匹配
    for i, t in enumerate(texts):
        pos = t.find(value)
        if pos != -1:
            return i, pos
    return None


def recognize(path: str, llm_cfg: Optional[dict] = None,
              sensor_map: Optional[dict] = None) -> dict:
    """解析报告并给图片/数字打上 replace / keep / review 结论。

    sensor_map: 传感器对照表 {编号: {监测部位/名称/类别/特征编码...}}，
        用于“编号(特征)_图型”行反查监测部位，生成位置化图表占位符
        （如 vibration_随州侧边跨跨中截面上游_trend_1），
        避免用传感器编号作占位符导致运行时难以匹配图库。

    两轮筛选流程：
      第一轮 — 关键词打分：classify_number / classify_image 基于关键词和文本
              长度给出初步 verdict（replace / keep / review）。
      第二轮 — LLM 完整性校验（llm_classifier.verify_and_complete）：
              1. 检查第一轮提取结果是否漏掉动态值（missed）
              2. 检查是否有静态值被误判为动态（wrong）
              3. 对待确认项（review）给出最终判定
              4. 识别季度/时间表述并建议动态化（text_replacements）
              若全部正确，LLM 只回复"是"。
      第二轮不可用（无 API Key / 调用失败）时，自动降级为文本长度启发式。
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        parsed = parse_docx(path)
    elif ext == ".pdf":
        parsed = parse_pdf(path)
    else:
        raise ValueError(f"仅支持 .docx / .pdf 报告，收到: {path}")

    # 第一轮：关键词打分
    for img in parsed["images"]:
        info = classify_image(
            img["caption"], img.get("w_in", 0), img.get("h_in", 0),
            img.get("in_header_footer", False),
        )
        img.update(info)
        img["_image_index"] = img.get("index", 0)

    for n in parsed["numbers"]:
        # 占位符规范化：统一为 {{stats.<metric>.<stat>}} 标准
        ph = n.get("placeholder")
        if ph and not re.match(r"^stats\.[a-zA-Z_]+\.(max|min|avg|median|std|range)$", ph):
            n["placeholder"] = None

    # 原文占位标记（[A-MAX]、[A-MAX-LOC] 等）：按上下文映射为 stats.* 占位符
    texts_all = parsed.get("texts") or []
    for tag in parsed.get("tags", []) or []:
        _classify_tag(tag, texts_all)
    log.info("原文占位标记识别：%d 个（如 A-MAX/A-MAX-LOC）", len(parsed.get("tags", []) or []))

    # 静态数字保护：位置名/塔号/跨号/测点数量/节标题编号等一律保留
    n_static = _protect_static_numbers(parsed)
    if n_static:
        log.info("静态数字保护：%d 个数字改回保留（位置/塔号/跨号/数量词/标题编号）", n_static)

    # 第二轮：LLM 完整性校验
    classifier = LLMClassifier(llm_cfg)

    review_numbers = [n for n in parsed["numbers"] if n["verdict"] == "review"]
    review_images = [img for img in parsed["images"] if img["verdict"] == "review"]

    # review 项过多时只把最模糊的（置信度接近 0.5）发给 LLM，其余本地启发式兜底
    MAX_LLM_REVIEW = 120
    review_numbers.sort(key=lambda n: abs(n.get("confidence", 0.5) - 0.5))
    llm_batch_reviews = review_numbers[:MAX_LLM_REVIEW]
    heuristic_rest = review_numbers[MAX_LLM_REVIEW:]
    for i, n in enumerate(llm_batch_reviews):
        n["_review_index"] = i
    for i, img in enumerate(review_images):
        img["_review_index"] = i

    extracted = [
        {
            "value": n.get("value", ""),
            "snippet": n.get("snippet", ""),
            "placeholder": n.get("placeholder") or "",
            "paragraph": n.get("paragraph"),
        }
        for n in parsed["numbers"]
        if n["verdict"] == "replace"
    ]

    # 第一轮：按节分批识别 文字/表格 的动态数字（避免整份报告一次塞给大模型）
    chunks = _build_doc_chunks(parsed, max_chars=10000)
    log.info("LLM 数字识别按节分批：共 %d 段", len(chunks))
    chunk_results = []
    for p0, p1, chunk_text in chunks:
        def _in_range(item):
            try:
                return p0 <= int(item.get("paragraph")) < p1
            except (TypeError, ValueError):
                return False
        extracted_chunk = [e for e in extracted if _in_range(e)]
        review_chunk = [n for n in llm_batch_reviews if _in_range(n)]
        r = classifier.verify_and_complete(
            doc_text=chunk_text,
            extracted=extracted_chunk,
            review_items=review_chunk,
            images=[],
            review_images=[],
        )
        chunk_results.append(r)

    # 合并各节结果
    llm_result = {
        "complete": all(r.get("complete") for r in chunk_results) if chunk_results else False,
        "missed": [m for r in chunk_results for m in (r.get("missed") or [])],
        "wrong": [w for r in chunk_results for w in (r.get("wrong") or [])],
        "review_decisions": {k: v for r in chunk_results
                             for k, v in (r.get("review_decisions") or {}).items()},
        "text_replacements": [t for r in chunk_results
                              for t in (r.get("text_replacements") or [])],
        "raw": "\n".join(r.get("raw") or "" for r in chunk_results),
    }
    doc_text = _build_doc_text(parsed)
    # 第二轮：单独识别图片
    images_result = classifier.verify_images(
        doc_text=doc_text,
        images=parsed["images"],
        review_images=review_images,
    )

    # 2a-1. 应用 LLM 对 review 数字的判定
    for n in llm_batch_reviews:
        idx = n.pop("_review_index")
        dec = llm_result["review_decisions"].get(idx)
        if dec and dec["verdict"] != "review":
            n["verdict"] = dec["verdict"]
            n["confidence"] = dec.get("confidence", 0.85)
            n["reasons"].append(dec.get("reason", "LLM 判定"))
            if dec["verdict"] == "replace":
                n["placeholder"] = dec.get("placeholder") or suggest_placeholder(
                    (n.get("snippet") or "")[:12], (n.get("snippet") or "")[12:16]
                )

    # 2a-2. 未发送给 LLM 的 review 项：文本长度启发式兜底
    for n in heuristic_rest:
        r = _text_length_heuristic(n)
        if r["verdict"] != "review":
            n["verdict"] = r["verdict"]
            n["confidence"] = r["confidence"]
            n["reasons"].append(r["reason"])
            if n["verdict"] == "replace" and not n.get("placeholder"):
                n["placeholder"] = suggest_placeholder(
                    (n.get("snippet") or "")[:12], (n.get("snippet") or "")[12:16]
                )

    # 2b. 补充漏掉的动态值（LLM 找到的 missed）
    texts = parsed.get("texts") or []
    data_fallback_counter = [sum(1 for n in parsed["numbers"] if n["verdict"] == "replace")]
    for m in llm_result["missed"]:
        # 跳过原文占位标记（[A-MAX] 等）：已由本地上下文映射处理，避免重复替换
        if TAG_TOKEN_RE.fullmatch(str(m.get("value", ""))):
            continue
        # 跳过非数字值（如 LLM 把 “RE24-RE25” 节点板编号误当漏掉的动态值）
        if not re.fullmatch(r"[\d.+\-–—~～eE%％,，\s]+", str(m.get("value", "")).strip()):
            continue
        loc = _locate_value(m.get("value", ""), m.get("snippet", ""), texts)
        if loc is None:
            continue
        p_idx, pos = loc
        # 去重：同一位置已存在 replace 项则跳过
        if any(
            n.get("paragraph") == p_idx and n.get("position") == pos
            for n in parsed["numbers"]
        ):
            continue
        data_fallback_counter[0] += 1
        marker = m.get("placeholder") or f"{{{{data.{data_fallback_counter[0]}}}}}"
        parsed["numbers"].append({
            "verdict": "replace",
            "confidence": 0.90,
            "reasons": [f"LLM 补充识别：{m.get('reason', '')}"],
            "placeholder": marker,
            "value": m.get("value", ""),
            "snippet": m.get("snippet", ""),
            "position": pos,
            "page": 1,
            "paragraph": p_idx,
            "table": None,
        })

    # 2c. 纠正误判（LLM 认为应为 keep 的）
    wrong_values = {(w.get("value", ""), w.get("snippet", "")[:30]) for w in llm_result["wrong"]}
    for n in parsed["numbers"]:
        key = (n.get("value", ""), (n.get("snippet") or "")[:30])
        if n["verdict"] == "replace" and key in wrong_values:
            n["verdict"] = "keep"
            n["confidence"] = 0.85
            n["reasons"].append("LLM 判定为静态值，改回保留")
            n["placeholder"] = None

    # 2d. 应用图片 review 判定与误判纠正（第二轮图片识别结果）
    for img in review_images:
        idx = img.pop("_review_index")
        v = images_result["images_review"].get(idx)
        if v in ("replace", "keep"):
            img["verdict"] = v
            img["confidence"] = 0.85
            img["reasons"].append("LLM 判定（图片轮）")
    for img in parsed["images"]:
        if img["verdict"] == "replace" and img["_image_index"] in images_result["images_wrong"]:
            img["verdict"] = "keep"
            img["reasons"].append("LLM 判定为固定图（图片轮），改回保留")
    for img in parsed["images"]:
        img.pop("_image_index", None)

    summary = {
        "images": {
            "replace": sum(1 for i in parsed["images"] if i["verdict"] == "replace"),
            "keep": sum(1 for i in parsed["images"] if i["verdict"] == "keep"),
            "review": sum(1 for i in parsed["images"] if i["verdict"] == "review"),
        },
        "numbers": {
            "replace": sum(1 for n in parsed["numbers"] if n["verdict"] == "replace"),
            "keep": sum(1 for n in parsed["numbers"] if n["verdict"] == "keep"),
            "review": sum(1 for n in parsed["numbers"] if n["verdict"] == "review"),
        },
        "chart_texts": len(parsed.get("chart_texts", [])),
        "llm": {
            "enabled": classifier.available(),
            "complete": llm_result.get("complete", False),
            "missed": len(llm_result.get("missed", [])),
            "wrong": len(llm_result.get("wrong", [])),
            "text_replacements": len(llm_result.get("text_replacements", [])),
        },
    }
    return {
        "source": os.path.abspath(path),
        "type": ext.lstrip("."),
        "parsed_at": dt.datetime.now().isoformat(timespec="seconds"),
        "summary": summary,
        "images": parsed["images"],
        "numbers": parsed["numbers"],
        "tags": parsed.get("tags", []),
        "chart_texts": parsed.get("chart_texts", []),
        "text_replacements": llm_result.get("text_replacements", []),
        "texts": parsed.get("texts", []),
        "sensor_map": sensor_map or {},
    }


def save_analysis(analysis: dict, out_path: str) -> str:
    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(analysis, f, ensure_ascii=False, indent=2)
    return out_path


def print_summary(analysis: dict) -> None:
    print("=" * 64)
    print("报告解析识别结果")
    print("=" * 64)
    print(f"来源: {analysis['source']}（{analysis['type']}）")
    s = analysis["summary"]
    llm_info = s.get("llm", {})
    if llm_info:
        status = "已启用" if llm_info.get("enabled") else "不可用（降级）"
        complete = "全部正确（回复'是'）" if llm_info.get("complete") else "有补充修改"
        print(
            f"LLM 二次筛选: {status} | {complete} | "
            f"补漏 {llm_info.get('missed', 0)} | 纠错 {llm_info.get('wrong', 0)} | "
            f"文本替换 {llm_info.get('text_replacements', 0)}"
        )
    print(f"\n图片: 建议替换 {s['images']['replace']} / 保留 {s['images']['keep']} / 待确认 {s['images']['review']}")
    if s.get("chart_texts"):
        print(f"图表文本占位: {s['chart_texts']} 处（_time_series/_histogram 行，将替换为 {{chart.*}}）")
    for i in analysis["images"]:
        cap = i.get("caption") or "（无图题）"
        mark = {"replace": "[替换]", "keep": "[保留]", "review": "[待确认]"}[i["verdict"]]
        print(f"  {mark} 图#{i['index']} 页{i['page']} {i.get('w_in', 0):.1f}x{i.get('h_in', 0):.1f}in  {cap}")

    print(f"\n数字: 建议替换 {s['numbers']['replace']} / 保留 {s['numbers']['keep']} / 待确认 {s['numbers']['review']}")
    shown = 0
    for n in analysis["numbers"]:
        if n["verdict"] != "replace":
            continue
        ph = f"  ->  {n['placeholder']}" if n.get("placeholder") else ""
        print(f"  [替换] {n['value']:>8}  置信 {n['confidence']:.2f}{ph}")
        print(f"          {n['snippet']}")
        shown += 1
        if shown >= 15:
            print("  ... 其余替换项见 JSON")
            break
    kept = [n for n in analysis["numbers"] if n["verdict"] == "keep"]
    if kept:
        print(f"\n保留的固定数值示例（前 8 条）:")
        for n in kept[:8]:
            print(f"  [保留] {n['value']:>8}  {n['snippet']}")
    print(f"\n完整结果: 请查看输出的 JSON 文件")


# ---------------------------------------------------------------------------
# DOCX 标注草稿：动态项 -> 占位符
# ---------------------------------------------------------------------------

def _suggest_chart_id(caption: str, fallback: str) -> str:
    caption = caption or ""
    for kw, cid in [
        ("直方", "histogram"),
        ("柱状", "daily_bars"),
        ("箱线", "boxplot"),
        ("折线", "trend"),
        ("曲线", "trend"),
        ("趋势", "trend"),
        ("分布", "histogram"),
        ("对比", "daily_bars"),
    ]:
        if kw in caption:
            return cid
    return fallback


def _infer_chart_location(paragraph, cell_ref_paras: Dict[int, tuple],
                          heading_paras=None, lookahead: int = 120) -> str:
    """推断图表占位符对应的监测位置。

    cell_ref_paras: {段落索引: (表标题, row_label)}。
    取图表段落之后“最近的表格”（中间不能跨节标题）的全部 row_label；
    只有单位置时才返回该位置，多位置（如跨中断面 3 个箱内/桥面测点）
    留给运行时按表格全量补图，返回空串。
    """
    if not isinstance(paragraph, int) or not cell_ref_paras:
        return ""
    nxt = [p for p in cell_ref_paras if p > paragraph]
    if not nxt:
        return ""
    p_first = min(nxt)
    # 图表段与表格之间不能出现新的节标题（说明表格属于下一节，如风荷载表在图上）
    if heading_paras:
        for hp in heading_paras:
            if paragraph < hp < p_first:
                return ""
    title_first = cell_ref_paras[p_first][0]
    rows = []
    for p in sorted(cell_ref_paras):
        if p_first <= p <= p_first + lookahead and cell_ref_paras[p][0] == title_first:
            rows.append(cell_ref_paras[p][1])
    uniq = sorted({str(x).strip() for x in rows if str(x).strip()})
    return uniq[0] if len(uniq) == 1 else ""


def _chart_block_locations(paragraph, cell_ref_paras: Dict[int, tuple],
                           heading_paras=None, lookahead: int = 160,
                           texts: Optional[list] = None) -> List[str]:
    """图表块下方“最近的表格”的监测部位集合（保持表格行顺序，去重）。

    规则：
      - 优先向下找最近的表格；中间不能跨节标题；
      - 向下找不到时向上找（图表块可能位于表格之后，如 风速节“下图为…”在统计表后）；
      - 同标题表格多张相邻时（源报告常把不同节表格都命名为“xx监测统计”），
        以最近的节标题为界，只取本节的表格，避免把后面几节的监测部位全收进来。
    """
    if not isinstance(paragraph, int) or not cell_ref_paras:
        return []
    nxt = [p for p in cell_ref_paras if p > paragraph]
    p_first = min(nxt) if nxt else None
    forward_ok = p_first is not None
    if forward_ok and heading_paras:
        for hp in heading_paras:
            if paragraph < hp < p_first:
                forward_ok = False  # 最近的表被节标题隔开，向下查找作废
                break
    rows = []
    if forward_ok:
        title_first = cell_ref_paras[p_first][0]
        hi = p_first + lookahead
        if heading_paras:
            for hp in heading_paras:
                if p_first < hp < hi:
                    hi = hp
                    break
        for p in sorted(cell_ref_paras):
            if p_first <= p < hi and cell_ref_paras[p][0] == title_first:
                rows.append(cell_ref_paras[p][1])
    if not rows:
        # 向下没有归属表格：向上找同节最近的表格（图表块在表格之后）
        prev = [p for p in cell_ref_paras if p < paragraph]
        h_prev = 0
        if heading_paras:
            for hp in heading_paras:
                if hp < paragraph:
                    h_prev = hp
                else:
                    break
            prev = [p for p in prev if p > h_prev]
        if prev:
            p_last = max(prev)
            title_prev = cell_ref_paras[p_last][0]
            lo = h_prev
            for p in sorted(cell_ref_paras):
                if lo < p <= p_last and cell_ref_paras[p][0] == title_prev:
                    rows.append(cell_ref_paras[p][1])
    uniq = []
    for r in rows:
        r = str(r).strip()
        # 保留“左侧/右侧”等方向行——它们是两个特征（如左支座/右支座），
        # 需要各自展开 时程图+直方图（由 _expand_chart_blocks 用节标题位置补全）
        if r and len(r) >= 1 and r not in uniq:
            uniq.append(r)
    if not uniq:
        return uniq
    # X/Y 是同一特征的方向（如 支座右侧X/支座右侧Y、梁端左侧X/左侧Y），
    # 不单独成图：全部行都以 X/Y 结尾时折叠回基础位置（右侧/左侧…）
    xy_rows = [r for r in uniq if re.fullmatch(r".+[XY]", r)]
    if len(xy_rows) == len(uniq):
        bases = []
        for r in uniq:
            b = r[:-1]
            if b and b not in bases:
                bases.append(b)
        if bases:
            uniq = bases
    # “顶板测点N / 底板测点N”：顶板/底板是不同的特征部位，测点N只是该特征下
    # 的传感器序号。折叠成特征前缀，并用表格标题补全成完整位置
    # （如 上游炎陵侧边跨跨中截面顶板 / …底板），避免 5 个测点行展开成 10 张图。
    prefixes = []
    ok = True
    for r in uniq:
        m = re.match(r"^(.+?)测点\s*\d*\s*$", r)
        if not m:
            ok = False
            break
        p = m.group(1).strip()
        if p and p not in prefixes:
            prefixes.append(p)
    if ok and prefixes:
        base = _position_from_title(title_first)
        # 表标题太泛（如“结构温度监测统计”）时提取不到基座，回退到节标题
        if not base and heading_paras and texts:
            for hp in reversed(heading_paras):
                if hp < paragraph:
                    base = _position_from_title(str(texts[hp]))
                    if base:
                        break
        if base:
            out = []
            for p in prefixes:
                # 前缀已是完整位置（如 行标签“吉首侧索塔中截面测点N”的前缀
                # “吉首侧索塔中截面”），不能再和基座拼接重复
                if base in p or p in base:
                    composed = p if len(p) >= len(base) else base
                else:
                    composed = base + p
                if composed not in out:
                    out.append(composed)
            return out
    # 节标题方位词补全：标题带“上游侧/下游侧”等方位，而表格位置未带时，
    # 把方位词拼进位置（湘江：3.1.1.1上游侧箱梁内环境温度 + 表“随州侧边跨跨中截面”
    # → “随州侧边跨跨中截面上游”），否则运行时图库/名称对照匹配不到。
    if heading_paras and texts:
        heading_dir = ""
        for hp in reversed(heading_paras):
            if hp < paragraph:
                ht = str(texts[hp])
                # 只取最近的节标题；若它本身没有方位词，就不继续往回找
                heading_dir = next((w for w in ("上游", "下游") if w in ht), "")
                break
        if heading_dir:
            uniq = [r if heading_dir in r else r + heading_dir for r in uniq]
    return uniq


def _position_from_title(title: str) -> str:
    """从节标题/表标题提取监测位置（去掉章节号、指标词、监测/统计等）。"""
    # 只剥掉章节号（如 3.3.6.1），不能吃掉墩号数字（如 “2#墩…” 开头的 2）
    t = re.sub(r"^\d+(?:\.\d+){0,3}(?![.\d#])\s*", "", str(title or ""))
    for w in ("结构应变监测", "环境湿度监测", "环境温度监测", "结构温度监测",
              "挠度监测", "位移监测", "倾角监测", "裂缝监测", "索力监测",
              "风速风向监测", "风速监测", "支座位移监测", "交通监测",
              "空间变位监测", "空间变位", "变位监测", "变位",
              "地震监测", "振动监测", "应力监测", "承台监测",
              "振动的", "位移的", "温度的", "湿度的", "挠度的", "应变的",
              "时程曲线图", "频率分布直方图", "相关性散点图", "散点图",
              "曲线图", "直方图", "如下图所示", "时程",
              "结构应变", "结构温度", "环境温度", "环境湿度",
              "挠度", "位移", "应变", "倾角", "裂缝", "索力",
              "风速", "风向", "支座", "监测", "数据分析", "统计",
              "结果如下表", "结果", "如下表"):
        t = t.replace(w, "")
    t = re.sub(r"[\s：:。]+", "", t)
    t = t.rstrip("表")  # 表标题末尾的“表”字（如 …监测统计表）
    t = t.strip("、，,和及")
    return t if len(t) >= 2 else ""


def _cell_position_from_title(title: str) -> str:
    """从表格标题提取 cell 占位符的位置基座（比 _position_from_title 更宽，
    保留“上游/下游/左/右”等方位，供行标签拼接）。

    例：“上游位置环境温度监测统计” -> “上游”
        “上游随州侧边跨跨中箱梁结构温度监测统计” -> “上游随州侧边跨跨中箱梁”
        “随州侧边跨跨中截面环境温度监测统计” -> “随州侧边跨跨中截面”
    """
    t = re.sub(r"^\d+(?:\.\d+){0,3}(?![.\d#])\s*", "", str(title or ""))
    for w in ("结构温度监测", "环境温度监测", "环境湿度监测", "应变监测",
              "挠度监测", "位移监测", "倾角监测", "振动监测", "地震监测",
              "索力监测", "裂缝监测", "风速监测", "风向监测",
              "支座位移监测", "结构应变监测", "结构振动监测",
              "时程曲线图", "频率分布直方图", "监测统计", "监测数据",
              "监测结果", "统计结果", "如下表", "监测", "统计",
              "结果", "数据分析", "平均温度", "最高温度", "最低温度",
              "最大温差", "平均应变", "最大应变", "最小应变",
              "最大应变差", "平均", "最高", "最低", "最大", "最小",
              "差值", "剔除温度", "相关性系数", "均方根", "标准差"):
        t = t.replace(w, "")
    t = re.sub(r"[\s：:。]+", "", t)
    t = t.rstrip("表")
    t = t.strip("、，,和及")
    return t if len(t) >= 1 else ""


def _norm(text: str) -> str:
    """轻量归一化：小写、去空格、全角转半角（与 bridge_source._norm 一致）。"""
    if not text:
        return ""
    out = []
    for ch in str(text):
        code = ord(ch)
        if code == 0x3000:
            out.append(" ")
        elif 0xFF01 <= code <= 0xFF5E:
            out.append(chr(code - 0xFEE0))
        else:
            out.append(ch)
    return "".join(out).strip().lower().replace(" ", "")


def _scatter_position_from_caption(caption: str, fallback: str = "") -> str:
    """从“第五跨L/4处主梁截面测点1结构应变-温度相关性散点图”提取位置。"""
    t = str(caption or "")
    for w in ("结构应变-温度相关性散点图", "应变-温度相关性散点图",
              "结构应变-温度相关性", "-温度相关性散点图", "相关性散点图",
              "应变-温度", "-温度", r"截面测点\d+"):
        t = re.sub(w, "", t)
    t = re.sub(r"截面测点\d+", "", t)
    t = re.sub(r"测点\d+", "", t)
    t = t.replace("截面", "").replace("测点", "").strip("、，,和及")
    return t if len(t) >= 2 else fallback


def _expand_chart_blocks(analysis: dict, cell_ref_paras: Dict[int, tuple],
                         heading_paras=None) -> tuple:
    """把图表文本占位行按“监测部位 × 图型”展开为带位置的占位符。

    例如“第6、7跨跨中断面环境温度监测…如下图所示：”+ 表格 3 个部位
    → 展开 6 个 {{chart.temperature_<部位>_<图型>_<n>}}。
    图型集合取“占位行后缀 + 描述行关键词”（时程曲线图/频率分布直方图），
    保证“锚固区环境湿度…时程曲线图、频率分布直方图”这类两种图都生成。
    表格行是“测点N”时，位置从节标题回推（如“第五跨L/4处主梁”）。
    “相关性散点图”图题单独生成 scatter 占位符。

    返回 (targets, inserts, new_entries, removed_paras)：
      targets:      {段落索引: marker}
      inserts:      {锚点段落索引: [marker, ...]}  需要插入的新占位段落
      new_entries:  展开后的 chart_texts 条目（替换原 explicit_suffix 条目）
      removed_paras: 块内多余、应清空的段落索引
    """
    cts = analysis.get("chart_texts", []) or []
    texts = analysis.get("texts", []) or []
    suffix_items = sorted(
        [ct for ct in cts if ct.get("source") == "explicit_suffix"],
        key=lambda c: c.get("paragraph") or 0,
    )
    # 相邻段落（<=3）归为同一个图表块
    blocks = []
    for item in suffix_items:
        p = item.get("paragraph")
        if blocks and isinstance(p, int) and isinstance(blocks[-1][-1].get("paragraph"), int) \
                and p - blocks[-1][-1]["paragraph"] <= 3:
            blocks[-1].append(item)
        else:
            blocks.append([item])

    targets: Dict[int, str] = {}
    inserts: Dict[int, List[str]] = {}
    new_entries = []
    removed: set = set()
    counter: Dict[str, int] = {}

    for block in blocks:
        paras = sorted({c["paragraph"] for c in block if isinstance(c.get("paragraph"), int)})
        if not paras:
            continue
        p_min = paras[0]
        kinds = []
        for c in block:
            k = c.get("kind")
            norm = "trend" if k in ("time_series", "trend", "curve") else "histogram"
            if norm not in kinds:
                kinds.append(norm)
        # 描述行关键词（“…时程曲线图、频率分布直方图如下图所示：”）
        for pi in range(p_min - 1, max(p_min - 5, -1), -1):
            t = str(texts[pi]) if 0 <= pi < len(texts) else ""
            if ("如下图所示" in t or "下图" in t) and ("曲线图" in t or "直方图" in t):
                if "时程曲线图" in t or "时间序列" in t:
                    if "trend" not in kinds:
                        kinds.append("trend")
                if "频率分布直方图" in t or "直方图" in t:
                    if "histogram" not in kinds:
                        kinds.append("histogram")
                break
        if "trend" in kinds and "histogram" in kinds:
            kinds = ["trend", "histogram"]
        elif not kinds:
            kinds = ["trend"]
        metric = str(block[0].get("metric") or "chart")
        locations = _chart_block_locations(p_min, cell_ref_paras, heading_paras,
                                           texts=texts)
        if metric == "chart" and locations:
            # 占位行没有中文关键词（如 new_sensor_group_1）时，从下方表格标题回推指标
            nxt = [p for p in cell_ref_paras if p > p_min]
            if nxt:
                metric = _metric_from_chart_text(cell_ref_paras[min(nxt)][0]) or metric
        title_fallback = False
        if not locations:
            # 表格行是“测点N”等无位置行时，从上方节标题回推单位置
            for pi in range(p_min - 1, max(p_min - 6, -1), -1):
                t = str(texts[pi]) if 0 <= pi < len(texts) else ""
                # 标题即使含“如下图所示”（如 “1/2主跨…结构温度时程曲线图、频率分布直方图如下图所示:”）
                # 也是标题，应作为位置来源
                if re.match(r"^\d+(?:\.\d+){0,3}\.?\s*[^\d]", t.strip()):
                    pos = _position_from_title(t)
                    if pos:
                        locations = [pos]
                        title_fallback = True
                        if "结构温度" in t:
                            metric = "structure_temperature"
                        break
        _DIRECTION_NORM = {"左侧", "右侧", "上游", "下游", "上游侧", "下游侧",
                           "左", "右", "左x", "右x", "左y", "右y",
                           "左幅", "右幅", "上游幅", "下游幅"}
        if locations and all(_norm(str(l)) in _DIRECTION_NORM for l in locations):
            # 行标签是方向词（如 支座位移表 的“左侧/右侧”、倾角表 的“左X/右X”），
            # 表示两个特征方向，用节标题位置补全成完整位置后各自展开 时程+直方图
            base_pos = ""
            for pi in range(p_min - 1, max(p_min - 6, -1), -1):
                t = str(texts[pi]) if 0 <= pi < len(texts) else ""
                if re.match(r"^\d+(?:\.\d+){0,3}\.?\s*[^\d]", t.strip()):
                    base_pos = _position_from_title(t)
                    if base_pos:
                        if "结构温度" in t:
                            metric = "structure_temperature"
                        break
            if base_pos:
                def _compose(d):
                    dn = _norm(str(d))
                    if len(dn) == 2 and dn[1] in ("x", "y"):
                        side = "左侧" if dn[0] == "左" else "右侧"
                        return base_pos + side + dn[1].upper()
                    return base_pos + str(d)
                locations = [_compose(l) for l in locations]

        # 块内可替换段落：explicit_suffix 段 + 块范围内的 bare_caption 图题段
        block_paras = set(paras)
        for ct in cts:
            if ct.get("source") == "bare_caption" and isinstance(ct.get("paragraph"), int):
                if p_min <= ct["paragraph"] <= paras[-1]:
                    block_paras.add(ct["paragraph"])
        block_paras = sorted(block_paras)

        if locations:
            # 标题回推单位置：每个占位行（含“第2组”等额外行）都保留并带上位置，
            # 而不是折叠成“位置×图型”后丢掉多余的占位行
            if title_fallback and len(locations) == 1:
                loc0 = locations[0]
                for c in block:
                    counter[metric] = counter.get(metric, 0) + 1
                    k = ("trend" if c.get("kind") in ("time_series", "trend", "curve")
                         else "histogram")
                    uid = f"{metric}_{loc0}_{k}_{counter[metric]}"
                    targets[c["paragraph"]] = f"{{{{chart.{uid}}}}}"
                    new_entries.append({
                        "paragraph": c.get("paragraph"),
                        "kind": k,
                        "chart_id": "trend" if k == "trend" else "histogram",
                        "metric": metric,
                        "text": c.get("text", ""),
                        "source": "expanded_block",
                        "_unique_chart_id": uid,
                        "location": loc0,
                    })
                continue
            markers = []
            meta = []
            for loc in locations:
                for k in kinds:
                    counter[metric] = counter.get(metric, 0) + 1
                    uid = f"{metric}_{loc}_{k}_{counter[metric]}"
                    markers.append(f"{{{{chart.{uid}}}}}")
                    meta.append({"uid": uid, "loc": loc, "kind": k, "metric": metric})
            for i, m in enumerate(markers):
                if i < len(block_paras):
                    targets[block_paras[i]] = m
                    meta[i]["paragraph"] = block_paras[i]
                else:
                    inserts.setdefault(block_paras[-1], []).append(m)
                    # 插入段落没有独立段落索引，用锚点段落（块内最后一段），
                    # 保证运行时按节聚类时这些图表归入同一节，避免被误判为缺图重复补图
                    meta[i]["paragraph"] = block_paras[-1]
            for p in block_paras[len(markers):]:
                removed.add(p)
            for m in meta:
                new_entries.append({
                    "paragraph": m["paragraph"],
                    "kind": m["kind"],
                    "chart_id": "trend" if m["kind"] == "trend" else "histogram",
                    "metric": m["metric"],
                    "text": "",
                    "source": "expanded_block",
                    "_unique_chart_id": m["uid"],
                    "location": m["loc"],
                })
        else:
            # 找不到表格：保持原格式（每个占位行一个 marker）
            for c in block:
                counter[metric] = counter.get(metric, 0) + 1
                k = ("trend" if c.get("kind") in ("time_series", "trend", "curve")
                     else "histogram")
                uid = f"{metric}_{k}_{counter[metric]}"
                targets[c["paragraph"]] = f"{{{{chart.{uid}}}}}"
                new_entries.append({
                    "paragraph": c.get("paragraph"),
                    "kind": c.get("kind"),
                    "chart_id": c.get("chart_id"),
                    "metric": metric,
                    "text": c.get("text", ""),
                    "source": "explicit_suffix",
                    "_unique_chart_id": uid,
                    "location": "",
                })

    # 相关性散点图图题 -> scatter 占位符（原文明确提到就要插入）
    for ct in cts:
        if ct.get("source") != "bare_caption":
            continue
        cap = str(ct.get("text") or "")
        if "相关性散点图" not in cap:
            continue
        p = ct.get("paragraph")
        if not isinstance(p, int) or p in targets or p in removed:
            continue
        pos = _scatter_position_from_caption(cap)
        if not pos:
            for pi in range(p - 1, max(p - 6, -1), -1):
                t = str(texts[pi]) if 0 <= pi < len(texts) else ""
                if ("如下图所示" not in t
                        and re.match(r"^\d+(?:\.\d+){0,3}\.?\s*[^\d]", t.strip())):
                    pos = _position_from_title(t)
                    if pos:
                        break
        if not pos:
            continue
        # 顶板/底板/左幅/右幅等是同一特征下的监测位置，不是不同的特征变量：
        # 位置之间没有相关性可画（如 应变监测节 顶板/底板 不生成 应变-温度 散点图），
        # 只有当同一位置有两个特征变量（如 温度-湿度）时才需要相关性图
        if any(w in pos for w in ("顶板", "底板", "左幅", "右幅", "腹板", "翼板", "侧板")):
            continue
        metric = _metric_from_chart_text(cap) or "chart"
        counter[metric] = counter.get(metric, 0) + 1
        uid = f"{metric}_{pos}_scatter_{counter[metric]}"
        targets[p] = f"{{{{chart.{uid}}}}}"
        new_entries.append({
            "paragraph": p, "kind": "scatter", "chart_id": "scatter",
            "metric": metric, "text": cap, "source": "expanded_scatter",
            "_unique_chart_id": uid, "location": pos,
        })
    return targets, inserts, new_entries, removed


def _special_strain_positions(texts, pidx):
    """从“4#、5#墩底部结构应变监测”类标题提取两个墩底位置的原文（保留实际字符）。"""
    for i in range(pidx - 1, max(pidx - 6, -1), -1):
        t = str(texts[i]) if 0 <= i < len(texts) else ""
        m = re.search(r"(\d+#)\u3001(\d+#)(.+?)\u7ed3\u6784\u5e94\u53d8\u76d1\u6d4b", t)
        if m:
            suffix = m.group(3)
            return m.group(1) + suffix, m.group(2) + suffix
    return "4#\u58a9\u5e95\u90e8", "5#\u58a9\u5e95\u90e8"


def _flatten_omml(p_el) -> None:
    """把段落里的 OMML 数学公式（如 m/s2）转成纯文本并原位替换。

    python-docx 的 paragraph.text 不含 OMML；模板替换 run 时公式节点可能悬空，
    导致单位渲染到段落外面（如 “振动绝对最大值为[I-MAX]。” + 孤立的 m/s2）。

    - 原位替换：公式在句中什么位置，单位文本就留在什么位置
      （如 “绝对最大值为[D-MAX]m/s2，…” 单位紧跟数值，不跑到句末）。
    - 表头形如 “绝对最大值（）” + m/s2 时，把单位塞回空括号内：
      “绝对最大值（m/s²）”。
    - 上标统一转 Unicode 上标（m/s2 -> m/s²），保持 Word 公式的观感。
    """
    M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    ommls = [c for c in p_el if c.tag == M + "oMath"]
    for omp in p_el.findall(M + "oMathPara"):
        om = omp.find(M + "oMath")
        if om is not None and om not in ommls:
            ommls.append(om)
    if not ommls:
        return
    from docx.oxml import OxmlElement as _OE
    import copy as _copy

    for om in ommls:
        unit = "".join((t.text or "") for t in om.iter(M + "t"))
        parent = om.getparent()
        if parent is None:
            continue
        if not unit:
            parent.remove(om)
            continue
        unit_txt = _pretty_unit_text(unit)
        prev = om.getprevious()
        # 表头 “绝对最大值（）” + 公式：单位塞回空括号内
        if prev is not None:
            prev_t = prev.find(W + "t")
            if prev_t is not None and prev_t.text:
                m = re.search(r"([（(])\s*([)）])\s*$", prev_t.text)
                if m:
                    prev_t.text = prev_t.text[:m.start()] + m.group(1) + unit_txt + m.group(2)
                    parent.remove(om)
                    continue
        r = _OE("w:r")
        prev_rpr = prev.find(W + "rPr") if prev is not None else None
        if prev_rpr is not None:
            r.append(_copy.deepcopy(prev_rpr))
        t = _OE("w:t")
        t.text = unit_txt
        r.append(t)
        parent.replace(om, r)


def _pretty_unit_text(unit: str) -> str:
    """单位里的数字转 Unicode 上标：m/s2 -> m/s²、m/s^2 -> m/s²。"""
    sup = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
           "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹"}
    u = str(unit or "").replace("^", "")
    return re.sub(r"(?<=[A-Za-z0-9/·])-?(\d+)",
                  lambda m: "".join(sup.get(c, c) for c in m.group(1)), u)


def _collect_unit_inserts(p) -> List[Tuple[int, str]]:
    """返回段落里 oMath 单位在【无单位文本】中的插入偏移与美化文本。

    flatten 会把 oMath 转成普通 run，使原文坐标整体后移；记录每个公式在
    原文本中的偏移，供主路径的 num/tag/cell 位置编辑统一平移修正。
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    out = []
    off = 0
    for child in p._p:
        tag = child.tag
        if tag == W + "r":
            t = child.find(W + "t")
            if t is not None and t.text:
                off += len(t.text)
        elif tag == M + "oMath":
            unit = "".join((x.text or "") for x in child.iter(M + "t"))
            if unit:
                out.append((off, _pretty_unit_text(unit)))
        elif tag == M + "oMathPara":
            om = child.find(M + "oMath")
            if om is not None:
                unit = "".join((x.text or "") for x in om.iter(M + "t"))
                if unit:
                    out.append((off, _pretty_unit_text(unit)))
    return out


def annotate_docx(src: str, dst: str, llm_cfg: Optional[dict] = None,
                  analysis: Optional[dict] = None,
                  sensor_map: Optional[dict] = None) -> dict:
    """把识别出的动态项改写为占位符，生成标注草稿（仅 DOCX）。

    - 动态数字 -> {{stats.<指标>.<统计>}}（无法映射时 {{data.N}}）
    - 动态图片 -> 该段替换为 {{chart.<图表ID>}}（按图题关键词推测，可手动改名）
    - 图表文本占位（xxx_time_series / xxx_histogram 行）-> {{chart.<ID>}}
    - 季度/时间表述 -> {{date.period_label}} / {{date.period_label_cn}}
    - 固定数字 / 固定图片（CAD 图等）原样保留

    llm_cfg: LLM 配置字典，传入则启用 LLM 二次筛选
    analysis: 可传入已计算好的识别结果，避免重复调用 LLM（内部会重新 recognize）
    sensor_map: 传感器对照表 {编号: {...}}，用于“编号(特征)_图型”行反查
        监测部位并生成位置化图表占位符（避免 chart_sensor_<编号>_<图型>）。
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from .report_builder import iter_block_items

    if analysis is None:
        analysis = recognize(src, llm_cfg=llm_cfg, sensor_map=sensor_map)
    elif sensor_map and not analysis.get("sensor_map"):
        analysis["sensor_map"] = sensor_map
    doc = Document(src)

    num_targets = {}   # paragraph index -> {pos: marker}
    minus_positions = {}   # paragraph index -> set(数字起始位置)：该数字前导 '-' 一并替换
    data_counter = [0]
    data_values = {}   # data.N -> original value (for fallback resolution)
    texts_all = analysis.get("texts", []) or []

    # 故障/异常位置判断句 -> {{summary.<metric>}} 占位符
    summary_targets = detect_summary_placeholders(texts_all)
    if summary_targets:
        analysis["summary_placeholders"] = [
            {
                "placeholder": marker,
                "metric": marker.strip("{}").split(".", 1)[1],
                "paragraph": idx,
                "snippet": str(texts_all[idx])[max(0, s - 20):e + 20],
            }
            for idx, spans in summary_targets.items()
            for s, e, marker in spans
        ]

    def _canon_placeholder(raw: str) -> str:
        """把 LLM/启发式给出的 stats 键规范成运行时指标键。"""
        m = re.match(r"^\{\{stats\.([a-zA-Z_]+)\.([a-zA-Z_]+)\}\}$", raw)
        if m:
            mm = {"structural_temp": "structure_temperature"}.get(m.group(1), m.group(1))
            return f"{{{{stats.{mm}.{m.group(2)}}}}}"
        return raw

    def _dash_sign(para_idx, pos, value) -> bool:
        """判断数字前导 '-' 是否为负号（'-' 前是中文/℃/% 等量词）。
        传感器编号（如 6RX(S)-22）前是 ')'，不吸收。"""
        if not isinstance(para_idx, int) or not isinstance(pos, int) or pos <= 0:
            return False
        t = str(texts_all[para_idx]) if 0 <= para_idx < len(texts_all) else ""
        if pos > len(t) or t[pos - 1] != "-":
            return False
        prev = t[pos - 2] if pos >= 2 else ""
        return bool(prev and ("\u4e00" <= prev <= "\u9fff" or prev in "℃%％ "))

    # 总结/结论段里 data.N 升级为 stats.* 的白名单指标（有真实传感器统计，
    # 车辆数/荷载等无统计值的指标不升级，保持原文回填）
    UPGRADE_METRICS = {
        "temperature", "humidity", "structure_temperature", "wind_speed",
        "strain", "deflection", "rotation", "displacement", "bearing_displacement",
        "cable_force", "crack",
    }

    def _try_upgrade_stats(para_idx, pos, value, marker) -> str:
        """“最高/最低/最大/最小 + 指标词”的 data.N 数字升级为 stats.*（重算），
        避免总结段出现“最高值回填旧值、最低值重算新值”的混用。"""
        if not marker.startswith("{{data."):
            return marker
        if not isinstance(para_idx, int) or not isinstance(pos, int):
            return marker
        t = str(texts_all[para_idx]) if 0 <= para_idx < len(texts_all) else ""
        if not t:
            return marker
        ctx = t[max(0, pos - 30):pos + len(value) + 30]
        if not any(s in ctx for s in ("最高", "最低", "最底", "最大", "最小", "平均", "均值")):
            return marker
        # 指标词只在同一句里找（结论段常见“索力最大值为…，对应测点位置为…；最小值为”
        # 长句，指标词离数字很远；但跨句的“结构温度…，地震…”不能串扰）
        sent_start = max(t.rfind("。", 0, pos), t.rfind("！", 0, pos),
                         t.rfind("？", 0, pos)) + 1
        sentence = t[sent_start:pos]
        semi = sentence.rfind("；")
        if semi != -1:
            clause = sentence[semi + 1:]
            clause_metric = _nearest_word(clause, len(clause), 60, METRIC_WORDS,
                                          skip={"load", "vehicle_count"})
            if clause_metric:
                sentence = clause
        best_metric = _nearest_word(sentence, len(sentence), 60, METRIC_WORDS,
                                    skip={"load", "vehicle_count"})
        best_stat = _nearest_word(sentence, len(sentence), 12, STAT_WORDS)
        if not best_metric or not best_stat or best_metric not in UPGRADE_METRICS:
            return marker
        return f"{{{{stats.{best_metric}.{best_stat}}}}}"

    def _fix_stat_key(para_idx, pos, marker) -> str:
        """按数字前的上下文纠正 stats 占位符的 max/min/range 键。
        LLM 可能把“最低湿度”错标成 max，这里用最近统计词兜底纠错。"""
        if not marker.startswith("{{stats."):
            return marker
        m = re.match(r"^\{\{stats\.([a-zA-Z_]+)\.([a-zA-Z_]+)\}\}$", marker)
        if not m or not isinstance(para_idx, int) or not isinstance(pos, int):
            return marker
        metric, stat = m.group(1), m.group(2)
        t = str(texts_all[para_idx]) if 0 <= para_idx < len(texts_all) else ""
        win = t[max(0, pos - 12):pos]
        if "最大差值" in win or "最小差值" in win:
            expect = "range"
        elif "最高" in win or "最大" in win:
            expect = "max"
        elif "最低" in win or "最底" in win or "最小" in win:
            expect = "min"
        else:
            return marker
        if stat != expect:
            return f"{{{{stats.{metric}.{expect}}}}}"
        return marker

    # 第一遍：LLM 补充识别（value 自带负号，如 "-2.08"）优先占用数字起始位置
    dash_wins = {}   # (paragraph, 数字起始位置) -> marker
    for n in analysis["numbers"]:
        if n["verdict"] != "replace":
            continue
        value = str(n.get("value", ""))
        pos = n.get("position")
        if not value.startswith("-") or not isinstance(pos, int):
            continue
        digit_pos = pos + 1
        marker = _canon_placeholder(n.get("placeholder") or "")
        dash_wins[(n.get("paragraph"), digit_pos)] = marker

    for n in analysis["numbers"]:
        if n["verdict"] != "replace":
            continue
        marker = n.get("placeholder")
        if marker is None:
            data_counter[0] += 1
            data_key = f"data.{data_counter[0]}"
            marker = f"{{{{data.{data_counter[0]}}}}}"
        elif not marker.startswith("{{"):
            # 统一补花括号（suggest_placeholder 返回无花括号的 key）
            marker = f"{{{{{marker}}}}}"
        else:
            marker = _canon_placeholder(marker)
        para_idx = n.get("paragraph")
        pos = n.get("position")
        value = str(n.get("value", ""))
        # 带负号的 LLM 条目：数字实际起始位置 = pos+1，吸收前导 '-'；
        # 同时跳过同位置普通条目，避免 data.N 覆盖 stats 占位符
        if value.startswith("-") and isinstance(pos, int):
            pos = pos + 1
            absorb = True
            marker = dash_wins.get((para_idx, pos), marker)
            marker = _fix_stat_key(para_idx, pos, marker)
        else:
            marker = _try_upgrade_stats(para_idx, pos, value, marker)
            marker = _fix_stat_key(para_idx, pos, marker)
            absorb = isinstance(pos, int) and _dash_sign(para_idx, pos, value)
        if isinstance(para_idx, int) and isinstance(pos, int):
            if (para_idx, pos) in dash_wins and not value.startswith("-"):
                continue  # 已被带负号的 LLM 条目占用
            num_targets.setdefault(para_idx, {})[pos] = marker
            if absorb:
                minus_positions.setdefault(para_idx, set()).add(pos)
            # data.N 原始值（负号吸收时存带负号的值；升级为 stats 后不再回填）
            if marker.startswith("{{data.") and (para_idx, pos) not in dash_wins:
                data_key = marker.strip("{}")
                data_values.setdefault(data_key, f"-{value}" if absorb else value)

    # 将 data_values 存入 analysis，供后续 report_builder 使用
    analysis["data_values"] = data_values

    # 原文占位标记（[A-MAX]、A-MAX-LOC 等）→ 占位符
    tag_targets = {}   # paragraph index -> {pos: (end, marker)}
    for tag in analysis.get("tags", []) or []:
        if tag.get("verdict") != "replace":
            continue
        marker = tag.get("placeholder")
        if not marker:
            # 推断不到指标/统计量的标记：用 data.N 回填（保留原标记文本，血缘日志标注未重算）
            data_counter[0] += 1
            data_key = f"data.{data_counter[0]}"
            data_values[data_key] = tag.get("value", "")
            marker = f"{{{{data.{data_counter[0]}}}}}"
        pp = tag.get("paragraph")
        pos = tag.get("position")
        end = tag.get("end")
        if isinstance(pp, int) and isinstance(pos, int) and isinstance(end, int):
            tag_targets.setdefault(pp, {})[pos] = (end, marker)

    img_targets = {}   # paragraph index -> [(image_index, marker)]
    for im in analysis["images"]:
        if im["verdict"] != "replace":
            continue
        fallback = f"chart_{im['index']}"
        marker = f"{{{{chart.{_suggest_chart_id(im.get('caption'), fallback)}}}}}"
        img_targets.setdefault(im["paragraph"], []).append((im["index"], marker))

    # 图表文本占位：xxx_time_series / xxx_histogram 行 / 纯图题 -> {{chart.<id>}}
    # 表格单元格引用：H(行,列) / M(行,列) / J(行,列) -> {{cell.<metric>.<column>.<stat>}}
    chart_text_targets = {}   # paragraph index -> marker
    chart_text_counter = {}
    cell_ref_targets = {}     # paragraph index -> {pos: marker}
    vehicle_cell_targets = {} # paragraph index -> marker（整段替换）
    cell_seq_rows = {}        # (metric, table_letter, table_title, row_label) -> 已出现的行号集合
    cell_ref_paras: Dict[int, tuple] = {}   # paragraph -> (表标题, row_label)，用于图表占位符位置推断
    # bare_caption 图题文本段（仅做识别提示，不替换为占位符；由 build_report 阶段补上编号图注）
    bare_caption_paras: set = set()

    # “编号(特征)_图型”行（如 184(xJsd)_时程曲线）：
    # 直接给出 传感器编号 + 特征 + 图型，识别为精确图表占位符
    CHART_LINE_RE = re.compile(
        r"^\s*(\d{1,5})\s*\(([^()]+)\)\s*_(时程曲线|频率分布直方图|时间序列图|时间序列|直方图|时程|曲线)\s*$"
    )
    chart_line_entries = []    # 额外写入 analysis.chart_texts 的条目
    cell_ref_columns_map = {
        # 表格列 → 统计类型
        # H 表 (索夹)：列2=avg, 列3=max, 列4=min, 列5=abs_max, 列6=rms, 列7=range
        "H": {2: "avg", 3: "max", 4: "min", 5: "abs_max", 6: "rms", 7: "range"},
        # M 表 (转角)：列2=avg, 列3=max, 列4=min, 列5=abs_max, 列6=rms, 列7=range
        "M": {2: "avg", 3: "max", 4: "min", 5: "abs_max", 6: "rms", 7: "range"},
        # J 表 (风速)：列2=avg, 列3=max, 列4=min, 列5=abs_max, 列6=rms, 列7=range
        # 注意：实际文档中 J 表无列5，但保持一致以便扩展
        "J": {2: "avg", 3: "max", 4: "min", 5: "abs_max", 6: "rms", 7: "range"},
    }
    # 行 → 测点列名（根据用户实际表格：第 2 行起为数据行）
    cable_clamp_row_to_col = {
        2: "h_87L_1", 3: "h_87R_1", 4: "h_88L_1", 5: "h_88R_1",
    }
    # 索夹第 2 组（如有）
    cable_clamp_row2_to_col = {
        6: "h_87L_2", 7: "h_87R_2", 8: "h_88L_2", 9: "h_88R_2",
    }
    rotation_row_to_col = {
        2: "m_junshan_x", 3: "m_junshan_y", 4: "m_yueyang_x", 5: "m_yueyang_y",
    }
    # J 表（风速）按列号取：col 2=avg, 3=max, 4=min, 6=rms, 7=range
    # 行号 2-5 对应：君山塔顶、岳阳塔顶、1/2 上游、1/2 下游
    wind_row_to_col = {
        2: "j_junshan_top", 3: "j_yueyang_top",
        4: "j_half_upstream", 5: "j_half_downstream",
    }

    # 通用行映射：表格第 N 行（从2开始） → 第 N-1 个测点列
    metric_point_map = {
        "cable_clamp": {**cable_clamp_row_to_col, **cable_clamp_row2_to_col},
        "rotation": rotation_row_to_col,
        "wind_speed": wind_row_to_col,
    }

    # 预扫描：先收集全部 cell_ref 段落的 (表标题, row_label)，
    # 供图表占位符做“位置关联”（图表段在其后表格的上方，遍历顺序不一定按段落序）
    for ct in analysis.get("chart_texts", []):
        if ct.get("source") == "cell_ref":
            p = ct.get("paragraph")
            rl = str(ct.get("row_label") or "").strip()
            if isinstance(p, int) and rl and not rl.startswith("测点"):
                cell_ref_paras[p] = (str(ct.get("table_title") or ""), rl)

    # 节标题段落索引（数字编号标题），用于图表位置推断的“不跨节”约束
    _texts_all = analysis.get("texts", []) or []
    heading_paras = [
        i for i, t in enumerate(_texts_all)
        if re.match(r"^\d+(?:\.\d+){0,3}\.?(?=[\u4e00-\u9fa5\s])", str(t).strip())
        and len(str(t).strip()) <= 60
    ]

    # “编号(特征)_图型”行已由 parse_docx 转为 explicit_suffix，
    # 由 _expand_chart_blocks 用下方表格的监测部位（上游/下游等）展开成位置化占位符，
    # 不再生成 chart_sensor_<编号>_<图型>。
    cleared_sensor_lines = set()

    # 图表文本占位行按“监测部位 × 图型”展开（带位置），并更新 chart_texts 条目
    chart_block_targets, chart_block_inserts, chart_block_entries, chart_block_removed = \
        _expand_chart_blocks(analysis, cell_ref_paras, heading_paras)
    chart_text_targets.update(chart_block_targets)
    if chart_block_entries:
        kept = [ct for ct in analysis.get("chart_texts", [])
                if ct.get("source") != "explicit_suffix"]
        kept.extend(chart_block_entries)
        analysis["chart_texts"] = kept

    for ct in analysis.get("chart_texts", []):
        if ct.get("source") == "cell_ref":
            # 静态表的 cell_ref 不替换（保持原值）
            if ct.get("verdict") == "keep":
                continue
            # 表格单元格引用占位符
            metric = ct.get("metric", "")
            row = ct.get("row", 0)
            col = ct.get("col", 0)
            table_letter = ct.get("table_letter", "")
            # 优先用识别阶段记录的 row_label / col_header（更稳定）
            row_label = ct.get("row_label", "")
            col_header = ct.get("col_header", "")
            row_slug = slugify_label(row_label) or f"r{row}"
            col_slug = slugify_label(col_header) or f"c{col}"
            # 列 → stat：用列头优先，再用 letter 映射兜底
            stat = stat_from_col_header(col_header)
            if stat == "value" and col_header == "":
                stat_map = cell_ref_columns_map.get(table_letter, {})
                stat = stat_map.get(col, f"col{col}")
            # 行 → 测点列名：先用识别阶段的 row_slug，再用 letter 模板映射
            row_map = metric_point_map.get(metric, {})
            column_from_map = row_map.get(row)
            if column_from_map and not row_label:
                column = column_from_map
            elif column_from_map and row_slug in column_from_map:
                column = column_from_map
            else:
                column = row_slug or column_from_map or f"unknown_{metric}_r{row}"
            # 结合表格标题/上下文补全位置：
            # - 行标签是“顶板测点1/底板测点3”等部位词时，从标题取位置基座拼接
            #   （如 标题“上游随州侧边跨跨中箱梁结构温度监测统计”+“顶板测点1”
            #     -> column=“上游随州侧边跨跨中箱梁顶板测点1”）
            # - 行标签已是完整位置（含 截面/箱梁/墩/跨 等）时保留，仅在
            #   标题带“上游/下游/左/右”方位且行标签缺方位时补方位。
            table_title_ctx = str(ct.get("table_title") or "")
            if row_label and not column_from_map:
                base_pos = _cell_position_from_title(table_title_ctx)
                rn = _norm(row_label)
                loc_words = ("截面", "箱梁", "墩", "跨", "断面", "梁段",
                             "索塔", "塔", "锚固", "桥面")
                has_loc = any(w in rn for w in loc_words)
                if not has_loc:
                    # 行标签只是“顶板测点1/底板测点3/测点2”时拼位置基座
                    if base_pos:
                        column = base_pos + row_label
                    else:
                        # 表格标题太泛（如“结构温度监测统计”）时，
                        # 从上方最近的小节标题继承位置
                        # （如 3.2.1.5上游湘潭侧中跨1/4箱梁结构温度监测）
                        para_idx = ct.get("paragraph")
                        if isinstance(para_idx, int):
                            _texts = analysis.get("texts", []) or []
                            # 先定位表格标题段（数据 cell 段落可能在表格内部，
                            # 直接向上扫不到节标题），再向上找最近的小节标题
                            start = para_idx - 1
                            if table_title_ctx:
                                for pi in range(para_idx - 1,
                                                max(para_idx - 40, -1), -1):
                                    _tt = str(_texts[pi]) if 0 <= pi < len(_texts) else ""
                                    if _tt.strip() == table_title_ctx.strip():
                                        start = pi - 1
                                        break
                            for pi in range(start,
                                            max(start - 30, -1), -1):
                                _t = str(_texts[pi]) if 0 <= pi < len(_texts) else ""
                                if not re.match(r"^\d+(?:\.\d+){0,3}",
                                                _t.strip()):
                                    continue
                                _bp = _cell_position_from_title(_t)
                                if _bp:
                                    column = _bp + row_label
                                    break
                else:
                    # 行标签已有位置，标题带方位且行标签缺方位时补方位
                    for side in ("上游", "下游", "左侧", "右侧", "左", "右"):
                        if side in _norm(table_title_ctx) and side not in rn:
                            column = side + row_label
                            break
            # 同一表格同一位置的多个测点行：占位符统一加 #N 行号索引
            # （如 ...顶板测点1.avg#1 / ...顶板测点2.avg#2），运行时按 #N
            # 精确取对应传感器。
            # 按“表格 + 位置基座”分组计数：位置基座去掉“测点N”后缀，
            # 让 顶板测点1/顶板测点2/底板测点1... 在同一位置组内按行号递增；
            # 不能含完整 column/row_label，否则每行都从 #1 开始，
            # 同一位置的多个测点会取到同一个传感器。
            # 位置基座保留“顶板/底板/腹板/翼板”等部位词（顶板与底板是
            # 表格映射里的不同位置组，row_index 须各自从 0 开始），
            # 只去掉“测点N”后缀
            col_base = re.sub(r"测点\s*\d+$", "", column).strip(" 、，,和及") \
                or column
            seq_key = (metric, table_letter, ct.get("table_title", ""),
                       col_base)
            rows = cell_seq_rows.setdefault(seq_key, set())
            rows.add(row)
            seq = len(rows)
            ct["_cell_seq"] = seq
            suffix = f"#{seq}"
            marker = f"{{{{cell.{metric}.{column}.{stat}{suffix}}}}}"
            # 车辆累计表：整段是一个数字，直接整段替换为占位符
            if ct.get("vehicle"):
                vehicle_cell_targets[ct["paragraph"]] = marker
                continue
            cell_ref_targets.setdefault(ct["paragraph"], {})[ct.get("position", 0)] = marker
            if row_label and not row_label.startswith("测点"):
                cell_ref_paras[ct["paragraph"]] = (str(ct.get("table_title") or ""), row_label)
        else:
            # 图表文本占位符
            # explicit_suffix / expanded_block 已由 _expand_chart_blocks
            # 按“监测部位×图型”展开，不能再重复编号
            if ct.get("source") in ("explicit_suffix", "expanded_block",
                                    "expanded_scatter", "special_strain"):
                continue
            # bare_caption 是图题文本（如“跨中断面环境温度时程曲线图”），源头 DOCX
            # 在 _time_series / _histogram 占位符之后紧接着就是图题段。annotate_docx
            # 不应再为图题段生成 {{chart.*}} 占位符——否则 build_report 会重复插入
            # 同一张图，导致图片覆盖文字。
            if ct.get("source") == "bare_caption":
                if ct["paragraph"] in chart_block_targets:
                    continue  # 已被图表块展开占用（该段将替换为带位置的占位符）
                bare_caption_paras.add(ct["paragraph"])
                continue
            metric = ct.get("metric", "chart")
            cid = ct.get("chart_id", "trend")
            chart_text_counter[metric] = chart_text_counter.get(metric, 0) + 1
            # 图表占位符带监测位置：与下方表格的传感器位置关联（单位置时写入）
            loc = _infer_chart_location(ct.get("paragraph"), cell_ref_paras,
                                        heading_paras=heading_paras)
            unique_cid = (f"{metric}_{loc}_{cid}_{chart_text_counter[metric]}" if loc
                          else f"{metric}_{cid}_{chart_text_counter[metric]}")
            chart_text_targets[ct["paragraph"]] = f"{{{{chart.{unique_cid}}}}}"
            ct["_unique_chart_id"] = unique_cid  # 供运行时 agent 引用
            if loc:
                ct["location"] = loc

    # 特殊应变行：特殊应变传感器_修正后时程曲线_2x2 / _频率分布_2x2
    # -> 4#、5#墩底部应变图占位符（图库有 4#/5#墩底部 YB(rsg) 合并图）
    special_strain_paras = {}   # pidx -> [第2个位置占位符]
    for pidx, t in enumerate(analysis.get("texts", []) or []):
        m = re.match(r"^\s*特殊应变传感器_修正后(时程曲线|频率分布)_2x2\s*$", str(t))
        if not m:
            continue
        kind = "trend" if "时程" in m.group(1) else "histogram"
        pos4, pos5 = _special_strain_positions(analysis.get("texts", []) or [], pidx)
        cid4 = f"strain_{pos4}_{kind}_1"
        cid5 = f"strain_{pos5}_{kind}_1"
        chart_text_targets[pidx] = f"{{{{chart.{cid4}}}}}"
        special_strain_paras[pidx] = [f"{{{{chart.{cid5}}}}}"]
        analysis.setdefault("chart_texts", []).extend([
            {
                "paragraph": pidx,
                "kind": "time_series" if kind == "trend" else "histogram",
                "chart_id": kind, "metric": "strain", "text": str(t),
                "source": "special_strain",
                "_unique_chart_id": cid4, "location": pos4,
            },
            {
                # 5# 占位符是插入段落，运行时同样需要条目（paragraph 指向同一节）
                "paragraph": pidx,
                "kind": "time_series" if kind == "trend" else "histogram",
                "chart_id": kind, "metric": "strain", "text": str(t),
                "source": "special_strain",
                "_unique_chart_id": cid5, "location": pos5,
            },
        ])
    if special_strain_paras:
        log.info("识别 %d 行特殊应变图表（4#/5#墩底部应变）", len(special_strain_paras))

    # 文本级替换（季度/时间表述动态化）
    text_targets = {}   # paragraph index -> [(original, replacement)]
    # 默认替换规则（不依赖 LLM，始终应用）
    default_text_replacements = [
        ("XXXX年第X季度（xx月-xx月）", "{{date.period_label_cn}}（{{date.period_label}}）"),
        ("XXXX年第X季度", "{{date.period_label_cn}}"),
        # 落款日期（封面/签字页）
        ("xxxx年xx月xx日", "{{date.signature}}"),
        ("XXXX年XX月XX日", "{{date.signature}}"),
        ("xxxx年XX月XX日", "{{date.signature}}"),
        ("XXXX年xx月xx日", "{{date.signature}}"),
        # 结论段报告期范围
        ("xx月-xx月", "{{date.period_start_month}}-{{date.period_end_month}}"),
    ]
    all_text_replacements = default_text_replacements + list(
        analysis.get("text_replacements", [])
    )
    for tr in all_text_replacements:
        if isinstance(tr, dict):
            original = tr.get("original", "")
            replacement = tr.get("replacement", "")
        else:
            original, replacement = tr
        if not original or not replacement:
            continue
        for i, t in enumerate(analysis.get("texts", []) or []):
            if original in t:
                text_targets.setdefault(i, []).append((original, replacement))

    replaced_numbers = 0
    skipped_numbers = 0
    replaced_images = 0
    replaced_texts = 0
    replaced_chart_texts = 0
    replaced_cell_refs = 0
    replaced_summaries = 0
    pending_inserts = []   # [(anchor 段落元素, [marker, ...])] 遍历后统一插入

    def _clean_chart_para(para) -> None:
        """去掉图名(af3)样式并居中，避免图片段样式不一致导致 Word 渲染重叠。"""
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            st = pPr.find(qn("w:pStyle"))
            if st is not None:
                pPr.remove(st)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def process_paragraph(p, idx):
        nonlocal replaced_numbers, skipped_numbers, replaced_images, replaced_texts, replaced_chart_texts, replaced_cell_refs, replaced_summaries
        if idx in cleared_sensor_lines:
            # 清空整段前先转换公式节点，避免 oMath 悬空渲染
            _flatten_omml(p._p)
            _clear_paragraph_text(p)
            return
        # 图表块展开后多余的占位行/图题段：清空
        if idx in chart_block_removed:
            _flatten_omml(p._p)
            _clear_paragraph_text(p)
            return
        # bare_caption 图题段：原 caption 文本会被 build_report 的 auto-caption 取代，
        # 清空段落文本，避免重复 caption
        if idx in bare_caption_paras:
            _flatten_omml(p._p)
            _clear_paragraph_text(p)
            return
        # 车辆累计表单元格：整段数字 → {{cell.vehicle_count.车道X.数值#N}}
        if idx in vehicle_cell_targets:
            _flatten_omml(p._p)
            _replace_whole_paragraph(p, vehicle_cell_targets[idx])
            replaced_cell_refs += 1
            return
        # 图表文本占位优先（整段替换为 {{chart.<id>}}）
        if idx in chart_text_targets:
            _flatten_omml(p._p)
            _replace_whole_paragraph(p, chart_text_targets[idx])
            _clean_chart_para(p)
            replaced_chart_texts += 1
            # 同一锚点段要插入的多个占位符必须合并成一条，应用时 reversed()
            # 才会还原成原始顺序；拆成多条会逐条插到锚点段之后导致顺序反转。
            extra_markers = (list(special_strain_paras.get(idx, []))
                             + list(chart_block_inserts.get(idx, [])))
            if extra_markers:
                pending_inserts.append((p._p, extra_markers))
            return
        # 数字 / 原文标记 / 表格单元格引用：合并成一次位置替换，
        # 避免同一段多次替换后位置错位（坐标均基于原文）
        # OMML 公式（如 m/s2）flatten 成普通 run 会让原文坐标后移：
        # 先记录公式偏移，flatten 后统一平移 num/tag/cell 的位置。
        unit_inserts = _collect_unit_inserts(p)
        _flatten_omml(p._p)
        num_t = num_targets.get(idx, {})
        tag_t = tag_targets.get(idx, {})
        cell_t = cell_ref_targets.get(idx, {})
        summary_t = summary_targets.get(idx, [])
        if num_t or tag_t or cell_t or summary_t:
            full = "".join(r.text for r in p.runs)
            if unit_inserts:
                # 起始坐标：字符本身在单位之后则平移；结束边界：单位恰好在
                # 边界处（如 “[D-MAX]m/s²” 的 end=37）时边界在单位之前，不平移。
                def _shift_start(pos):
                    return pos + sum(len(u) for off, u in unit_inserts if off <= pos)

                def _shift_end(pos):
                    return pos + sum(len(u) for off, u in unit_inserts if off < pos)
                num_t = {_shift_start(k): v for k, v in num_t.items()}
                tag_t = {_shift_start(k): (_shift_end(e), m) for k, (e, m) in tag_t.items()}
                cell_t = {_shift_start(k): v for k, v in cell_t.items()}
                summary_t = [(_shift_start(s), _shift_end(e), m)
                             for s, e, m in summary_t]
                if minus_positions.get(idx):
                    minus_positions[idx] = {_shift_start(k) for k in minus_positions[idx]}
            edits = []
            if num_t:
                num_t = _filter_location_numbers(full, num_t)
                minus = minus_positions.get(idx, set())
                for m in NUMBER_RE.finditer(full):
                    if m.start() not in num_t:
                        continue
                    absorb = 1 if (m.start() in minus and m.start() > 0
                                   and full[m.start() - 1] == "-") else 0
                    edits.append((m.start() - absorb, m.end(), num_t[m.start()]))
                    replaced_numbers += 1
            if tag_t:
                for s, (e, marker) in tag_t.items():
                    if 0 <= s < e <= len(full):
                        edits.append((s, e, marker))
                        replaced_numbers += 1
            if cell_t:
                for m in CELL_REF_RE.finditer(full):
                    if m.start() in cell_t:
                        edits.append((m.start(), m.end(), cell_t[m.start()]))
                        replaced_cell_refs += 1
            # 故障/异常位置判断句整体替换为 {{summary.<metric>}}
            # （区间起点早于句内数字，_apply_position_edits 会先应用并覆盖内部编辑）
            for s, e, marker in summary_t:
                if 0 <= s < e <= len(full):
                    edits.append((s, e, marker))
                    replaced_summaries += 1
            _apply_position_edits(p, edits, full)
        if img_targets.get(idx):
            _replace_image_paragraph(p, img_targets[idx][0][1])
            replaced_images += 1
        if text_targets.get(idx):
            for original, replacement in text_targets[idx]:
                # 年份吸收：“2025年第一季度”→{{date.period_label_cn}}（period_label_cn
                # 自带年份，前面的“2025年”应一并吞掉，避免输出“2025年2026年第1季度”）
                if replacement == "{{date.period_label_cn}}" and re.match(
                        r"^第?.{0,3}季度$", original.strip()):
                    full2 = "".join(r.text for r in p.runs)
                    m2 = re.search(r"(20\d{2}年)" + re.escape(original.strip())
                                   + r"(?![0-9])", full2)
                    if m2:
                        original = m2.group(1) + original.strip()
                if _replace_text_in_paragraph(p, original, replacement):
                    replaced_texts += 1

    idx = 0
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            process_paragraph(item, idx)
            idx += 1
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p, idx)
                        idx += 1
    for section in doc.sections:
        for p in section.header.paragraphs:
            process_paragraph(p, idx)
            idx += 1
        for p in section.footer.paragraphs:
            process_paragraph(p, idx)
            idx += 1

    # 图表块展开需要额外插入的占位段落（统一在遍历后插入，避免干扰段落迭代）
    for anchor_el, markers in pending_inserts:
        for marker in reversed(markers):
            new_p = OxmlElement("w:p")
            anchor_el.addnext(new_p)
            np = Paragraph(new_p, doc)
            np.add_run(marker)
            _clean_chart_para(np)

    os.makedirs(os.path.dirname(os.path.abspath(dst)) or ".", exist_ok=True)
    doc.save(dst)
    return {
        "output": os.path.abspath(dst),
        "replaced_numbers": replaced_numbers,
        "skipped_numbers_split_runs": skipped_numbers,
        "replaced_images": replaced_images,
        "replaced_texts": replaced_texts,
        "replaced_chart_texts": replaced_chart_texts,
        "replaced_cell_refs": replaced_cell_refs,
        "replaced_summaries": replaced_summaries,
        "llm_summary": analysis["summary"].get("llm", {}),
    }


def _replace_whole_paragraph(p, marker: str) -> None:
    """把段落整体替换为占位符文本。"""
    runs = p.runs
    if runs:
        runs[0].text = marker
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(marker)


def _clear_paragraph_text(p) -> None:
    """清空段落所有 run 的文本，保留段落本身用于间距。"""
    for r in p.runs:
        r.text = ""


def _replace_text_in_paragraph(p, original: str, replacement: str) -> bool:
    """在段落中按文本片段替换（合并 run，保留首个 run 格式）。

    返回是否替换成功。
    """
    runs = p.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    if original not in full:
        return False
    new_text = full.replace(original, replacement)
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _replace_numbers_in_paragraph(p, targets: dict, minus_positions: set = None):
    """在段落里按全局位置替换数字。

    当数字跨多个 run 时（如 '3' + '9.93' 分属不同 run），
    先合并所有 run 再替换，避免跨 run 匹配被跳过。
    minus_positions: 数字起始位置集合；命中时把数字紧邻的前导 '-' 一并替换
    （如 “最低温度为-2.08℃” -> “最低温度为{{stats.…}}℃”），
    避免运行时负值输出成 “--3.6”。
    """
    runs = p.runs
    if not runs:
        return 0, 0
    full = "".join(r.text for r in runs)
    minus_positions = minus_positions or set()
    offsets = []
    cur = 0
    for r in runs:
        offsets.append(cur)
        cur += len(r.text)

    hits = []
    for m in NUMBER_RE.finditer(full):
        if m.start() in targets:
            hits.append((m.start(), m.end(), targets[m.start()]))
    if not hits:
        return 0, 0

    # 检查是否有匹配跨越多个 run
    needs_merge = False
    for start, end, marker in hits:
        ri = max(i for i, off in enumerate(offsets) if off <= start)
        run_end = offsets[ri] + len(runs[ri].text)
        if end > run_end:
            needs_merge = True
            break

    if needs_merge:
        # 合并所有 run 到第一个 run，保留首个 run 的格式
        runs[0].text = full
        for r in runs[1:]:
            r.text = ""
        offsets = [0]
        runs = [runs[0]]

    ok = 0
    skip = 0
    for start, end, marker in sorted(hits, key=lambda x: -x[0]):
        ri = max(i for i, off in enumerate(offsets) if off <= start)
        run_end = offsets[ri] + len(runs[ri].text)
        if end > run_end:
            skip += 1
            continue
        rel = start - offsets[ri]
        # 负号吸收：识别阶段标记的 minus_positions（前导 '-' 属于负号而非编号），
        # 把 '-' 一并替换进占位符。
        absorb = 0
        if start in minus_positions and start > 0 and full[start - 1] == "-":
            absorb = 1
        runs[ri].text = (runs[ri].text[:rel - absorb] + marker
                         + runs[ri].text[rel + (end - start):])
        ok += 1
    return ok, skip


def _replace_tags_in_paragraph(p, targets: dict) -> int:
    """按绝对位置替换原文占位标记（[A-MAX] 等），targets: {start: (end, marker)}。"""
    runs = p.runs
    if not runs or not targets:
        return 0
    full = "".join(r.text for r in runs)
    offsets = []
    cur = 0
    for r in runs:
        offsets.append(cur)
        cur += len(r.text)
    hits = sorted(
        ((s, e, m) for s, (e, m) in targets.items() if 0 <= s < e <= len(full)),
        key=lambda x: -x[0],
    )
    # 标记跨 run 时先合并
    needs_merge = False
    for s, e, _m in hits:
        ri = max(i for i, off in enumerate(offsets) if off <= s)
        if e > offsets[ri] + len(runs[ri].text):
            needs_merge = True
            break
    if needs_merge:
        runs[0].text = full
        for r in runs[1:]:
            r.text = ""
        offsets = [0]
        runs = [runs[0]]
    ok = 0
    for s, e, marker in hits:
        ri = max(i for i, off in enumerate(offsets) if off <= s)
        rel = s - offsets[ri]
        runs[ri].text = runs[ri].text[:rel] + marker + runs[ri].text[rel + (e - s):]
        ok += 1
    return ok


def _apply_position_edits(p, edits: list, full: str) -> None:
    """把 (start, end, replacement) 编辑一次应用到段落全文。

    数字/标记/单元格引用混在同一段时，各自的位置基于【原文】计算；
    若分多次替换，前一次的长度变化会让后一次错位（如 标记 与
    {{stats.*}} 数字叠加后出现乱码）。统一合并成一次替换可避免。
    """
    if not edits:
        return
    # 按起点升序处理：低起点（通常区间更大，如 单元格引用 I(2,2)）先应用，
    # 内部数字等重叠编辑被跳过
    edits = sorted(edits, key=lambda e: e[0])
    pieces = []
    last = 0
    applied = []   # 已应用的区间 (start, end)，用于跳过重叠编辑
    for s, e, rep in edits:
        # 与已应用区间重叠时跳过（如 I(2,2) 内部的数字 2、2 应被整体替换覆盖）
        if any(s < ae and e > as_ for as_, ae in applied):
            continue
        pieces.append(full[last:s])
        pieces.append(rep)
        applied.append((s, e))
        last = s
        last = e
    pieces.append(full[last:])
    new_text = "".join(pieces)
    # 源报告里不成对的方括号（如 “[D-MAX ，…对应测点位置为[D-MAX-LOC]。]”）
    # 替换后留下的孤立 “[{{” / “}}]” 一并清掉
    new_text = new_text.replace("[{{", "{{").replace("}}]", "}}")
    if new_text == full:
        return
    runs = p.runs
    if not runs:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# 静态数字保护：位置名/跨号/塔号/测点数量/节标题编号中的数字不应替换成占位符。
# 同时用于 annotate 阶段剔除替换目标（_filter_location_numbers）。
STATIC_NUMBER_RE = re.compile(
    r"第\s*[\d、，,和及]+\s*跨(?:\s*跨中|\s*断面)?|"          # 第6、7、8跨跨中 / 第5、9跨跨中
    r"L\s*[\d、，,和及/]+\s*(?:断面|处)|"                      # L3/8断面 / L/4处
    r"\d+\s*L\s*[\d、，,和及/]+\s*(?:断面|处)|"                # 3L/4断面
    r"\d+\s*#\s*(?:[、，,和及]+\s*\d+\s*#\s*)*(?:索塔|塔|墩)|"  # 4#墩 / 5#、6#索塔
    r"\d+\s*号\s*(?:索塔|塔|墩)|"                              # 6号塔
    r"\d+\s*个\s*(?:GNSS监测点|GNSS测点|测点|监测点|基站|监测站)|"  # 布设2个GNSS测点 / 15个监测点
    r"共计\s*\d+\s*个|共\s*\d+\s*个|"                             # 共计1个 / 共117对
    r"车道\s*\d+|"                                                # 车道1、车道2（静态车道号）
    r"\d+\s*min(?![A-Za-z0-9])|"                                  # 10min平均风速（时间窗，非动态值）
    r"\d+\s*个\s*螺栓|"                                           # 抽取61个螺栓（检查数量，静态）
    r"\d+(?=\s*(?:主桁架|节点板))|"                               # 节点板编号（如 2774主桁架…节点板）
    r"\d+\s*[～~-]\s*\d+\s*N\s*·?\s*m|"                           # 合格范围 62～63N·m（阈值）
    r"\d+\s*N\s*·?\s*m\b|"                                        # 单值 N·m（扭矩阈值）
    r"\d+\s*m\s*处|"                                              # 30m处（监测位置）
    r"(?<![\d.])\d+\s*/\s*\d+(?=\s*(?:主跨|跨中|边跨|跨|桥面|断面|截面|钢桁|箱梁|处))|"  # 1/2主跨、1/4处、2/4跨
    r"^\d+(?:\.\d+){0,3}\.?(?=[\u4e00-\u9fa5\s])"              # 节标题编号（4.监测结论 / 3.5.1梁端倾角…）
)
LOCATION_SPAN_RE = STATIC_NUMBER_RE


def _protect_static_numbers(parsed: dict) -> int:
    """把静态数字（位置名/塔号/跨号/测点数量/节标题编号等）从替换目标中剔除。

    在 LLM 二次筛选之前调用，避免把“第6、7、8跨跨中布设2个GNSS测点”这类
    设计常量交给 LLM 或替换成占位符。
    """
    texts = parsed.get("texts") or []
    changed = 0
    for num in parsed.get("numbers", []):
        if num.get("verdict") not in ("replace", "review"):
            continue
        p = num.get("paragraph")
        if not isinstance(p, int) or not (0 <= p < len(texts)):
            continue
        p_text = str(texts[p])
        # 节标题整体保护：编号标题（3.3.1.2  2/4跨主梁截面挠度监测）里的数字
        # 都是位置/编号（跨号、墩号、1/4、30m 等），一律不替换。
        # 年份/季度若出现在标题里，由文本替换（text_replacements）另行处理，不受影响。
        if re.match(r"^\d+(?:\.\d+){0,3}\.?\s*[\u4e00-\u9fa5A-Za-z#\d]", p_text):
            num["verdict"] = "keep"
            num["confidence"] = 0.95
            num["reasons"].append("节标题数字整体保护")
            num["placeholder"] = None
            changed += 1
            continue
        pos = num.get("position", -1)
        if pos < 0:
            continue
        hit = False
        for m in STATIC_NUMBER_RE.finditer(p_text):
            if m.start() <= pos < m.end():
                hit = True
                break
        if hit:
            num["verdict"] = "keep"
            num["confidence"] = 0.92
            num["reasons"].append("静态位置/塔号/跨号/数量词/标题编号")
            num["placeholder"] = None
            changed += 1
    return changed


def _filter_location_numbers(full_text: str, targets: dict) -> dict:
    """把位置名里的数字从替换目标中剔除（不替换为 data.N 占位符）。"""
    if not targets:
        return targets
    protected = [(m.start(), m.end()) for m in LOCATION_SPAN_RE.finditer(full_text)]
    if not protected:
        return targets
    return {pos: marker for pos, marker in targets.items()
            if not any(s <= pos < e for s, e in protected)}


def _replace_cell_refs_in_paragraph(p, targets: dict):
    """在段落里按全局位置替换 H(行,列) / M(行,列) / J(行,列) 单元格引用占位符。

    targets: {start_pos: marker} 字典

    当 cell_ref 文本跨多个 run 时（如 'J' + '(2,2)' 分属不同 run），
    先合并所有 run 再替换，避免跨 run 匹配被跳过。
    """
    runs = p.runs
    if not runs:
        return 0, 0
    full = "".join(r.text for r in runs)

    # 在全文里找所有 cell_ref 引用，按位置匹配 targets
    matches = []
    for m in CELL_REF_RE.finditer(full):
        if m.start() in targets:
            matches.append((m.start(), m.end(), targets[m.start()]))
    if not matches:
        return 0, 0

    # 计算 run 偏移表
    offsets = []
    cur = 0
    for r in runs:
        offsets.append(cur)
        cur += len(r.text)

    # 检查是否有匹配跨越多个 run
    needs_merge = False
    for start, end, marker in matches:
        ri = max(i for i, off in enumerate(offsets) if off <= start)
        run_end = offsets[ri] + len(runs[ri].text)
        if end > run_end:
            needs_merge = True
            break

    if needs_merge:
        # 合并所有 run 到第一个 run，保留首个 run 的格式
        runs[0].text = full
        for r in runs[1:]:
            r.text = ""
        # 重置为单 run 模式
        offsets = [0]
        runs = [runs[0]]

    # 从右向左替换（保持左侧位置不变）
    ok = 0
    skip = 0
    for start, end, marker in sorted(matches, key=lambda x: -x[0]):
        ri = max(i for i, off in enumerate(offsets) if off <= start)
        run_end = offsets[ri] + len(runs[ri].text)
        if end > run_end:
            skip += 1
            continue
        rel = start - offsets[ri]
        runs[ri].text = runs[ri].text[:rel] + marker + runs[ri].text[rel + (end - start):]
        ok += 1
    return ok, skip


def _replace_image_paragraph(p, marker: str) -> None:
    """把段落里的图片元素移除，并将段落文本替换为图表占位符。"""
    from docx.oxml.ns import qn

    for tag in ("w:drawing", "w:pict", "w:object"):
        for node in p._p.findall(".//" + qn(tag)):
            node.getparent().remove(node)
    runs = p.runs
    for r in runs[1:]:
        r.text = ""
    if runs:
        runs[0].text = marker
    else:
        p.add_run(marker)
