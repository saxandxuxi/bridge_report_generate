# -*- coding: utf-8 -*-
"""报告生成：把统计值、日期、图表填充进 Word 模板。

流程：
  1. 展开表格中的可重复行（{{rows.<数据集>}} + {{col.<字段>}}）；
  2. 处理条件渲染块（{{?condition}}...{{?}}）；
  3. 替换所有正文/表格/页眉页脚里的 {{stats.*}}、{{date.*}} 占位符；
  4. 处理带默认值的占位符（{{key|default:值}}）和表达式（{{expr:...}}）；
  5. 把 {{chart.<ID>}} 独占段落替换为图表图片；
  6. 最后检查是否还有未替换的占位符。

支持的占位符语法：
  {{stats.<指标>.<统计>}}              基本统计值
  {{stats.<指标>.<统计>:0.1f}}         带格式说明符
  {{date.<字段>}}                       报告期日期
  {{chart.<ID>}}                         图表位置
  {{rows.<数据集>}} + {{col.<字段>}}    表格可重复行
  {{key|default:默认值}}                带默认值的占位符（找不到时用默认值）
  {{expr:a - b}}                         简单表达式计算（+ - * /）
  {{?condition}}...{{?}}                 条件渲染块
"""

import copy
import datetime as dt
import hashlib
import io
import json
import logging
import os
import re
import shutil
import tempfile
import zipfile
from typing import Callable, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

log = logging.getLogger("report-agent.report_builder")

# 占位符正则：key + 可选格式说明符 + 可选默认值
# key 用 [^:|}]+ 匹配除分隔符外的任意字符（含中文、#、（）等），避免特殊字符占位符"隐形"
MARKER_RE = re.compile(r"\{\{([^:|}]+)(?::([^}|]+))?(?:\|default:([^}]*))?\}\}")
# 条件块正则：{{?condition}}content{{?}}
COND_RE = re.compile(r"\{\{\?(.+?)\}\}(.*?)\{\{\?\}\}", re.DOTALL)
# 表达式前缀
EXPR_PREFIX = "expr:"


def _eval_simple_expr(expr: str, stats: Dict) -> Optional[float]:
    """安全地计算简单算术表达式（仅支持 + - * / 和数字、stats 值引用）。

    如 "stats.temperature.max - stats.temperature.min" → 9.5
    如果无法计算返回 None。
    """
    # 提取所有标识符引用
    refs = re.findall(r"stats\.[a-zA-Z0-9_.]+", expr)
    replacements = {}
    for ref in refs:
        parts = ref.split(".")
        node = stats
        try:
            for p in parts[1:]:
                node = node[p]
            if isinstance(node, (int, float)):
                replacements[ref] = str(node)
            else:
                return None
        except (KeyError, TypeError):
            return None
    # 替换引用为数值
    for ref, val in replacements.items():
        expr = expr.replace(ref, val)
    # 验证：只允许数字、运算符和空格
    if not re.fullmatch(r"[\d\s+\-*/.()]+", expr):
        return None
    try:
        result = eval(expr, {"__builtins__": {}}, {})
        return float(result)
    except Exception:  # noqa: BLE001
        return None


def build_value_resolver(stats: Dict, period: Dict,
                         data_registry=None, data_values: Dict = None,
                         bridge=None, missing_sink: Optional[list] = None,
                         lineage: Optional[list] = None,
                         data_meta: Optional[Dict] = None,
                         llm_cfg: Optional[Dict] = None,
                         missing_marker: str = "—") -> Callable[[str], str]:
    """根据统计结果和报告期构建占位符解析函数。

    支持以下占位符类型：
      stats.<指标>.<统计>     — 统计值
      date.<字段>             — 报告期日期
      expr:<表达式>           — 简单算术表达式
      cell.<metric>.<column>.<stat>  — 表格单元格（多数据源 + 测点级查询）
      data.N                 — 通用数据占位符（回填原始值）
      summary.<metric>       — 按季度/年度统计生成的结论性总结（LLM 或规则兜底）

    bridge: 可选的真实监测数据适配器（report_agent.bridge_source.BridgeData），
            优先于 data_registry / stats 解析 cell 与 stats 占位符。
    missing_sink: 可选的列表，解析不到的 cell/stats 键会追加进去（供 Web 待补清单）。
    lineage: 可选的列表，每次解析把 占位符/来源/计算链路/最终值 追加进去（供数据血缘日志）。
    data_meta: {{data.N}} 的原文上下文（来自 analysis numbers），供血缘日志引用。
    missing_marker: 解析不到时填入文档的标记（默认 “—”），并写入血缘日志说明未找到。
    """

    def _log(entry: Dict) -> None:
        if lineage is not None:
            lineage.append(entry)

    def resolve(key: str, table_title: str = "", row_index: int = 0) -> str:
        # 表达式计算
        if key.startswith(EXPR_PREFIX):
            expr = key[len(EXPR_PREFIX):]
            result = _eval_simple_expr(expr, stats)
            if result is not None:
                return f"{result:.1f}"
            raise KeyError(f"表达式无法计算: {key}")
        # 表格单元格：cell.metric.column.stat
        if key.startswith("cell."):
            from .data_loader import resolve_cell
            # 同一位置多测点行的索引后缀（如 cell.crack.5#塔底部.avg#2 -> row_index=1）
            row_override = None
            m_sfx = re.match(r"^(.*)#(\d+)$", key)
            if m_sfx:
                key = m_sfx.group(1)
                row_override = int(m_sfx.group(2)) - 1
            parts = key.split(".")
            if len(parts) != 4:
                raise KeyError(f"cell 占位符格式错误（期望 cell.metric.column.stat）: {key}")
            _, metric, column, stat = parts
            cell_row = row_override if row_override is not None else row_index
            value = None
            if bridge is not None:
                try:
                    value, detail = bridge.resolve_cell_detail(
                        metric, column, stat, period,
                        table_title=table_title or "", row_index=cell_row)
                    if value is not None:
                        _log({**detail, "类型": "cell重算",
                              "输出": _format_cell_value(stat, value)})
                    else:
                        _log({"占位符": key, "类型": "cell重算", "结果": "未找到", **detail})
                except Exception as exc:  # noqa: BLE001
                    log.warning("桥数据解析 cell 失败 %s: %s", key, exc)
            if value is None and bridge is None and data_registry is not None:
                value = resolve_cell(data_registry, metric, column, stat, period)
            if value is None:
                if missing_sink is not None:
                    missing_sink.append(key)
                # 桥模式下解析不到 -> 填占位符（真实数据源已尝试过）
                if bridge is not None:
                    return missing_marker
                # 指标无数据源（如 strain/displacement）——返回占位符而非崩溃
                if data_registry is not None and data_registry.get(metric) is None:
                    import logging
                    logging.getLogger("report-agent.builder").warning(
                        "指标 %s 无数据源，cell.%s 填入占位符", metric, key
                    )
                    return "—"
                # 未知列（如表格行超出已映射测点范围）——返回占位符而非崩溃
                if column.startswith("unknown_"):
                    import logging
                    logging.getLogger("report-agent.builder").warning(
                        "未知测点列，cell.%s 填入占位符", key
                    )
                    return "—"
                raise KeyError(f"无法计算单元格: {key}（数据源 {metric} 或列 {column} 不存在）")
            return _format_cell_value(stat, value)
        if key.startswith("stats."):
            parts = key.split(".")
            # {{stats.<metric>.<stat>.loc}}：最值对应的监测部位（如 A-MAX-LOC）
            if len(parts) == 4 and parts[3] == "loc" and bridge is not None:
                _metric, _stat = parts[1], parts[2]
                _value, _detail = bridge.resolve_metric_stat_detail(_metric, _stat, period)
                if _detail and _detail.get("位置"):
                    _log({"占位符": key, "类型": "stats最值位置", "值": _detail["位置"],
                          "输出": _detail["位置"],
                          "关联": f"stats.{_metric}.{_stat}",
                          "说明": "取达到最值/最大差值的传感器监测部位"})
                    return str(_detail["位置"])
                _log({"占位符": key, "类型": "stats最值位置", "结果": "未找到",
                      "原因": f"stats.{_metric}.{_stat} 无最值位置可推断"})
                if missing_sink is not None:
                    missing_sink.append(key)
                return missing_marker
            if len(parts) == 3 and bridge is not None:
                _, metric, stat = parts
                # 总结段落状态句（缺失数据位置自动生成）
                if stat == "data_status":
                    return bridge.resolve_data_status(metric, period)
                if stat == "abnormal_clause":
                    return bridge.resolve_abnormal_clause(metric, period)
                value = None
                try:
                    value, detail = bridge.resolve_metric_stat_detail(metric, stat, period)
                    if value is not None:
                        _log({**detail, "类型": "stats重算",
                              "输出": _format_stat_value(key, value)})
                    else:
                        _log({"占位符": key, "类型": "stats重算", "结果": "未找到", **detail})
                except Exception as exc:  # noqa: BLE001
                    log.warning("桥数据解析 stats 失败 %s: %s", key, exc)
                if value is not None:
                    return _format_stat_value(key, value)
                if missing_sink is not None:
                    missing_sink.append(key)
                # 只有“桥模式真实指标”解析不到才标缺失；
                # stats.days 等非桥指标继续走下方 computed 兜底
                if bridge is not None and metric in bridge.metrics:
                    return missing_marker
            node = stats
            found = True
            for p in parts[1:]:
                if not isinstance(node, dict) or p not in node:
                    found = False
                    break
                node = node[p]
            if found:
                return _format_stat_value(key, node)
            # 兼容多数据源：尝试从 data_registry 加载
            if data_registry is not None and len(parts) == 3:
                from .data_loader import resolve_metric_stat
                metric, col, stat = parts[1], None, parts[2]
                value = resolve_metric_stat(data_registry, metric, stat, period)
                if value is not None:
                    return _format_stat_value(key, value)
            # 指标无数据源——返回占位符而非崩溃
            _metric = parts[1] if len(parts) > 1 else "?"
            if data_registry is not None and _metric not in data_registry.all_metrics():
                import logging
                logging.getLogger("report-agent.builder").warning(
                    "指标 %s 无数据源，填入占位符: %s", _metric, key
                )
                return "—"
            raise KeyError(f"统计量不存在: {key}")
        if key.startswith("date."):
            field = key.split(".", 1)[1]
            # 落款日期：2026年8月6日
            if field == "signature":
                sig = _signature_date(period)
                return f"{sig.year}年{sig.month}月{sig.day}日"
            # 报告期起止月份（结论段“xx月-xx月”用）
            if field == "period_start_month":
                start = period.get("start")
                return f"{start.month}月" if isinstance(start, dt.date) else ""
            if field == "period_end_month":
                end = period.get("end")
                return f"{end.month}月" if isinstance(end, dt.date) else ""
            # 周期标签：period_label（2026.1~3）/ period_label_cn（2026年第一季度）
            if field in ("period_label", "label"):
                value = period.get("label") or period.get("label_cn")
                return str(value) if value else ""
            if field in ("period_label_cn", "label_cn"):
                value = period.get("label_cn") or period.get("label")
                return str(value) if value else ""
            alias = {"period_start": "start", "period_end": "end"}.get(field, field)
            value = period.get(alias)
            if value is None:
                raise KeyError(f"日期字段不存在: {key}")
            if isinstance(value, dt.date):
                return value.strftime("%Y-%m-%d")
            if isinstance(value, dt.datetime):
                return value.strftime("%Y-%m-%d %H:%M")
            return str(value)
        # 结论性总结：{{summary.<metric>}} —— 基于季度/年度统计生成
        if key.startswith("summary."):
            metric = key.split(".", 1)[1] if "." in key else ""
            if bridge is not None and metric and metric in bridge.metrics:
                text = bridge.build_feature_summary(metric, period,
                                                    llm_cfg=llm_cfg)
                if text:
                    _log({
                        "占位符": key,
                        "类型": "特征总结",
                        "指标": metric,
                        "输出": text,
                        "说明": "基于季度/年度统计的极值位置与缺失情况生成",
                    })
                    return text
                if missing_sink is not None:
                    missing_sink.append(key)
                return missing_marker
            raise KeyError(f"不支持的总结占位符: {key}")
        # 通用数据占位符：回填 annotate_docx 阶段保存的原始值
        if key.startswith("data."):
            if data_values and key in data_values:
                raw = str(data_values[key])
                meta = (data_meta or {}).get(key, {})
                _log({
                    "占位符": key,
                    "类型": "data原文回填",
                    "值": raw,
                    "原文上下文": meta.get("context", ""),
                    "原文段落": meta.get("paragraph"),
                    "说明": "回填分析阶段记录的原文数值，未重算（如需重算应使用 stats.* 占位符）",
                })
                return raw
            _log({"占位符": key, "类型": "data原文回填", "结果": "未找到",
                  "原因": "analysis data_values 中没有该键的原文值"})
            if missing_sink is not None:
                missing_sink.append(key)
            return missing_marker
        raise KeyError(f"不支持的占位符: {key}")

    resolve.lineage = lineage or []
    return resolve


def _format_cell_value(stat: str, value: float) -> str:
    """格式化表格单元格值。"""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        # 整数化的值（如 0.0）显示为整数
        if value == int(value) and abs(value) < 1e9:
            return str(int(value))
        # 只有真正极小（绝对值 < 1e-4，如 9.7e-05）才用科学计数法；
        # 0.0001~0.1 之间的小量（如 -2.810e-04、1.357e-03）用普通小数
        # 展示，避免表格里 0.000281 这类数据写成科学计数。
        if abs(value) < 1e-4 and value != 0.0:
            return f"{value:.3e}"
        if abs(value) < 0.1:
            return f"{value:.4g}"
        # 标准差用 2 位小数
        if stat in ("std", "均方根", "rms"):
            return f"{value:.2f}"
        return f"{value:.2f}"
    return str(value)


def _format_stat_value(key: str, value) -> str:
    if isinstance(value, dt.date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        if key.endswith(".std"):
            return f"{value:.2f}"
        if abs(value) < 1e-4 and value != 0.0:
            return f"{value:.3e}"
        if abs(value) < 0.1:
            return f"{value:.4g}"
        return f"{value:.1f}"
    return str(value)


def _format_col_value(field: str, value) -> str:
    if isinstance(value, dt.date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float):
        if field == "deviation":
            return f"{value:+.1f}"
        return f"{value:.1f}"
    if isinstance(value, int):
        return str(value)
    return str(value)


def _fill_paragraph(paragraph: Paragraph, resolver: Callable[[str], str],
                    table_title: str = "", row_index: int = 0) -> int:
    """替换段落内所有占位符，返回替换数量。合并多个 run，保留首个 run 的格式。

    支持语法：
      {{key}}               基本替换
      {{key:0.1f}}          带格式说明符
      {{key|default:值}}    带默认值（key 不存在时用默认值）
      {{key:0.1f|default:0}} 带格式和默认值
    """
    runs = paragraph.runs
    if not runs:
        return 0
    full = "".join(r.text for r in runs)
    matches = list(MARKER_RE.finditer(full))
    if not matches:
        return 0

    parts = []
    last = 0
    replaced = 0
    for m in matches:
        key, spec, default = m.group(1), m.group(2), m.group(3)
        parts.append(full[last:m.start()])
        try:
            value = resolver(key, table_title=table_title, row_index=row_index)
            if spec:
                try:
                    if isinstance(value, (int, float)):
                        value = format(value, spec.strip())
                    elif isinstance(value, str) and value.replace(".", "").replace("-", "").isdigit():
                        value = format(float(value), spec.strip())
                except (ValueError, TypeError):
                    pass  # 格式化失败则保持原值
            parts.append(str(value))
            replaced += 1
        except KeyError as exc:
            if default is not None:
                # 有默认值，使用默认值替代
                parts.append(default)
                replaced += 1
            else:
                raise KeyError(f"{exc}（位置：\"{full.strip()[:40]}\"）") from exc
        last = m.end()
    parts.append(full[last:])

    runs[0].text = "".join(parts)
    for r in runs[1:]:
        r.text = ""
    return replaced


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(r.text for r in paragraph.runs)


def _ensure_rgb(path: str) -> str:
    """把 RGBA/调色板 PNG 转成 RGB（WPS/Word 对 RGBA 渲染兼容性差，可能显示空白）。"""
    try:
        from PIL import Image
        img = Image.open(path)
        if img.mode not in ("RGBA", "LA", "P", "PA"):
            return path
        cache = os.path.join(tempfile.gettempdir(), "report_rgb")
        os.makedirs(cache, exist_ok=True)
        # 缓存键包含文件修改时间/大小，图库重新生成后不会返回旧缓存图
        mtime = os.path.getmtime(path) if os.path.isfile(path) else 0
        fsize = os.path.getsize(path) if os.path.isfile(path) else 0
        key = hashlib.sha1(
            f"{os.path.abspath(path)}|{mtime}|{fsize}".encode("utf-8")
        ).hexdigest()[:16]
        out = os.path.join(cache, key + ".png")
        if os.path.isfile(out):
            return out
        bg = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode in ("RGBA", "LA", "PA"):
            rgba = img.convert("RGBA")
            bg.paste(rgba, mask=rgba.split()[-1])
        else:
            bg.paste(img.convert("RGB"))
        bg.save(out, "PNG")
        return out
    except Exception as exc:  # noqa: BLE001
        log.warning("PNG 转 RGB 失败 %s: %s", path, exc)
        return path


def _set_auto_line_spacing(paragraph: Paragraph) -> None:
    """覆盖默认样式的固定行距(lineRule=exact)。

    模板的 Normal 样式常带 w:line=500 lineRule=exact（固定行高约 25pt），
    inline 大图会溢出固定行高、压住上下内容造成重叠。
    这里显式改为单倍自动行距，让行随图片高度自动扩展。
    """
    try:
        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
    except Exception as exc:  # noqa: BLE001
        log.warning("设置图片段落自动行距失败: %s", exc)


def _insert_chart(paragraph: Paragraph, png_path: str, width_inches: float) -> None:
    """把段落清空并插入居中图片。"""
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_auto_line_spacing(paragraph)
    run = paragraph.add_run()
    w, h = _chart_insert_size(png_path, width_inches)
    if w:
        run.add_picture(_ensure_rgb(png_path), width=w)
    else:
        run.add_picture(_ensure_rgb(png_path), height=h)


def _chart_insert_size(png_path: str, width_inches: float,
                       max_height_inches: float = 6.5):
    """计算插入尺寸：优先按宽度；图过高时按最大高度限制（保持比例）。"""
    try:
        from PIL import Image
        img = Image.open(png_path)
        w, h = img.size
        ratio = h / w if w else 1.0
        if ratio * width_inches > max_height_inches:
            return None, Inches(max_height_inches)
    except Exception:  # noqa: BLE001
        pass
    return Inches(width_inches), None


def iter_block_items(parent):
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        body = parent.element.body
    else:
        body = parent._tc
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _expand_row_tables(
    doc: Document,
    row_datasets: Dict[str, List[Dict]],
    base_resolver: Callable[[str], str],
) -> int:
    """展开 {{rows.<数据集>}} 模板行，返回展开的行数。"""
    expanded = 0
    for table in doc.tables:
        for row in list(table.rows):
            dataset = None
            for cell in row.cells:
                m = re.search(r"\{\{rows\.([a-zA-Z0-9_.]+)\}\}", cell.text)
                if m:
                    dataset = m.group(1)
                    break
            if not dataset:
                continue

            records = row_datasets.get(dataset)
            if not records:
                raise ValueError(f"行数据集 '{dataset}' 没有可用数据")

            tr = row._tr
            anchor = tr
            for rec in records:
                new_tr = copy.deepcopy(tr)
                for tc in new_tr.findall(qn("w:tc")):
                    for p_elm in tc.findall(qn("w:p")):
                        para = Paragraph(p_elm, None)

                        def col_resolver(key: str, _rec=rec) -> str:
                            if key.startswith("rows."):
                                return ""  # 行模板指令，填充时移除
                            if key.startswith("col."):
                                field = key.split(".", 1)[1]
                                if field not in _rec:
                                    raise KeyError(f"记录缺少字段: {field}")
                                return _format_col_value(field, _rec[field])
                            return base_resolver(key)

                        _fill_paragraph(para, col_resolver)
                anchor.addnext(new_tr)
                anchor = new_tr
                expanded += 1
            tr.getparent().remove(tr)
    return expanded


def _walk_paragraphs(doc: Document):
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            yield item
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def _eval_condition(condition: str, resolver: Callable[[str], str]) -> bool:
    """评估条件表达式，返回 True/False。

    支持的格式：
      stats.temperature.days_above_30>0     — 比较
      stats.temperature.max>=35            — 大于等于
      stats.temperature.min<10             — 小于
      stats.humidity.avg                   — 仅检查存在性和真值
      not stats.temperature.days_above_30   — 取反
    """
    condition = condition.strip()

    # 取反
    negate = False
    if condition.startswith("not "):
        negate = True
        condition = condition[4:].strip()

    # 尝试匹配比较运算符
    for op in (">=", "<=", "!=", "==", ">", "<"):
        if op in condition:
            left, right = condition.split(op, 1)
            left = left.strip()
            right = right.strip()
            try:
                left_val = float(resolver(left)) if left.startswith(("stats.", "date.", "expr:")) else float(left)
                right_val = float(right) if not right.startswith(("stats.", "date.")) else float(resolver(right))
            except (KeyError, ValueError):
                return negate  # 无法解析的值默认 False（取反后 True）

            if op == ">":
                result = left_val > right_val
            elif op == "<":
                result = left_val < right_val
            elif op == ">=":
                result = left_val >= right_val
            elif op == "<=":
                result = left_val <= right_val
            elif op == "!=":
                result = left_val != right_val
            elif op == "==":
                result = left_val == right_val
            return (not result) if negate else result

    # 无运算符：检查存在性和真值
    try:
        val = resolver(condition)
        result = bool(val) and val != "0" and val != "0.0"
        return (not result) if negate else result
    except KeyError:
        return negate  # 不存在 → False（取反后 True）


def _process_conditional_blocks(paragraph: Paragraph, resolver: Callable[[str], str]) -> None:
    """处理段落中的条件渲染块 {{?condition}}content{{?}}。

    条件为真时保留 content（移除条件标记），条件为假时移除整个块。
    """
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text for r in runs)
    if "{{?" not in full:
        return

    def replace_block(m: re.Match) -> str:
        condition = m.group(1).strip()
        content = m.group(2)
        if _eval_condition(condition, resolver):
            return content
        return ""

    new_text = COND_RE.sub(replace_block, full)
    if new_text != full:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""


# ---------------------------------------------------------------------------
# 报告期文字修正：页眉年份/季度、落款日期、目录页码刷新
# ---------------------------------------------------------------------------

def _signature_date(period: Dict) -> dt.datetime:
    """落款日期：优先用报告期生成时间，其次报告期结束日。"""
    val = period.get("generated_at")
    if isinstance(val, dt.datetime):
        return val
    if isinstance(val, dt.date):
        return dt.datetime(val.year, val.month, val.day)
    end = period.get("end")
    if isinstance(end, dt.date):
        return dt.datetime(end.year, end.month, end.day)
    return dt.datetime.now()


def _quarter_no(period: Dict) -> int:
    end = period.get("end") or period.get("generated_at")
    if isinstance(end, dt.datetime):
        end = end.date()
    if isinstance(end, dt.date):
        return (end.month - 1) // 3 + 1
    return 0


def _period_text_replacements(period: Dict):
    """返回 (正则, 替换串) 列表，按顺序应用。

    覆盖模板中残留的字面量日期/季度写法：
      xxxx年xx月xx日 / XXXX年XX月XX日 / ××××年××月××日
      xxxx年 / XXXX年
      第X季度 / 第x季度
      xx月xx日
    """
    sig = _signature_date(period)
    y, m, d = sig.year, sig.month, sig.day
    q = _quarter_no(period)
    x4 = r"[xX×]{4}"
    x2 = r"[xX×]{2}"
    start = period.get("start")
    end = period.get("end")
    sm = getattr(start, "month", m)
    em = getattr(end, "month", m)
    def _prev_year_quarter(m):
        """把“上一年度的第X季度”改成当前报告期（如 2025年第一季度 -> 2026年第一季度）。"""
        y = int(m.group(1))
        if y == sig.year - 1:
            return f"{sig.year}年{m.group(2)}"
        return m.group(0)

    reps = [
        (re.compile(rf"{x4}年{x2}月{x2}日"), f"{y}年{m}月{d}日"),
        # 报告期范围：xx月-xx月（如结论段“2026年xx月-xx月”）
        (re.compile(rf"{x2}月[-—–~～]{x2}月"), f"{sm}月-{em}月"),
        (re.compile(rf"{x4}年"), f"{y}年"),
        (re.compile(rf"第[xX×]季度"), f"第{q}季度" if q else "第X季度"),
        (re.compile(rf"{x2}月{x2}日"), f"{m}月{d}日"),
        # 签字表等处的裸月/日占位（顺序在范围之后，避免误伤）
        (re.compile(rf"{x2}月"), f"{m}月"),
        (re.compile(rf"{x2}日"), f"{d}日"),
        # 上一年度的“第X季度” -> 当前报告期（正文字面量，如“2025年第一季度”）
        (re.compile(r"(20\d{2})年(第[一二三四1-4]季度)"), _prev_year_quarter),
        # 兜底：模板残留 “2025年{{date.period_label_cn}}” 解析后变成
        # “2025年2026年第1季度” → 去掉旧年份
        (re.compile(r"20\d{2}年(\d{4}年)"), r"\1"),
    ]
    # 年度报告：季度字眼全部改为年度
    if period.get("mode") == "yearly":
        reps += [
            (re.compile(r"本季度"), "本年度"),
            # “2025年第一季度” -> “2026年度”（顺序在 _prev_year_quarter 之后）
            (re.compile(r"(20\d{2})年(第[一二三四1-4]季度)"),
             lambda m: f"{y}年度"),
            # 页眉等 “xxxx年第X季度” -> “2026年”
            (re.compile(rf"第[xX×]季度"), ""),
            # 封面 “2026年度（2026年）” 冗余括注 -> “2026年度”
            (re.compile(rf"（{y}年）"), ""),
        ]
    return reps


def _apply_period_text_fixes(doc: Document, period: Dict) -> int:
    """把页眉/页脚/正文里的字面量日期与季度替换为实际报告期。"""
    replacements = _period_text_replacements(period)
    changed = 0
    for paragraph in _walk_paragraphs(doc):
        text = _paragraph_text(paragraph)
        new_text = text
        for pat, rep in replacements:
            new_text = pat.sub(rep, new_text)
        if new_text != text:
            runs = paragraph.runs
            if not runs:
                continue
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""
            changed += 1
    return changed


def _enable_update_fields(doc: Document) -> None:
    """在 settings.xml 中写入 <w:updateFields/>，让 Word 打开时自动刷新目录页码。"""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    # 按 CT_Settings 顺序插入：characterSpacingControl 之后、hdrShapeDefaults 之前
    anchor = settings.find(qn("w:hdrShapeDefaults"))
    if anchor is None:
        anchor = settings.find(qn("w:footnotePr"))
    if anchor is None:
        anchor = settings.find(qn("w:compat"))
    if anchor is not None:
        anchor.addprevious(el)
    else:
        settings.append(el)


def _flatten_rgba_in_docx(path: str) -> int:
    """把 docx 包内所有 RGBA/调色板 PNG 原位转成 RGB（WPS/Word 兼容性兜底）。

    返回转换的图片数量。转换失败不影响文档本身。
    """
    tmp = path + ".rgb.tmp"
    converted = 0
    try:
        with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/media/") and data.startswith(b"\x89PNG"):
                    try:
                        from PIL import Image
                        img = Image.open(io.BytesIO(data))
                        if img.mode in ("RGBA", "LA", "PA", "P"):
                            bg = Image.new("RGB", img.size, (255, 255, 255))
                            rgba = img.convert("RGBA")
                            bg.paste(rgba, mask=rgba.split()[-1])
                            buf = io.BytesIO()
                            bg.save(buf, "PNG")
                            data = buf.getvalue()
                            converted += 1
                    except Exception:  # noqa: BLE001
                        pass
                zout.writestr(item, data)
        shutil.move(tmp, path)
    except Exception as exc:  # noqa: BLE001
        log.warning("输出图 RGB 平铺失败: %s", exc)
        if os.path.isfile(tmp):
            try:
                os.remove(tmp)
            except OSError:
                pass
    return converted


def _copy_ppr(target_el, source_el) -> None:
    """把源段落的段落属性（pPr，含样式/间距/行高）复制到目标段落。"""
    try:
        src = source_el.find(qn("w:pPr"))
        if src is None:
            return
        import copy as _copy
        tgt = target_el.find(qn("w:pPr"))
        if tgt is not None:
            target_el.remove(tgt)
        target_el.insert(0, _copy.deepcopy(src))
    except Exception:  # noqa: BLE001
        pass


def _p_has_blip(p_el) -> bool:
    return len(p_el.findall(".//" + qn("a:blip"))) > 0


def _p_is_caption(p_el) -> bool:
    text = "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()
    return bool(re.match(r"^图\d+(\.\d+)*-\d+\s", text))


def _p_has_content(p_el) -> bool:
    if _p_has_blip(p_el):
        return True
    text = "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()
    return bool(text)


def _reorder_chart_blocks(doc: Document) -> int:
    """同一节内：时程曲线图放在前面，频率分布直方图放在后面（保持同类内原顺序）。

    图表块 = 图片段 + 紧随其后的图注段。块之间的空段忽略；
    两个块之间出现 非空正文段/表格 视为节边界。
    返回重排的块数。
    """
    body = doc.element.body
    children = list(body.iterchildren())

    blocks = []
    i = 0
    while i < len(children):
        el = children[i]
        if el.tag == qn("w:p") and _p_has_blip(el):
            els = [el]
            kind = "trend"
            j = i + 1
            if j < len(children) and children[j].tag == qn("w:p") and _p_is_caption(children[j]):
                els.append(children[j])
                cap_text = "".join(t.text or "" for t in children[j].iter(qn("w:t")))
                kind = "hist" if "直方图" in cap_text else "trend"
                j += 1
            blocks.append((i, j - 1, kind, els))
            i = j
        else:
            i += 1
    if len(blocks) < 2:
        return 0

    groups = []
    cur = []
    for k, (start, end, kind, els) in enumerate(blocks):
        if cur:
            prev_end = blocks[k - 1][1]
            between = children[prev_end + 1:start]
            boundary = any(
                c.tag == qn("w:tbl")
                or (c.tag == qn("w:p") and _p_has_content(c) and not _p_is_caption(c))
                for c in between
            )
            if boundary:
                groups.append(cur)
                cur = []
        cur.append((kind, els))
    if cur:
        groups.append(cur)

    moved = 0
    for group in groups:
        kinds = [b[0] for b in group]
        if kinds == sorted(kinds, key=lambda x: 0 if x == "trend" else 1):
            continue
        ordered = sorted(group, key=lambda b: 0 if b[0] == "trend" else 1)
        flat = [el for _, els in ordered for el in els]
        prev = flat[0].getprevious()
        for el in flat:
            if prev is None:
                body.insert(0, el)
            else:
                prev.addnext(el)
            prev = el
        moved += len(group)
    if moved:
        log.info("图表块按图型重排：%d 块（时程曲线图在前，直方图在后）", moved)
    return moved


def _renumber_and_unify_captions(doc: Document) -> int:
    """重排后按新顺序重编图注号（图3.1.1-N），并统一图注字号为 10.5pt。"""
    counters = {}
    changed = 0
    for p in _walk_paragraphs(doc):
        text = _paragraph_text(p)
        m = re.match(r"^(图\d+(?:\.\d+){0,2})-(\d+)(\s.*)$", text)
        if not m:
            continue
        prefix = m.group(1)[1:]
        counters[prefix] = counters.get(prefix, 0) + 1
        new_text = f"图{prefix}-{counters[prefix]}{m.group(3)}"
        runs = p.runs
        if not runs:
            continue
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        runs[0].font.size = Pt(10.5)
        try:
            rPr = runs[0]._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:  # noqa: BLE001
            pass
        changed += 1
    return changed


def _add_caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    """在图表段落之后插入一行居中图注（图X.X.X 图名）。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    _copy_ppr(new_p, paragraph._p)  # 继承图表段落的样式/间距，避免行高塌陷
    p = Paragraph(new_p, paragraph._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_auto_line_spacing(p)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def _new_paragraph_after(anchor_el, parent, source_el=None) -> Paragraph:
    """在指定元素后新建一个空段落，返回 Paragraph 包装。"""
    new_p = OxmlElement("w:p")
    anchor_el.addnext(new_p)
    if source_el is not None:
        _copy_ppr(new_p, source_el)
    return Paragraph(new_p, parent)


def build_report(
    template_path: str,
    output_path: str,
    resolver: Callable[[str], str],
    chart_images: Dict[str, str],
    row_datasets: Optional[Dict[str, List[Dict]]] = None,
    chart_width_inches: float = 5.8,
    strict: bool = True,
    period: Optional[Dict] = None,
    chart_captions: Optional[Dict[str, str]] = None,
    extra_charts: Optional[Dict[str, List[Dict]]] = None,
    text_replace: Optional[Dict[str, str]] = None,
) -> List[str]:
    """基于模板生成报告，返回未替换的占位符列表（strict=True 时抛出异常）。"""
    doc = Document(template_path)
    row_datasets = row_datasets or {}
    chart_captions = chart_captions or {}
    extra_charts = extra_charts or {}

    if row_datasets:
        _expand_row_tables(doc, row_datasets, resolver)

    # 正文：按块顺序处理；记录最近的短文本作为表格标题上下文，
    # 供 cell 占位符按“表格标题 -> 断面位置 -> 测点N”解析。
    current_title = ""
    current_prefix = ""          # 章节号（如 3.1.1）
    fig_counters: Dict[str, int] = {}
    global_fig_no = 0
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            _process_conditional_blocks(item, resolver)
            text = _paragraph_text(item).strip()
            # 检测章节标题（如 “3.1.1.1  第6、7跨跨中断面环境温度”），用于图号前缀
            hm = re.match(r"^(\d+(?:\.\d+){1,3})(?![\d.])", text)
            if hm and len(text) <= 60:
                parts = hm.group(1).split(".")
                current_prefix = ".".join(parts[:3])
            chart_match = MARKER_RE.fullmatch(text)
            if chart_match and chart_match.group(1).startswith("chart."):
                chart_id = chart_match.group(1).split(".", 1)[1]
                if chart_id not in chart_images:
                    raise KeyError(f"缺少图表 {chart_id} 的图片文件")
                _insert_chart(item, chart_images[chart_id], chart_width_inches)
                # 图号 + 图名（图3.1.1-1 第6跨跨中断面主梁箱内环境温度时程曲线图）
                caption = chart_captions.get(chart_id, "")
                if caption:
                    if current_prefix:
                        fig_counters[current_prefix] = fig_counters.get(current_prefix, 0) + 1
                        fig_no = f"图{current_prefix}-{fig_counters[current_prefix]}"
                    else:
                        global_fig_no += 1
                        fig_no = f"图{global_fig_no}"
                    anchor_cap_p = _add_caption_after(item, f"{fig_no} {caption}")
                # 自动补齐的缺图（同一节监测部位多、模板占位符少）：
                # 必须插在锚点图注之后，避免图注顺序错乱
                anchor_el = anchor_cap_p._p if caption else item._p
                for ex in extra_charts.get(chart_id, []):
                    ex_p = _new_paragraph_after(anchor_el, item._parent, source_el=item._p)
                    ex_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_auto_line_spacing(ex_p)
                    w, h = _chart_insert_size(ex["path"], chart_width_inches)
                    if w:
                        ex_p.add_run().add_picture(_ensure_rgb(ex["path"]), width=w)
                    else:
                        ex_p.add_run().add_picture(_ensure_rgb(ex["path"]), height=h)
                    ex_anchor = ex_p._p
                    ex_caption = ex.get("caption", "")
                    if ex_caption:
                        if current_prefix:
                            fig_counters[current_prefix] = fig_counters.get(current_prefix, 0) + 1
                            ex_fig = f"图{current_prefix}-{fig_counters[current_prefix]}"
                        else:
                            global_fig_no += 1
                            ex_fig = f"图{global_fig_no}"
                        ex_cap_p = _add_caption_after(ex_p, f"{ex_fig} {ex_caption}")
                        ex_anchor = ex_cap_p._p
                    anchor_el = ex_anchor
            else:
                _fill_paragraph(item, resolver)
            new_text = _paragraph_text(item).strip()
            if new_text and len(new_text) <= 80:
                current_title = new_text
        elif isinstance(item, Table):
            for ri, row in enumerate(item.rows):
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _process_conditional_blocks(p, resolver)
                        _fill_paragraph(p, resolver, table_title=current_title, row_index=ri)

    # 页眉页脚
    for section in doc.sections:
        for p in section.header.paragraphs:
            _process_conditional_blocks(p, resolver)
            _fill_paragraph(p, resolver)
        for p in section.footer.paragraphs:
            _process_conditional_blocks(p, resolver)
            _fill_paragraph(p, resolver)

    # 页眉/落款等字面量日期季度修正 + 目录页码刷新标记
    fixed = 0
    if period:
        fixed = _apply_period_text_fixes(doc, period)
    _enable_update_fields(doc)
    if fixed:
        log.info("报告期文字修正 %d 处（页眉/落款日期/季度）", fixed)

    unfilled = _remaining_markers(doc)
    if unfilled and strict:
        raise RuntimeError(
            "模板中存在未替换的占位符:\n" + "\n".join(f"  {u}" for u in unfilled)
        )

    import os

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    # 同节内图表按图型重排（曲线图在前，直方图在后）+ 图注重编号 + 统一字号
    _reorder_chart_blocks(doc)
    _renumber_and_unify_captions(doc)
    # 兜底：所有含图片的段落统一自动行距，防止模板 Normal 固定行距(lineRule=exact)
    # 导致 inline 大图溢出行高、与上下内容重叠
    for pa in _walk_paragraphs(doc):
        if pa._p.findall('.//' + qn("w:drawing")):
            _set_auto_line_spacing(pa)
    if text_replace:
        _apply_text_replacements(doc, text_replace)
    doc.save(output_path)
    rgb_n = _flatten_rgba_in_docx(output_path)
    if rgb_n:
        log.info("输出文档 RGBA 图片转 RGB：%d 张", rgb_n)
    return unfilled


def _write_data_lineage(lineage: List[Dict], logs_dir: str, period: Dict) -> str:
    """把本次生成所有占位符的数据链路写入 logs/data_lineage_<label>.json 与 .log。

    返回写出的 .log 路径。找不到的项会以“结果=未找到”明确记录，不会静默填值。
    """
    os.makedirs(logs_dir, exist_ok=True)
    label = (period.get("label") or "report").replace(" ", "_")
    json_path = os.path.join(logs_dir, f"data_lineage_{label}.json")
    log_path = os.path.join(logs_dir, f"data_lineage_{label}.log")

    payload = {
        "报告期": {
            "mode": period.get("mode"),
            "start": str(period.get("start")),
            "end": str(period.get("end")),
            "label": label,
        },
        "汇总": {
            "总条目": len(lineage),
            "stats重算": sum(1 for e in lineage if e.get("类型") == "stats重算"),
            "cell重算": sum(1 for e in lineage if e.get("类型") == "cell重算"),
            "data原文回填": sum(1 for e in lineage if e.get("类型") == "data原文回填"),
            "未找到": sum(1 for e in lineage if e.get("结果") == "未找到"),
        },
        "条目": lineage,
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2, default=str)

    lines = [
        f"数据链路日志 报告期={payload['报告期']}",
        f"汇总: {payload['汇总']}",
        "=" * 70,
    ]
    for e in lineage:
        ph = e.get("占位符") or "?"
        typ = e.get("类型") or "?"
        if e.get("结果") == "未找到":
            lines.append(f"[未找到] {ph} ({typ}) 原因: {e.get('原因', '')}")
            continue
        out = e.get("输出")
        if typ == "data原文回填":
            lines.append(f"[{ph}] {typ} 值={e.get('值')}")
        else:
            lines.append(f"[{ph}] {typ} 输出={out}")
        if typ == "stats重算":
            lines.append(f"    指标={e.get('指标')} 统计量={e.get('统计量')} "
                         f"报告期={e.get('报告期')} 聚合规则={e.get('聚合规则')}")
            for d in e.get("逐传感器", []):
                lines.append(f"    传感器{d.get('传感器编号')} "
                             f"{d.get('监测部位')} {d.get('特征')} "
                             f"天数={d.get('天数')} 值={d.get('值')} "
                             f"文件={d.get('统计文件')}")
        elif typ == "cell重算":
            lines.append(f"    分支={e.get('分支')} 监测部位={e.get('监测部位')} "
                         f"表格={e.get('表格标题')} 行号={e.get('表格行号')}")
            d = e.get("传感器")
            if isinstance(d, dict):
                lines.append(f"    传感器{d.get('传感器编号')} {d.get('监测部位')} "
                             f"天数={d.get('天数')} 值={d.get('值')} "
                             f"文件={d.get('统计文件')}")
            elif e.get("聚合明细"):
                lines.append(f"    回退聚合: 规则={e['聚合明细'].get('聚合规则')} "
                             f"传感器数={e['聚合明细'].get('传感器数')}")
        elif typ == "data原文回填":
            lines.append(f"    值={e.get('值')} 原文上下文={e.get('原文上下文', '')[:60]} "
                         f"原文段落={e.get('原文段落')}")
        lines.append("-" * 70)
    with open(log_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return log_path


def verify_table_columns(output_path: str, lineage: Optional[List[Dict]] = None,
                         logs_dir: str = "", label: str = "report") -> List[Dict]:
    """填表后校验：找出“整列未解析”或“整列同值”的表格列。

    - 整列都是 “—”：该列所有单元格都没解析到传感器（应检查占位符/统计值）。
    - 整列数值完全相同：结合血缘里这些测点对应的传感器判断——
      若不同测点解析到了不同传感器却填成同值，说明是填充错误；
      若传感器本就相同，则提示“数据本身如此”。
    """
    doc = Document(output_path)
    warnings = []

    def _num(v: str):
        t = v.replace("−", "-").replace("—", "").strip()
        if not t:
            return None
        try:
            return float(t)
        except ValueError:
            return None

    # 血缘里 指标/测点 -> 解析到的传感器集合
    sensor_by_cell = {}
    for e in (lineage or []):
        ph = str(e.get("占位符") or "")
        if not ph.startswith("cell."):
            continue
        sid = None
        s = e.get("传感器")
        if isinstance(s, dict):
            sid = s.get("传感器编号")
        if sid is None:
            sid = e.get("sensor_id")
        sensor_by_cell.setdefault(ph, set())
        if sid:
            sensor_by_cell[ph].add(str(sid))

    for t_idx, table in enumerate(doc.tables):
        if not table.rows or len(table.rows) < 2:
            continue
        n_cols = max(len(r.cells) for r in table.rows)
        for c_idx in range(n_cols):
            vals = []
            markers = []
            for row in table.rows[1:]:
                if c_idx >= len(row.cells):
                    continue
                txt = row.cells[c_idx].text.strip()
                vals.append(txt)
                markers.append(txt)
            nums = [_num(v) for v in vals]
            filled = [n for n in nums if n is not None]
            missing = sum(1 for i, n in enumerate(nums) if n is None and vals[i] in ("—", ""))
            if not filled and missing >= 2:
                warnings.append({
                    "表序号": t_idx, "列序号": c_idx, "问题": "整列未解析",
                    "说明": f"{missing} 个数据行全部为 “—”（未找到传感器/统计值）",
                    "表头": str(table.rows[0].cells[c_idx].text.strip())[:40],
                })
                continue
            # 整列同值：用绝对容差判断（避免 1e-4 量级的小数被
            # round(n, 4) 误判为相同，如 0.00016 vs 0.000151）。
            if len(filled) >= 2 and max(filled) - min(filled) < 1e-9:
                same = filled[0]
                # 用血缘判断：这些行解析到的传感器是否不同
                sensors = set()
                for e in (lineage or []):
                    ph = str(e.get("占位符") or "")
                    if not ph.startswith("cell."):
                        continue
                    if abs(float(e.get("输出") or 0)) == abs(same) and "值" in e:
                        pass
                # 简化：统计该列占位符涉及的传感器数（通过血缘中 cell 条目）
                col_sensors = set()
                for e in (lineage or []):
                    ph = str(e.get("占位符") or "")
                    s = e.get("传感器")
                    sid = s.get("传感器编号") if isinstance(s, dict) else None
                    if sid and str(e.get("输出") or "") == f"{same:.1f}":
                        col_sensors.add(str(sid))
                warnings.append({
                    "表序号": t_idx, "列序号": c_idx,
                    "问题": "整列同值",
                    "值": same,
                    "行数": len(filled),
                    "血缘命中传感器数": len(col_sensors) if col_sensors else "未知",
                    "说明": (f"整列数值相同（{same:.2f}）。若不同测点对应不同传感器则为填充错误；"
                             "若传感器本就相同则属数据本身。")
                })
    if warnings and logs_dir:
        os.makedirs(logs_dir, exist_ok=True)
        with open(os.path.join(logs_dir, f"verify_tables_{label}.log"),
                  "w", encoding="utf-8") as f:
            f.write(f"填表校验 {label}：发现问题 {len(warnings)} 处\n")
            for w in warnings:
                f.write(json.dumps(w, ensure_ascii=False) + "\n")
    return warnings


def _apply_text_replacements(doc: Document, replacements: Dict[str, str]) -> int:
    """对全文（正文/表格/页眉页脚）做字典级文字修正，如 堡->墩（源文档字库替换错误）。"""
    changed = 0
    for paragraph in _walk_paragraphs(doc):
        full = "".join(r.text for r in paragraph.runs)
        new = full
        for a, b in replacements.items():
            new = new.replace(a, b)
        if new != full:
            runs = paragraph.runs
            if not runs:
                continue
            runs[0].text = new
            for r in runs[1:]:
                r.text = ""
            changed += 1
    if changed:
        log.info("文字修正 %d 处：%s", changed, replacements)
    return changed


def _remaining_markers(doc: Document) -> List[str]:
    leftovers = []
    for paragraph in _walk_paragraphs(doc):
        text = _paragraph_text(paragraph)
        for m in MARKER_RE.finditer(text):
            leftovers.append(f"{m.group(0)}  <- \"{text.strip()[:50]}\"")
        # 检测未处理的条件块
        if "{{?" in text:
            leftovers.append(f"未处理的条件块  <- \"{text.strip()[:50]}\"")
    return leftovers
