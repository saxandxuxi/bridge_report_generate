# -*- coding: utf-8 -*-
from docx import Document

src = r"D:\Code\data_analysis\templates\矮寨大桥_template_v6.docx"
dst = r"D:\Code\data_analysis\templates\矮寨大桥_template_v7.docx"
doc = Document(src)
fixed = 0
for tbl in doc.tables:
    rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
    if any("车道1" in str(c) for r in rows for c in r):
        for ri, row in enumerate(tbl.rows):
            label = row.cells[0].text.strip().replace("％", "%")
            if "比例" in label:
                for ci in range(1, len(row.cells)):
                    lane = tbl.rows[0].cells[ci].text.strip()
                    marker = f"{{{{cell.vehicle_count.{lane}.ratio#2}}}}"
                    cell = row.cells[ci]
                    # 清空并写入占位符
                    for p in cell.paragraphs:
                        for r_ in p.runs:
                            r_.text = ""
                        if p.runs:
                            p.runs[0].text = marker
                        else:
                            p.add_run(marker)
                        fixed += 1
                        break
        break
doc.save(dst)
print("fixed cells:", fixed)
doc2 = Document(dst)
for tbl in doc2.tables:
    rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
    if any("车道1" in str(c) for r in rows for c in r):
        for r in rows:
            print(r)
        break
