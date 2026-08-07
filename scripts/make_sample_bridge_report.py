#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成“某跨河桥梁月度监测报告”样例（DOCX），用于验证识别器：
包含固定参数（桥长123米等）、动态统计值、趋势图、CAD 示意图等混合内容。
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from PIL import Image, ImageDraw

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE_DIR, "outputs", "sample")
OUT_DOCX = os.path.join(OUT_DIR, "桥梁监测报告样例.docx")


def make_line_chart(path: str) -> None:
    img = Image.new("RGB", (800, 420), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([40, 40, 760, 380], outline="#888888")
    pts = [(60 + i * 70, 340 - int(150 + 60 * ((i % 5) - 2))) for i in range(10)]
    d.line(pts, fill="#2E74B5", width=3)
    for x, y in pts:
        d.ellipse([x - 4, y - 4, x + 4, y + 4], fill="#2E74B5")
    img.save(path)


def make_histogram(path: str) -> None:
    img = Image.new("RGB", (800, 420), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([40, 40, 760, 380], outline="#888888")
    heights = [70, 120, 200, 280, 230, 150, 90]
    bw = 80
    for i, h in enumerate(heights):
        x0 = 90 + i * (bw + 12)
        d.rectangle([x0, 380 - h, x0 + bw, 380], fill="#2E74B5", outline="white")
    img.save(path)


def make_schematic(path: str) -> None:
    """画一张简单的桥梁 CAD 示意图风格图片。"""
    img = Image.new("RGB", (900, 420), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([40, 40, 860, 380], outline="#888888")
    # 桥面
    d.line([(80, 200), (820, 200)], fill="#333333", width=6)
    # 桥墩
    for x in (220, 450, 680):
        d.rectangle([x - 18, 200, x + 18, 330], outline="#333333", width=3)
    # 主梁
    d.line([(80, 180), (820, 180)], fill="#555555", width=4)
    # 标注线
    d.line([(80, 350), (820, 350)], fill="#999999", width=1)
    for x in (80, 820):
        d.line([(x, 345), (x, 355)], fill="#999999", width=1)
    img.save(path)


def build() -> None:
    os.makedirs(OUT_DIR, exist_ok=True)
    make_line_chart(os.path.join(OUT_DIR, "chart_trend.png"))
    make_histogram(os.path.join(OUT_DIR, "chart_histogram.png"))
    make_schematic(os.path.join(OUT_DIR, "schematic.png"))

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def add(text, size=11, bold=False, align=None, after=6, color=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        return p

    def heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        return p

    def caption(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        p.paragraph_format.space_after = Pt(8)
        return p

    add("某跨河桥梁月度监测报告", size=18, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add("报告日期：2026-08-03　　合同编号：HT-2026-003",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    heading("一、工程概况")
    add("本桥为预应力混凝土连续梁桥，桥长123米，桥宽24.5米，"
        "跨径组合（30+60+30）米，设计荷载公路-I级，设计时速80km/h，"
        "抗震设防烈度7度，桥梁桩号K12+345，主梁采用C50混凝土。")

    heading("二、监测指标统计")
    add("本周期共31天监测数据：最高温度26.0℃，最低温度22.5℃，"
        "平均温度24.3℃，温度标准差1.18℃；桥梁挠度最大2.3mm，平均1.1mm；"
        "温度超过30℃的天数为0天。")

    add("表1 关键监测指标", size=9, after=4)
    table = doc.add_table(rows=7, cols=2)
    table.style = doc.styles["Table Grid"]
    rows = [
        ("指标", "数值"),
        ("最高温度", "26.0℃"),
        ("最低温度", "22.5℃"),
        ("平均温度", "24.3℃"),
        ("最大挠度", "2.3mm"),
        ("平均挠度", "1.1mm"),
        ("数据天数", "31天"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    heading("三、图表分析")
    doc.add_picture(os.path.join(OUT_DIR, "chart_trend.png"), width=Inches(5.6))
    caption("图1 温度变化趋势图")
    doc.add_picture(os.path.join(OUT_DIR, "chart_histogram.png"), width=Inches(5.6))
    caption("图2 温度分布直方图")
    doc.add_picture(os.path.join(OUT_DIR, "schematic.png"), width=Inches(5.6))
    caption("图3 桥梁CAD示意图")

    heading("四、结论")
    add("本周期平均温度24.3℃，最高温度26.0℃，最低温度22.5℃，"
        "桥梁结构整体处于安全状态，最大挠度2.3mm未超限值。"
        "下阶段持续关注温度变化趋势与挠度发展。")

    doc.save(OUT_DOCX)
    print(f"样例报告已生成: {OUT_DOCX}")


if __name__ == "__main__":
    build()
