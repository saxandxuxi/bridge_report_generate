#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Check run structure of J-table paragraphs to find the cross-run issue."""

import os
import sys
import re

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from report_agent.recognizer import recognize, CELL_REF_RE
from report_agent.report_builder import iter_block_items
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = r"D:\五座桥数据分析文件\洞庭湖大桥.docx"


def check_runs(p, idx, j_paras, h_paras, m_paras):
    """Check run structure of a paragraph."""
    runs = p.runs
    full = "".join(r.text for r in runs)
    
    if idx in j_paras:
        table = "J"
    elif idx in h_paras:
        table = "H"
    elif idx in m_paras:
        table = "M"
    else:
        return
    
    matches = list(CELL_REF_RE.finditer(full))
    offsets = []
    cur = 0
    for r in runs:
        offsets.append(cur)
        cur += len(r.text)
    
    for m in matches:
        start, end = m.start(), m.end()
        ri = max(i for i, off in enumerate(offsets) if off <= start)
        run_end = offsets[ri] + len(runs[ri].text)
        spans = end > run_end
        
        if spans:
            print(f"  [{table}] para {idx}: '{m.group(0)}' SPANS RUNS!")
            print(f"    runs: {[r.text for r in runs]}")
            print(f"    offsets: {offsets}")
            print(f"    match: start={start}, end={end}")
            print(f"    run[{ri}]: '{runs[ri].text}', run_end={run_end}")
            print(f"    -> end({end}) > run_end({run_end}) = {spans}")
        else:
            if table == "J":
                print(f"  [{table}] para {idx}: '{m.group(0)}' OK (run[{ri}]='{runs[ri].text}', {len(runs)} runs)")
            elif table in ("H", "M"):
                first_para = min(h_paras) if table == "H" else min(m_paras)
                if idx == first_para:
                    print(f"  [{table}] para {idx}: '{m.group(0)}' OK (run[{ri}]='{runs[ri].text}', {len(runs)} runs)")


if __name__ == "__main__":
    print("Running recognize (no LLM)...")
    analysis = recognize(SRC, llm_cfg=None)

    j_replace = [ct for ct in analysis.get("chart_texts", [])
                 if ct.get("source") == "cell_ref" and ct.get("table_letter") == "J" and ct.get("verdict") == "replace"]
    h_replace = [ct for ct in analysis.get("chart_texts", [])
                 if ct.get("source") == "cell_ref" and ct.get("table_letter") == "H" and ct.get("verdict") == "replace"]
    m_replace = [ct for ct in analysis.get("chart_texts", [])
                 if ct.get("source") == "cell_ref" and ct.get("table_letter") == "M" and ct.get("verdict") == "replace"]
    j_paras = set(ct["paragraph"] for ct in j_replace)
    h_paras = set(ct["paragraph"] for ct in h_replace)
    m_paras = set(ct["paragraph"] for ct in m_replace)

    print(f"J paras: {len(j_paras)}, H paras: {len(h_paras)}, M paras: {len(m_paras)}")

    doc = Document(SRC)
    idx = 0
    all_paras = h_paras | m_paras | j_paras

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            if idx in all_paras:
                check_runs(item, idx, j_paras, h_paras, m_paras)
            idx += 1
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if idx in all_paras:
                            check_runs(p, idx, j_paras, h_paras, m_paras)
                        idx += 1

    for section in doc.sections:
        for p in section.header.paragraphs:
            idx += 1
        for p in section.footer.paragraphs:
            idx += 1

    print("\nDone.")
