#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""构建报告模板：templates/temperature_report_template.docx

模板中的动态内容通过占位符标记：
  {{stats.<指标>}}   统计值       {{date.<字段>}}   报告期日期
  {{chart.<ID>}}     图表位置     {{rows.<数据集>}} 表格可重复行
  {{col.<字段>}}     重复行内字段
  {{?condition}}...{{?}}           条件渲染块
  {{key|default:默认值}}           带默认值的占位符

支持通过命令行参数自定义报告标题、指标名称等，不再局限于温度场景：
  python build_template.py --title "湿度分析报告" --metric 湿度 --metric-en humidity --metric-unit "%%"
  python build_template.py --title "沉降监测报告" --metric 沉降 --metric-en settlement --metric-unit mm --no-humidity

设计遵循 standard_business_brief 预设：Calibri/微软雅黑、1 英寸页边距、
H1 #2E74B5 16pt、表格 9360 DXA 固定宽度、页眉页脚 0.492 英寸。
"""

import argparse
import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "temperature_report_template.docx")

LATIN_FONT = "Calibri"
EAST_FONT = "微软雅黑"

HEADING_BLUE = "#2E74B5"
HEADING_DARK = "#1F4D78"
MUTED_GRAY = "#595959"
LIGHT_FILL = "#F2F4F7"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_FONT)


def add_para(doc, text="", style=None, align=None, size=11, bold=False,
             color=None, before=None, after=None, line=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    return p


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = LATIN_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), LATIN_FONT)
        rfonts.set(qn("w:eastAsia"), EAST_FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def setup_page(doc: Document, title: str = "温度数据分析报告") -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run(title), size=9, color="#808080")

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run("第 "), size=9, color="#808080")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # 9pt
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    fp._p.append(fld)
    set_run(fp.add_run(" 页 · 生成时间 {{date.generated_at}}"), size=9, color="#808080")


def _ordered_insert(parent, tag: str) -> OxmlElement:
    """按 OOXML 规定顺序插入 tblPr 子元素。"""
    existing = parent.find(qn(tag))
    if existing is not None:
        return existing
    child = OxmlElement(tag)
    order = [
        "w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
        "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
        "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
        "w:tblLayout", "w:tblCellMar", "w:tblLook",
    ]
    idx = order.index(tag) if tag in order else len(order)
    insert_at = len(parent)
    for i, el in enumerate(parent):
        if el.tag in order and order.index(el.tag) > idx:
            insert_at = i
            break
    parent.insert(insert_at, child)
    return child


def set_table_geometry(table, widths_dxa, indent_dxa=120,
                       margins=(80, 80, 120, 120)) -> None:
    """设置表格固定宽度、缩进与单元格边距（单位 DXA）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    total = sum(widths_dxa)
    tblW = _ordered_insert(tblPr, "w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")

    tblInd = _ordered_insert(tblPr, "w:tblInd")
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")

    layout = _ordered_insert(tblPr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    cell_mar = _ordered_insert(tblPr, "w:tblCellMar")
    for tag, val in zip(("w:top", "w:start", "w:bottom", "w:end"), margins):
        el = cell_mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            cell_mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in grid.findall(qn("w:gridCol")):
            grid.remove(gc)
        for w in widths_dxa:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            grid.append(gc)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_dxa):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_header_row_repeat(table) -> None:
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell(cell, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             size=11, color=None, fill=None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)


def add_table(doc, rows, cols, widths_dxa, header_texts, body_rows,
              align_map=None):
    """创建带表头的表格并填充内容。"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = doc.styles["Table Grid"]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    align_map = align_map or {}

    for j, h in enumerate(header_texts):
        set_cell(table.rows[0].cells[j], h, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, fill=LIGHT_FILL)
    set_header_row_repeat(table)

    for i, row_vals in enumerate(body_rows, start=1):
        for j, val in enumerate(row_vals):
            align = align_map.get(j, WD_ALIGN_PARAGRAPH.LEFT)
            set_cell(table.rows[i].cells[j], val, align=align)

    set_table_geometry(table, widths_dxa)
    return table


def caption(doc, text: str) -> None:
    add_para(doc, text, size=10, color=MUTED_GRAY, before=4, after=4)


def build(
    title: str = "温度数据分析报告",
    metric: str = "温度",
    metric_en: str = "temperature",
    metric_unit: str = "℃",
    include_humidity: bool = True,
    thresholds: list = None,
    output_path: str = None,
) -> str:
    """构建报告模板。

    参数:
        title: 报告标题（如"温度数据分析报告"）
        metric: 指标中文名（如"温度"）
        metric_en: 指标英文名（用于占位符，如"temperature"）
        metric_unit: 指标单位（如"℃"）
        include_humidity: 是否包含湿度段落
        thresholds: 阈值列表（默认 [30, 35]）
        output_path: 输出路径（默认 templates/<metric_en>_report_template.docx）

    返回:
        生成的模板文件路径
    """
    thresholds = thresholds or [30, 35]
    if output_path is None:
        output_path = os.path.join(
            BASE_DIR, "templates", f"{metric_en}_report_template.docx"
        )

    doc = Document()
    setup_styles(doc)
    setup_page(doc, title=title)
    doc.core_properties.title = f"{title}模板"

    # ---- 标题区 ----
    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=20, bold=True, color=HEADING_DARK, after=2)
    add_para(doc,
             "数据区间：{{date.period_start}} 至 {{date.period_end}}　|　"
             "生成时间：{{date.generated_at}}",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=MUTED_GRAY, after=14)

    # ---- 一、报告摘要 ----
    doc.add_heading("一、报告摘要", level=1)
    summary_parts = [
        f"本报告基于 {{{{date.period_start}}}} 至 {{{{date.period_end}}}} 共 ",
        f"{{{{stats.days}}}} 天的监测数据，对{metric}进行统计分析。",
    ]
    add_para(doc, "".join(summary_parts))
    add_para(
        doc,
        f"本周期最高{metric} {{{{stats.{metric_en}.max}}}}{metric_unit}，"
        f"最低{metric} {{{{stats.{metric_en}.min}}}}{metric_unit}，"
        f"平均{metric} {{{{stats.{metric_en}.avg}}}}{metric_unit}，"
        f"{metric}标准差 {{{{stats.{metric_en}.std}}}}{metric_unit}。",
    )
    # 阈值段落
    threshold_parts = []
    for t in thresholds:
        threshold_parts.append(
            f"超过 {t}{metric_unit} 的天数为 {{{{stats.{metric_en}.days_above_{t}}}}} 天，"
        )
    if threshold_parts:
        add_para(doc, "期间" + "".join(threshold_parts)[:-1] + "。", after=10)

    # ---- 二、关键指标统计 ----
    doc.add_heading("二、关键指标统计", level=1)
    caption(doc, f"表1 关键统计指标（{{{{date.period_start}}}} ~ {{{{date.period_end}}}}）")
    stat_rows = [
        [f"最高{metric}", f"{{{{stats.{metric_en}.max}}}} {metric_unit}（{{{{stats.{metric_en}.max_date}}}}）"],
        [f"最低{metric}", f"{{{{stats.{metric_en}.min}}}} {metric_unit}（{{{{stats.{metric_en}.min_date}}}}）"],
        [f"平均{metric}", f"{{{{stats.{metric_en}.avg}}}} {metric_unit}"],
        [f"中位{metric}", f"{{{{stats.{metric_en}.median}}}} {metric_unit}"],
        [f"{metric}标准差", f"{{{{stats.{metric_en}.std}}}} {metric_unit}"],
        [f"{metric}极差", f"{{{{stats.{metric_en}.range}}}} {metric_unit}"],
    ]
    if include_humidity:
        stat_rows.append(["平均湿度", "{{stats.humidity.avg}} %"])
    for t in thresholds:
        stat_rows.append([f">{t}{metric_unit} 天数", f"{{{{stats.{metric_en}.days_above_{t}}}}} 天"])
    stat_rows.append(["数据天数", "{{stats.days}} 天"])

    add_table(
        doc,
        rows=len(stat_rows) + 1,
        cols=2,
        widths_dxa=[3168, 6192],
        header_texts=["指标", "数值"],
        body_rows=stat_rows,
        align_map={0: WD_ALIGN_PARAGRAPH.LEFT, 1: WD_ALIGN_PARAGRAPH.LEFT},
    )

    # ---- 三、图表分析 ----
    doc.add_heading("三、图表分析", level=1)

    doc.add_heading(f"3.1 {metric}变化趋势", level=2)
    add_para(doc, "{{chart.trend}}")
    caption(doc, f"图1 {metric}逐日变化趋势（{{{{date.period_start}}}} ~ {{{{date.period_end}}}}）")

    doc.add_heading(f"3.2 {metric}分布", level=2)
    add_para(doc, "{{chart.histogram}}")
    caption(doc, f"图2 {metric}分布直方图")

    doc.add_heading(f"3.3 每日{metric}对比", level=2)
    add_para(doc, "{{chart.daily_bars}}")
    caption(doc, f"图3 每日{metric}对比柱状图")

    doc.add_heading(f"3.4 {metric}分布特征", level=2)
    add_para(doc, "{{chart.boxplot}}")
    caption(doc, f"图4 {metric}分布箱线图")

    # ---- 四、逐日数据明细 ----
    doc.add_heading("四、逐日数据明细", level=1)
    caption(doc, f"表2 逐日{metric}明细（偏差为相对本周期平均{metric}的偏差）")
    table2 = doc.add_table(rows=2, cols=3)
    table2.style = doc.styles["Table Grid"]
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(["日期", f"{metric}（{metric_unit}）", f"偏差（{metric_unit}）"]):
        set_cell(table2.rows[0].cells[j], h, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, fill=LIGHT_FILL)
    set_header_row_repeat(table2)
    set_cell(table2.rows[1].cells[0], "{{rows.daily_records}}",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table2.rows[1].cells[1], f"{{{{col.{metric_en}:0.1f}}}}",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table2.rows[1].cells[2], "{{col.deviation:+0.1f}}",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table2, [3168, 3024, 3168])

    # ---- 五、结论与建议 ----
    doc.add_heading("五、结论与建议", level=1)
    add_para(
        doc,
        f"总体来看，本周期平均{metric}为 {{{{stats.{metric_en}.avg}}}}{metric_unit}，"
        f"最高 {{{{stats.{metric_en}.max}}}}{metric_unit}"
        f"（{{{{stats.{metric_en}.max_date}}}}），"
        f"最低 {{{{stats.{metric_en}.min}}}}{metric_unit}"
        f"（{{{{stats.{metric_en}.min_date}}}}）。"
        f"{metric}标准差为 {{{{stats.{metric_en}.std}}}}{metric_unit}，"
        f"反映出本周期{metric}的波动情况。",
    )
    t0 = thresholds[0] if thresholds else 30
    add_para(
        doc,
        f"建议持续关注高温（>{t0}{metric_unit}）天数的变化趋势，必要时启动相应措施；"
        f"加强对{metric}异常的监测与预警。",
        after=0,
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"模板已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="构建数据分析报告模板")
    parser.add_argument("--title", default="温度数据分析报告", help="报告标题")
    parser.add_argument("--metric", default="温度", help="指标中文名（如 温度/湿度/沉降）")
    parser.add_argument("--metric-en", default="temperature", help="指标英文名（用于占位符）")
    parser.add_argument("--metric-unit", default="℃", help="指标单位（如 ℃/mm/%%）")
    parser.add_argument("--no-humidity", action="store_true", help="不包含湿度段落")
    parser.add_argument("--thresholds", type=int, nargs="+", default=[30, 35],
                        help="阈值列表（空格分隔）")
    parser.add_argument("--output", default=None, help="输出路径（默认 templates/<metric_en>_report_template.docx）")
    args = parser.parse_args()

    build(
        title=args.title,
        metric=args.metric,
        metric_en=args.metric_en,
        metric_unit=args.metric_unit,
        include_humidity=not args.no_humidity,
        thresholds=args.thresholds,
        output_path=args.output,
    )
