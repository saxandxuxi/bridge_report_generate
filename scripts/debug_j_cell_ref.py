#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Debug script: trace why J-table (wind_speed) cell_refs are not replaced."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from report_agent.recognizer import parse_docx, CELL_REF_RE
from report_agent.report_builder import iter_block_items

SRC = r"D:\五座桥数据分析文件\洞庭湖大桥.docx"


def _check_para(p, idx, targets):
    """Check if cell_ref regex matches in run-concatenated text. Returns match count."""
    if not targets:
        return 0
    runs = p.runs
    if not runs:
        print(f"  para {idx}: NO RUNS, targets={targets}")
        return 0
    full = "".join(r.text for r in runs)
    parse_text = paragraphs[idx] if idx < len(paragraphs) else "???"
    matches = list(CELL_REF_RE.finditer(full))
    target_positions = set(targets.keys())
    match_positions = set(m.start() for m in matches)
    matched = target_positions & match_positions
    unmatched = target_positions - match_positions

    if full != parse_text:
        print(f"  para {idx}: TEXT MISMATCH!")
        print(f"    parse_text: '{parse_text[:80]}'")
        print(f"    run_concat: '{full[:80]}'")
    
    if matched and not unmatched:
        print(f"  para {idx}: OK - {len(matched)} matches, full='{full[:60]}'")
        return len(matched)
    if unmatched:
        print(f"  para {idx}: PROBLEM - {len(unmatched)}/{len(targets)} unmatched")
        print(f"    full: '{full[:80]}'")
        print(f"    targets: {targets}")
        print(f"    match_positions: {match_positions}")
        print(f"    target_positions: {target_positions}")
        for u in sorted(unmatched):
            if u < len(full):
                print(f"    text at pos {u}: '{full[u:u+10]}'")
            else:
                print(f"    pos {u} beyond text length {len(full)}")
    return len(matched)


if __name__ == "__main__":
    print("=" * 70)
    print("Step 1: parse_docx -> get chart_texts")
    print("=" * 70)
    parsed = parse_docx(SRC)
    paragraphs = parsed["texts"]
    chart_texts = parsed.get("chart_texts", [])

    # Filter cell_refs by table letter
    cell_refs_by_letter = {}
    for ct in chart_texts:
        if ct.get("source") != "cell_ref":
            continue
        letter = ct.get("table_letter", "?")
        cell_refs_by_letter.setdefault(letter, []).append(ct)

    print(f"\nTotal chart_texts: {len(chart_texts)}")
    print(f"Cell_refs by table letter:")
    for letter in sorted(cell_refs_by_letter.keys()):
        items = cell_refs_by_letter[letter]
        replace_count = sum(1 for c in items if c.get("verdict") == "replace")
        keep_count = sum(1 for c in items if c.get("verdict") == "keep")
        print(f"  {letter}: {len(items)} total ({replace_count} replace, {keep_count} keep)")

    # Focus on J-table
    j_refs = cell_refs_by_letter.get("J", [])
    j_replace = [c for c in j_refs if c.get("verdict") == "replace"]
    j_keep = [c for c in j_refs if c.get("verdict") == "keep"]
    print(f"\n--- J-table cell_refs: {len(j_replace)} replace, {len(j_keep)} keep ---")
    for ct in j_replace[:5]:
        print(f"  [replace] para={ct['paragraph']}, pos={ct.get('position', 0)}, "
              f"text='{ct['text']}', row={ct['row']}, col={ct['col']}, "
              f"title='{ct.get('table_title', '')[:40]}'")
    for ct in j_keep[:3]:
        print(f"  [keep] para={ct['paragraph']}, pos={ct.get('position', 0)}, "
              f"text='{ct['text']}', row={ct['row']}, col={ct['col']}, "
              f"title='{ct.get('table_title', '')[:40]}'")

    # Also check H and M for comparison
    for letter in ("H", "M"):
        refs = cell_refs_by_letter.get(letter, [])
        rep = [c for c in refs if c.get("verdict") == "replace"]
        keep = [c for c in refs if c.get("verdict") == "keep"]
        print(f"\n--- {letter}-table: {len(rep)} replace, {len(keep)} keep ---")
        if rep:
            print(f"  first replace: para={rep[0]['paragraph']}, title='{rep[0].get('table_title', '')[:40]}'")
        if keep:
            print(f"  first keep: para={keep[0]['paragraph']}, title='{keep[0].get('table_title', '')[:40]}'")

    print("\n" + "=" * 70)
    print("Step 2: Simulate annotate_docx iteration - check idx alignment")
    print("=" * 70)

    # Build cell_ref_targets as annotate_docx would
    cell_ref_targets = {}
    for ct in j_replace:
        cell_ref_targets.setdefault(ct["paragraph"], {})[ct.get("position", 0)] = ct["text"]

    j_para_indices = set(cell_ref_targets.keys())
    print(f"\nJ-table replace paragraph indices: {sorted(j_para_indices)}")
    print(f"Total cell_ref_targets entries: {len(cell_ref_targets)}")

    # Simulate iteration
    doc = Document(SRC)
    idx = 0
    found_j_paras = set()
    idx_text_map = {}

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            idx_text_map[idx] = ("body", item.text.strip()[:60])
            if idx in j_para_indices:
                found_j_paras.add(idx)
            idx += 1
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        idx_text_map[idx] = ("table", p.text.strip()[:60])
                        if idx in j_para_indices:
                            found_j_paras.add(idx)
                        idx += 1

    for section in doc.sections:
        for p in section.header.paragraphs:
            idx_text_map[idx] = ("header", p.text.strip()[:60])
            if idx in j_para_indices:
                found_j_paras.add(idx)
            idx += 1
        for p in section.footer.paragraphs:
            idx_text_map[idx] = ("footer", p.text.strip()[:60])
            if idx in j_para_indices:
                found_j_paras.add(idx)
            idx += 1

    print(f"Total paragraphs iterated in annotate: {idx}")
    print(f"Total paragraphs from parse_docx: {len(paragraphs)}")
    print(f"J paragraphs found in annotate iteration: {len(found_j_paras)}/{len(j_para_indices)}")

    missing = j_para_indices - found_j_paras
    if missing:
        print(f"\n*** MISSING J paragraphs: {sorted(missing)} ***")
        for mp in sorted(missing):
            if mp < len(paragraphs):
                print(f"  para {mp} (from parse): '{paragraphs[mp][:60]}'")
            if mp in idx_text_map:
                print(f"  para {mp} (from annotate): {idx_text_map[mp]}")
    else:
        print("  All J paragraphs found!")

    # Show what's at those indices in the annotate iteration
    print(f"\nJ paragraph details in annotate iteration:")
    for pi in sorted(found_j_paras):
        location, text = idx_text_map.get(pi, ("?", ""))
        parse_text = paragraphs[pi][:60] if pi < len(paragraphs) else "???"
        match_str = "OK" if text == parse_text else "MISMATCH"
        print(f"  idx={pi} [{location}]: annotate='{text}' | parse='{parse_text}' [{match_str}]")

    print("\n" + "=" * 70)
    print("Step 3: Check run-concatenated text vs p.text for J paragraphs")
    print("=" * 70)

    doc3 = Document(SRC)
    idx3 = 0
    j_match_count = 0
    j_check_count = 0

    for item in iter_block_items(doc3):
        if isinstance(item, Paragraph):
            if idx3 in j_para_indices:
                j_match_count += _check_para(item, idx3, cell_ref_targets.get(idx3, {}))
                j_check_count += 1
            idx3 += 1
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if idx3 in j_para_indices:
                            j_match_count += _check_para(p, idx3, cell_ref_targets.get(idx3, {}))
                            j_check_count += 1
                        idx3 += 1

    for section in doc3.sections:
        for p in section.header.paragraphs:
            idx3 += 1
        for p in section.footer.paragraphs:
            idx3 += 1

    print(f"\nTotal J paragraphs checked: {j_check_count}")
    print(f"Total cell_refs matched: {j_match_count}")

    print("\n" + "=" * 70)
    print("Step 4: Compare paragraph count between parse and annotate")
    print("=" * 70)

    parse_count = len(paragraphs)
    annotate_count = idx3
    print(f"parse_docx paragraph count: {parse_count}")
    print(f"annotate_docx paragraph count: {annotate_count}")
    if parse_count != annotate_count:
        print("*** MISMATCH! Paragraph counts differ! ***")
        print("  This means idx in annotate_docx doesn't align with paragraph index in parse_docx")
    else:
        print("  Counts match - indices should be aligned")

    print("\nDone.")
