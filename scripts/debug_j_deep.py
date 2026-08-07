#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Deep debug: trace annotate_docx cell_ref_targets building for J-table."""

import os
import sys
import re

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from report_agent.recognizer import recognize, CELL_REF_RE
from report_agent.report_builder import iter_block_items
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = r"D:\五座桥数据分析文件\洞庭湖大桥.docx"

print("Running recognize (no LLM)...")
analysis = recognize(SRC, llm_cfg=None)

# Get J-table cell_refs
j_replace = [ct for ct in analysis.get("chart_texts", [])
             if ct.get("source") == "cell_ref" and ct.get("table_letter") == "J" and ct.get("verdict") == "replace"]
j_paras = set(ct["paragraph"] for ct in j_replace)
print(f"J-table replace cell_refs: {len(j_replace)}, paragraphs: {sorted(j_paras)}")

# === Simulate cell_ref_targets building ===
cell_ref_columns_map = {
    "H": {2: "avg", 3: "max", 4: "min", 5: "abs_max", 6: "rms", 7: "range"},
    "M": {2: "avg", 3: "max", 4: "min", 5: "abs_max", 6: "rms", 7: "range"},
    "J": {2: "avg", 3: "max", 4: "min", 5: "range", 6: "rms"},
}
cable_clamp_row_to_col = {2: "h_87L_1", 3: "h_87R_1", 4: "h_88L_1", 5: "h_88R_1"}
cable_clamp_row2_to_col = {6: "h_87L_2", 7: "h_87R_2", 8: "h_88L_2", 9: "h_88R_2"}
rotation_row_to_col = {2: "m_junshan_x", 3: "m_junshan_y", 4: "m_yueyang_x", 5: "m_yueyang_y"}
wind_row_to_col = {2: "j_junshan_top", 3: "j_yueyang_top", 4: "j_half_upstream", 5: "j_half_downstream"}
metric_point_map = {
    "cable_clamp": {**cable_clamp_row_to_col, **cable_clamp_row2_to_col},
    "rotation": rotation_row_to_col,
    "wind_speed": wind_row_to_col,
}

cell_ref_targets = {}
chart_text_targets = {}
chart_text_counter = {}

for ct in analysis.get("chart_texts", []):
    if ct.get("source") == "cell_ref":
        if ct.get("verdict") == "keep":
            continue
        metric = ct.get("metric", "")
        row = ct.get("row", 0)
        col = ct.get("col", 0)
        table_letter = ct.get("table_letter", "")
        row_map = metric_point_map.get(metric, {})
        column = row_map.get(row)
        if column is None:
            column = f"unknown_{metric}_r{row}"
        stat_map = cell_ref_columns_map.get(table_letter, {})
        stat = stat_map.get(col, f"col{col}")
        marker = f"{{{{cell.{metric}.{column}.{stat}}}}}"
        cell_ref_targets.setdefault(ct["paragraph"], {})[ct.get("position", 0)] = marker
    else:
        metric = ct.get("metric", "chart")
        cid = ct.get("chart_id", "trend")
        chart_text_counter[metric] = chart_text_counter.get(metric, 0) + 1
        unique_cid = f"{metric}_{cid}_{chart_text_counter[metric]}"
        chart_text_targets[ct["paragraph"]] = f"{{{{chart.{unique_cid}}}}}"

print(f"\ncell_ref_targets entries: {len(cell_ref_targets)}")
print(f"chart_text_targets entries: {len(chart_text_targets)}")

# Check J paragraphs in cell_ref_targets
print(f"\nJ paragraphs in cell_ref_targets:")
for pi in sorted(j_paras):
    in_crt = pi in cell_ref_targets
    in_ctt = pi in chart_text_targets
    print(f"  para {pi}: cell_ref_targets={in_crt}, chart_text_targets={in_ctt}")
    if in_crt:
        print(f"    cell_ref_targets[{pi}] = {cell_ref_targets[pi]}")
    if in_ctt:
        print(f"    chart_text_targets[{pi}] = {chart_text_targets[pi]}")

# Check overlap
overlap = j_paras & set(chart_text_targets.keys())
if overlap:
    print(f"\n*** OVERLAP! J paras also in chart_text_targets: {sorted(overlap)} ***")
    for pi in sorted(overlap):
        print(f"  para {pi}: chart_text_targets = '{chart_text_targets[pi]}'")
        # Find which chart_text caused this
        for ct in analysis.get("chart_texts", []):
            if ct.get("paragraph") == pi and ct.get("source") != "cell_ref":
                print(f"    chart_text: source={ct['source']}, text='{ct.get('text', '')[:60]}'")
else:
    print(f"\nNo overlap between J paras and chart_text_targets.")

# Now simulate process_paragraph for J paras
print("\n" + "=" * 60)
print("Simulating process_paragraph for J paras")
print("=" * 60)

doc = Document(SRC)
idx = 0
for item in iter_block_items(doc):
    if isinstance(item, Paragraph):
        if idx in j_paras:
            _process_j_para(item, idx, cell_ref_targets.get(idx, {}), chart_text_targets.get(idx))
        idx += 1
    elif isinstance(item, Table):
        for row in item.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if idx in j_paras:
                        _process_j_para(p, idx, cell_ref_targets.get(idx, {}), chart_text_targets.get(idx))
                    idx += 1

for section in doc.sections:
    for p in section.header.paragraphs:
        idx += 1
    for p in section.footer.paragraphs:
        idx += 1


def _process_j_para(p, idx, crt, ctt):
    """Simulate process_paragraph for a J paragraph."""
    print(f"\n  para {idx}: text='{p.text[:60]}'")
    print(f"    in chart_text_targets: {bool(ctt)}")
    print(f"    in cell_ref_targets: {bool(crt)}")
    
    if ctt:
        print(f"    -> Would be replaced as chart_text: '{ctt}'")
        print(f"    -> *** THIS IS THE BUG! chart_text_targets takes priority and returns early ***")
        return
    
    if crt:
        runs = p.runs
        if not runs:
            print(f"    -> NO RUNS")
            return
        full = "".join(r.text for r in runs)
        matches = list(CELL_REF_RE.finditer(full))
        target_positions = set(crt.keys())
        match_positions = set(m.start() for m in matches)
        matched = target_positions & match_positions
        print(f"    full='{full[:60]}', matches={[m.group(0) for m in matches]}")
        print(f"    targets={crt}")
        print(f"    matched={len(matched)}/{len(target_positions)}")
        
        # Actually do the replacement
        if matched:
            offsets = []
            cur = 0
            for r in runs:
                offsets.append(cur)
                cur += len(r.text)
            ok = 0
            for m in matches:
                if m.start() in crt:
                    ri = max(i for i, off in enumerate(offsets) if off <= m.start())
                    run_end = offsets[ri] + len(runs[ri].text)
                    if m.end() > run_end:
                        print(f"      SKIP: match spans multiple runs")
                        continue
                    rel = m.start() - offsets[ri]
                    marker = crt[m.start()]
                    runs[ri].text = runs[ri].text[:rel] + marker + runs[ri].text[rel + (m.end() - m.start()):]
                    ok += 1
                    print(f"      REPLACED: {m.group(0)} -> {marker}")
            print(f"    -> Replaced {ok} cell_refs")
